import csv
import json
import os
import platform
import re
import threading
import time
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox, scrolledtext, ttk
from typing import Dict, List, Optional, Tuple
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml import parse_xml


class ConfigManager:
    """Gerencia o carregamento e salvamento da configuração de campos"""
    
    DEFAULT_CONFIG = [
        {"label": "Campo1:", "key": "campo1"},
        {"label": "Campo2:", "key": "campo2"},
        {"label": "Campo3:", "key": "campo3"},
        {"label": "Campo4:", "key": "campo4"},
        {"label": "Campo5:", "key": "campo5"},
        {"label": "Campo6:", "key": "campo6"}
    ]

    def __init__(self, config_file: str = 'config_campos.json'):
        self.config_file = Path(config_file)

    def load_config(self) -> List[Dict]:
        """Carrega a configuração do arquivo JSON ou cria uma padrão"""
        try:
            if self.config_file.exists():
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    print(f"✅ Configuração carregada de '{self.config_file}'")
                    return config
            else:
                return self._create_default_config()
        except Exception as e:
            print(f"⚠️ Erro ao carregar configuração: {e}")
            return self.DEFAULT_CONFIG

    def _create_default_config(self) -> List[Dict]:
        """Cria arquivo de configuração padrão"""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.DEFAULT_CONFIG, f, indent=4, ensure_ascii=False)
            print(f"ℹ️ Arquivo '{self.config_file}' criado com configuração padrão")
            return self.DEFAULT_CONFIG
        except Exception as e:
            print(f"❌ Erro ao criar configuração padrão: {e}")
            return self.DEFAULT_CONFIG


class CSVReader:
    """Responsável pela leitura de arquivos CSV"""
    
    ENCODINGS = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252', 'windows-1252']

    @staticmethod
    def read_csv(file_path: str) -> Tuple[Optional[List[str]], Optional[List[str]], Optional[pd.DataFrame]]:
        """Lê um arquivo CSV e retorna a lista de nomes, colunas e o DataFrame completo"""
        try:
            return CSVReader._read_with_pandas(file_path)
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao ler o CSV: {e}")
            return None, None, None

    @staticmethod
    def _read_with_pandas(file_path: str) -> Tuple[Optional[List[str]], Optional[List[str]], Optional[pd.DataFrame]]:
        """Tenta ler o CSV usando pandas e retorna nomes, colunas e DataFrame"""
        for encoding in CSVReader.ENCODINGS:
            try:
                df = pd.read_csv(file_path, encoding=encoding, engine='python', 
                               on_bad_lines='skip')
                if 'Nome' in df.columns:
                    nomes = df['Nome'].dropna().str.strip()
                    nomes_list = nomes[nomes != ''].tolist()
                    colunas = [col for col in df.columns if col != 'Nome']
                    return nomes_list, colunas, df
                else:
                    # Se não encontrar coluna "Nome", usa a primeira coluna
                    primeira_coluna = df.columns[0]
                    nomes = df[primeira_coluna].dropna().str.strip()
                    nomes_list = nomes[nomes != ''].tolist()
                    colunas = [col for col in df.columns if col != primeira_coluna]
                    return nomes_list, colunas, df
            except Exception:
                continue
        return None, None, None

    @staticmethod
    def get_csv_columns(file_path: str) -> Optional[List[str]]:
        """Obtém apenas as colunas do arquivo CSV"""
        try:
            for encoding in CSVReader.ENCODINGS:
                try:
                    df = pd.read_csv(file_path, encoding=encoding, engine='python', 
                                   on_bad_lines='skip', nrows=1)
                    return df.columns.tolist()
                except Exception:
                    continue
            return None
        except Exception as e:
            print(f"Erro ao obter colunas do CSV: {e}")
            return None


class ColumnSelectionDialog:
    """Diálogo para seleção de colunas do CSV"""
    
    def __init__(self, parent, colunas: List[str]):
        self.parent = parent
        self.colunas = colunas
        self.selected_columns = []
        
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Selecionar Colunas do CSV")
        self.dialog.geometry("500x400")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        self._setup_ui()
        
    def _setup_ui(self):
        main_frame = ttk.Frame(self.dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="Selecione as colunas que deseja incluir no documento:", 
                 font=("Arial", 10, "bold")).pack(anchor=tk.W, pady=(0, 10))
        
        # Frame para a lista de colunas com scrollbar
        list_frame = ttk.Frame(main_frame)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Listbox com múltipla seleção
        self.listbox = tk.Listbox(list_frame, selectmode=tk.MULTIPLE, yscrollcommand=scrollbar.set)
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scrollbar.config(command=self.listbox.yview)
        
        # Preencher listbox com colunas
        for coluna in self.colunas:
            self.listbox.insert(tk.END, coluna)
        
        # Frame para botões
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)
        
        ttk.Button(button_frame, text="Selecionar Todas", 
                  command=self._select_all).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(button_frame, text="Desmarcar Todas", 
                  command=self._deselect_all).pack(side=tk.LEFT, padx=(0, 5))
        
        ttk.Button(button_frame, text="Confirmar", 
                  command=self._confirm).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(button_frame, text="Cancelar", 
                  command=self._cancel).pack(side=tk.RIGHT)
    
    def _select_all(self):
        """Seleciona todas as colunas"""
        self.listbox.select_set(0, tk.END)
    
    def _deselect_all(self):
        """Deseleciona todas as colunas"""
        self.listbox.select_clear(0, tk.END)
    
    def _confirm(self):
        """Confirma a seleção"""
        selections = self.listbox.curselection()
        self.selected_columns = [self.colunas[i] for i in selections]
        self.dialog.destroy()
    
    def _cancel(self):
        """Cancela a seleção"""
        self.selected_columns = []
        self.dialog.destroy()
    
    def show(self) -> List[str]:
        """Mostra o diálogo e retorna as colunas selecionadas"""
        self.parent.wait_window(self.dialog)
        return self.selected_columns


class DocumentProcessor:
    """Processa e gera documentos Word baseados em templates"""
    
    @staticmethod
    def clean_filename(filename: str, max_length: int = 100) -> str:
        """Limpa o nome do arquivo removendo caracteres inválidos"""
        cleaned = re.sub(r'[<>:"/\\|?*]', '_', filename)
        cleaned = cleaned.strip()[:max_length]
        return cleaned or "caso_teste"

    @staticmethod
    def fill_template(doc: Document, data: Dict[str, str], field_config: List[Dict], 
                     colunas_selecionadas: List[str] = None, dados_csv: Dict[str, str] = None) -> None:
        """Preenche o template com os dados fornecidos - AGORA ADICIONA APÓS CONTEÚDO EXISTENTE"""
        
        # ADICIONAR ESPAÇO REDUZIDO APÓS O TÍTULO
        espaco_apos_titulo = doc.add_paragraph()
        espaco_apos_titulo.paragraph_format.space_after = Pt(1)
        
       
        # Adicionar título da seção de dados
        titulo = doc.add_heading('Dados do Teste', level=1)
        for run in titulo.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  
        
        # Adicionar campos da configuração
        for campo_info in field_config:
            key = campo_info['key']
            label = campo_info['label'].rstrip(':')
            value = data.get(key, '')
            
            # Adicionar parágrafo com campo e valor
            campo_para = doc.add_paragraph()
            campo_run = campo_para.add_run(f"{label}: ")
            campo_run.bold = True
            campo_run.font.name = 'Arial'
            campo_run.font.size = Pt(12)
            campo_run.font.color.rgb = RGBColor(0, 0, 0)
            
            valor_run = campo_para.add_run(value)
            valor_run.bold = False
            valor_run.font.name = 'Arial'
            valor_run.font.size = Pt(12)
            valor_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Adicionar caso de teste
        caso_para = doc.add_paragraph()
        caso_run = caso_para.add_run("Caso de Teste: ")
        caso_run.bold = True
        caso_run.font.name = 'Arial'
        caso_run.font.size = Pt(12)
        caso_run.font.color.rgb = RGBColor(0, 0, 0)
        
        nome_run = caso_para.add_run(data.get('Caso de Teste', ''))
        nome_run.bold = False
        nome_run.font.name = 'Arial'
        nome_run.font.size = Pt(12)
        nome_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Adicionar tabela com dados do CSV se houver colunas selecionadas
        if colunas_selecionadas and dados_csv:
            DocumentProcessor._adicionar_tabela_csv(doc, colunas_selecionadas, dados_csv)

    @staticmethod
    def _adicionar_tabela_csv(doc: Document, colunas_selecionadas: List[str], dados_csv: Dict[str, str]) -> None:
        """Adiciona uma tabela com os dados do CSV ao documento"""
        if not colunas_selecionadas or not dados_csv:
            return
        
        # Adicionar um espaço antes da tabela
        doc.add_paragraph()
        
        # Adicionar título da tabela
        titulo = doc.add_heading('Dados Adicionais do Caso de Teste', level=2)
        
        for run in titulo.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Criar tabela
        tabela = doc.add_table(rows=len(colunas_selecionadas), cols=2)
        
        # Adicionar dados
        for i, coluna in enumerate(colunas_selecionadas):
            if coluna in dados_csv:
                row_cells = tabela.rows[i].cells
                
                # Primeira coluna: nome do campo em negrito
                row_cells[0].text = coluna
                for paragraph in row_cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Segunda coluna: valor sem negrito
                row_cells[1].text = str(dados_csv[coluna])
                for paragraph in row_cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        run.bold = False
                        run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Adicionar bordas pretas às células
                for cell in row_cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                          r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                          r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                          r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                          r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                          r'</w:tcBorders>')
                    tcPr.append(tcBorders)


class DefaultDocumentGenerator:
    """Gera documentos padrão quando nenhum template é fornecido"""
    
    @staticmethod
    def create_default_document(data: Dict[str, str], field_config: List[Dict], 
                               colunas_selecionadas: List[str] = None, 
                               dados_csv: Dict[str, str] = None) -> Document:
        """Cria um documento padrão com estrutura organizada - AGORA DINÂMICO BASEADO NA CONFIGURAÇÃO"""
        doc = Document()
        
        # Configurar estilos de fonte padrão
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        
        # Título do documento
        title = doc.add_heading('Evidências de Teste - Documentação', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Seção de informações do teste
        info_heading = doc.add_heading('Informações do Teste', level=1)
        for run in info_heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Adicionar campos dinamicamente baseados na configuração
        for campo_info in field_config:
            key = campo_info['key']
            label = campo_info['label'].rstrip(':')
            value = data.get(key, 'Não informado')
            
            campo_para = doc.add_paragraph()
            label_run = campo_para.add_run(f"{label}: ")
            label_run.bold = True
            label_run.font.name = 'Arial'
            label_run.font.size = Pt(12)
            label_run.font.color.rgb = RGBColor(0, 0, 0)
            
            value_run = campo_para.add_run(value)
            value_run.bold = False
            value_run.font.name = 'Arial'
            value_run.font.size = Pt(12)
            value_run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
        
        # Seção do caso de teste
        caso_para = doc.add_paragraph()
        caso_run = caso_para.add_run('Nome do Caso de Teste: ')
        caso_run.bold = True
        caso_run.font.name = 'Arial'
        caso_run.font.size = Pt(12)
        caso_run.font.color.rgb = RGBColor(0, 0, 0)
        
        nome_run = caso_para.add_run(data.get('Caso de Teste', 'Não informado'))
        nome_run.font.name = 'Arial'
        nome_run.font.size = Pt(12)
        nome_run.bold = False
        nome_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Adicionar tabela com dados do CSV se houver colunas selecionadas
        if colunas_selecionadas and dados_csv:
            DocumentProcessor._adicionar_tabela_csv(doc, colunas_selecionadas, dados_csv)
        
        doc.add_paragraph()
        
        # Seções fixas adicionais (mantidas do original)
        DefaultDocumentGenerator._add_standard_sections(doc)
        
        return doc

    @staticmethod
    def _add_standard_sections(doc: Document) -> None:
        """Adiciona seções padrão ao documento"""
        # Seção de descrição
        desc_heading = doc.add_heading('Descrição do Teste', level=2)
        for run in desc_heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        desc_para = doc.add_paragraph(
            "Esta seção deve conter a descrição detalhada do caso de teste executado, "
            "incluindo pré-condições, passos de execução e resultados esperados."
        )
        
        # Seção de evidências
        evid_heading = doc.add_heading('Evidências Coletadas', level=2)
        for run in evid_heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        evid_para = doc.add_paragraph("Registro das evidências coletadas durante a execução do teste:")
        
        # Tabela para evidências
        evidencias_table = doc.add_table(rows=6, cols=3)
        evidencias_table.style = 'Light Grid Accent 1'
        
        # Cabeçalho da tabela de evidências
        evidencias_header = evidencias_table.rows[0].cells
        headers = ['Etapa', 'Evidência', 'Resultado']
        for col, header in enumerate(headers):
            evidencias_header[col].text = header
            for paragraph in evidencias_header[col].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Linhas para preenchimento
        etapas = [
            'Pré-condições',
            'Configuração Inicial', 
            'Execução do Teste',
            'Pós-condições',
            'Resultado Final'
        ]
        
        for row, etapa in enumerate(etapas, 1):
            if row < len(evidencias_table.rows):
                row_cells = evidencias_table.rows[row].cells
                row_cells[0].text = etapa
                row_cells[1].text = "[Descreva a evidência coletada]"
                row_cells[2].text = "[Resultado obtido - OK/Erro]"
                
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            run.bold = False
                            run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
        
        # Seção de observações
        obs_heading = doc.add_heading('Observações e Comentários', level=2)
        for run in obs_heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        obs_para = doc.add_paragraph("Adicione observações relevantes sobre a execução do teste:")
        
        # Área para observações
        obs_list_para = doc.add_paragraph()
        obs_title_run = obs_list_para.add_run("Observações Gerais:\n")
        obs_title_run.bold = True
        obs_title_run.font.name = 'Arial'
        obs_title_run.font.size = Pt(12)
        obs_title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        obs_items = [
            "• [Insira observações sobre problemas encontrados]\n",
            "• [Comentários sobre o comportamento do sistema]\n",
            "• [Sugestões de melhorias]\n",
            "• [Outras informações relevantes]"
        ]
        
        for item in obs_items:
            item_run = obs_list_para.add_run(item)
            item_run.font.name = 'Arial'
            item_run.font.size = Pt(12)
            item_run.bold = False
            item_run.font.color.rgb = RGBColor(0, 0, 0)


class TemplateGeneratorModule:
    """Interface principal da aplicação"""
    
    def __init__(self, parent, settings=None):
        self.parent = parent
        self.settings = settings or {}
        self.window = None
        
        self.config_manager = ConfigManager()
        self.csv_reader = CSVReader()
        self.doc_processor = DocumentProcessor()
        self.default_doc_generator = DefaultDocumentGenerator()
        
        self.campos_config = self.config_manager.load_config()
        self.campos_entries: Dict[str, tk.Entry] = {}
        
        # Controle das colunas do CSV
        self.colunas_selecionadas: List[str] = []
        self.df_csv: Optional[pd.DataFrame] = None
        self.csv_colunas: List[str] = []
        
        # Variável para controle de diretório automático
        self.auto_directory_var = tk.BooleanVar(value=True)
    
    def show(self):
        """Mostra interface completa"""
        if self.window and self.window.winfo_exists():
            self.window.lift()
            return
        
        self.window = tk.Toplevel(self.parent)
        self.window.title("PrintF - Gerador de Templates")
        self.window.geometry("1000x800")
        self.window.minsize(900, 700)
        
        # Configurar cor de fundo
        try:
            from modules.styles import LiquidGlassStyle
            self.window.configure(bg=LiquidGlassStyle.BG_PRIMARY)
        except ImportError:
            self.window.configure(bg='#f0f0f0')
        
        self._create_complete_ui()

    def hide(self):
        """Esconde o módulo"""
        if self.window and self.window.winfo_exists():
            self.window.destroy()
        self.window = None

    def _create_complete_ui(self):
        """Cria a interface completa do módulo"""
        # Frame principal
        main_frame = ttk.Frame(self.window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        self._configure_grid_weights()
        self._create_title_section(main_frame)
        self._create_dynamic_fields_section(main_frame)
        self._create_file_section(main_frame)
        self._create_control_buttons(main_frame)
        self._create_progress_section(main_frame)
        self._create_log_section(main_frame)
        
        self._set_default_template()

    def _configure_grid_weights(self) -> None:
        """Configura os pesos do grid para redimensionamento"""
        self.window.columnconfigure(0, weight=1)
        self.window.rowconfigure(0, weight=1)

    def _create_title_section(self, parent) -> None:
        """Cria a seção do título"""
        titulo = ttk.Label(parent, text="📄 PrintF - Gerador de Templates", 
                          font=("Arial", 16, "bold"))
        titulo.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        ttk.Separator(parent, orient='horizontal').grid(
            row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))

    def _create_dynamic_fields_section(self, parent) -> None:
        """Cria os campos dinâmicos baseados na configuração - AGORA TOTALMENTE DINÂMICO"""
        ttk.Label(parent, text="Dados dos Testes:", 
                 font=("Arial", 12, "bold")).grid(
                     row=2, column=0, columnspan=3, sticky=tk.W, pady=(0, 10))

        # Criar campos dinamicamente baseados na configuração
        for i, campo_info in enumerate(self.campos_config):
            self._create_field_row(parent, campo_info, i)

    def _create_field_row(self, parent, campo_info: Dict, row_index: int) -> None:
        """Cria uma linha de campo na interface"""
        label_text = campo_info['label']
        campo_key = campo_info['key']
        
        ttk.Label(parent, text=label_text).grid(
            row=3 + row_index, column=0, sticky=tk.W, pady=2)
        
        entry = ttk.Entry(parent, width=40)
        entry.grid(row=3 + row_index, column=1, columnspan=2, 
                  sticky=(tk.W, tk.E), pady=2, padx=(5, 0))
        
        self.campos_entries[campo_key] = entry

    def _create_file_section(self, parent) -> None:
        """Cria a seção de seleção de arquivos"""
        next_row = 3 + len(self.campos_config)
        
        ttk.Separator(parent, orient='horizontal').grid(
            row=next_row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)

        ttk.Label(parent, text="Arquivos:", 
                 font=("Arial", 12, "bold")).grid(
                     row=next_row + 1, column=0, columnspan=3, sticky=tk.W, pady=(0, 10))

        # Campos de arquivo
        self.csv_entry = self._create_file_field(parent, "CSV:*", next_row + 2, self.selecionar_csv)
        self.template_entry = self._create_file_field(parent, "Template (Opcional):", next_row + 3, self.selecionar_template)
        
        # Opção para diretório automático
        directory_frame = ttk.Frame(parent)
        directory_frame.grid(row=next_row + 4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=2)
        
        ttk.Label(directory_frame, text="Pasta Saída:").pack(side=tk.LEFT)
        
        self.pasta_entry = ttk.Entry(directory_frame, width=40)
        self.pasta_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 10))
        
        ttk.Button(directory_frame, text="Procurar", command=self.selecionar_pasta).pack(side=tk.LEFT)
        
        # Checkbox para diretório automático
        self.auto_dir_check = ttk.Checkbutton(
            directory_frame, 
            text="Criar automaticamente com nome do template", 
            variable=self.auto_directory_var,
            command=self._toggle_auto_directory
        )
        self.auto_dir_check.pack(side=tk.LEFT, padx=(10, 0))

        # Info sobre campos obrigatórios
        info_label = ttk.Label(parent, text="* Campos obrigatórios", font=("Arial", 9), foreground="gray")
        info_label.grid(row=next_row + 5, column=0, columnspan=3, sticky=tk.W, pady=(5, 0))

    def _toggle_auto_directory(self):
        """Ativa/desativa campo de pasta quando usar diretório automático"""
        if self.auto_directory_var.get():
            self.pasta_entry.config(state='disabled')
            # Preencher automaticamente se tiver template
            template_path = self.template_entry.get().strip()
            if template_path and Path(template_path).exists():
                self._update_auto_directory()
        else:
            self.pasta_entry.config(state='normal')

    def _update_auto_directory(self):
        """Atualiza o diretório automático baseado no template"""
        if not self.auto_directory_var.get():
            return
            
        template_path = self.template_entry.get().strip()
        if template_path and Path(template_path).exists():
            template_name = Path(template_path).stem
            auto_dir = f"evidencias_{template_name}"
            self.pasta_entry.delete(0, tk.END)
            self.pasta_entry.insert(0, auto_dir)

    def _create_file_field(self, parent, label: str, row: int, command) -> ttk.Entry:
        """Cria um campo de seleção de arquivo"""
        ttk.Label(parent, text=label).grid(row=row, column=0, sticky=tk.W, pady=2)
        entry = ttk.Entry(parent, width=40)
        entry.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=2, padx=(5, 0))
        ttk.Button(parent, text="Procurar", command=command).grid(
            row=row, column=2, padx=(5, 0))
        return entry

    def _create_control_buttons(self, parent) -> None:
        """Cria os botões de controle"""
        next_row = 3 + len(self.campos_config) + 6
        
        ttk.Separator(parent, orient='horizontal').grid(
            row=next_row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)

        button_frame = ttk.Frame(parent)
        button_frame.grid(row=next_row + 1, column=0, columnspan=3, pady=10)
        
        self.gerar_btn = ttk.Button(button_frame, text="▶️ Gerar Documentos", 
                                   command=self.iniciar_processamento)
        self.gerar_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(button_frame, text="🔄 Limpar", command=self.limpar_campos).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="❌ Sair", command=self.hide).pack(side=tk.LEFT, padx=5)

    def _create_progress_section(self, parent) -> None:
        """Cria a seção de progresso"""
        next_row = 3 + len(self.campos_config) + 8
        
        ttk.Separator(parent, orient='horizontal').grid(
            row=next_row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)

        ttk.Label(parent, text="Progresso:").grid(
            row=next_row + 1, column=0, columnspan=3, sticky=tk.W, pady=(0, 5))
        
        self.progress = ttk.Progressbar(parent, mode='determinate')
        self.progress.grid(row=next_row + 2, column=0, columnspan=3, 
                          sticky=(tk.W, tk.E), pady=(0, 10))

    def _create_log_section(self, parent) -> None:
        """Cria a seção de log"""
        next_row = 3 + len(self.campos_config) + 10
        
        ttk.Label(parent, text="Log de Execução:").grid(
            row=next_row, column=0, columnspan=3, sticky=tk.W, pady=(0, 5))
        
        self.log_text = scrolledtext.ScrolledText(parent, width=70, height=15, state='disabled')
        self.log_text.grid(row=next_row + 1, column=0, columnspan=3, 
                          sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        
        parent.rowconfigure(next_row + 1, weight=1)

    def _set_default_template(self) -> None:
        """Preenche template padrão se existir - AGORA CRIA AUTOMATICAMENTE SE NÃO EXISTIR"""
        template_path = 'template_evidencias.docx'
        if not Path(template_path).exists():
            # Cria template automaticamente se não existir
            self.log("📝 Criando template padrão automaticamente...")
            if self._criar_template_exemplo_automatico():
                self.template_entry.insert(0, template_path)
                self.log("✅ Template padrão criado com sucesso!")
                # Atualizar diretório automático se estiver ativo
                if self.auto_directory_var.get():
                    self._update_auto_directory()
        else:
            self.template_entry.insert(0, template_path)
            if self.auto_directory_var.get():
                self._update_auto_directory()

    def _criar_template_exemplo_automatico(self) -> bool:
        """Cria template de exemplo automaticamente (sem interação do usuário)"""
        try:
            doc = Document()
            
            # Configurar estilo normal para Arial 12
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)
            font.color.rgb = RGBColor(0, 0, 0)
            
            # Título principal
            main_title = doc.add_heading('Evidências de Teste', level=1)
            for run in main_title.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(16)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            
            doc.add_paragraph()
            
            # Conteúdo de exemplo do template original
            content_para = doc.add_paragraph("Este é o template padrão para documentação de testes. ")
            content_para.add_run("Os dados específicos de cada teste serão adicionados após esta seção.").bold = True
            
            doc.add_paragraph()
            doc.add_paragraph("Instruções:")
            instructions = doc.add_paragraph()
            instructions.add_run("• Preencha os dados necessários no aplicativo\n")
            instructions.add_run("• Selecione o arquivo CSV com os casos de teste\n")
            instructions.add_run("• Os dados serão automaticamente inseridos após este conteúdo")
            
            doc.save('template_evidencias.docx')
            return True
        except Exception as e:
            self.log(f"⚠️ Aviso: Não foi possível criar template automático: {e}")
            return False

    # Métodos de seleção de arquivos
    def selecionar_csv(self) -> None:
        arquivo = filedialog.askopenfilename(title="Selecionar arquivo CSV", 
                                            filetypes=[("CSV Files", "*.csv"), ("Todos os arquivos", "*.*")])
        if arquivo:
            self.csv_entry.delete(0, tk.END)
            self.csv_entry.insert(0, arquivo)
            
            # Verificar se o CSV tem colunas além do Nome
            self._verificar_colunas_csv(arquivo)

    def _verificar_colunas_csv(self, arquivo_csv: str) -> None:
        """Verifica se o CSV tem colunas adicionais e pergunta ao usuário se deseja selecioná-las"""
        try:
            # Obter colunas do CSV
            colunas = self.csv_reader.get_csv_columns(arquivo_csv)
            
            if colunas and len(colunas) > 1:  # Tem colunas além do Nome (ou primeira coluna)
                # Remover a coluna de nome (assumindo que é a primeira)
                coluna_nome = colunas[0]
                colunas_adicionais = colunas[1:] if len(colunas) > 1 else []
                
                if colunas_adicionais:
                    self.log(f"📊 CSV possui {len(colunas_adicionais)} colunas adicionais")
                    
                    # Perguntar ao usuário se deseja selecionar colunas
                    resposta = messagebox.askyesno(
                        "Colunas Adicionais Encontradas", 
                        f"O arquivo CSV possui {len(colunas_adicionais)} colunas adicionais:\n\n" +
                        "\n".join(f"• {coluna}" for coluna in colunas_adicionais) +
                        "\n\nDeseja selecionar quais colunas incluir nos documentos?"
                    )
                    
                    if resposta:
                        # Mostrar diálogo de seleção de colunas
                        dialog = ColumnSelectionDialog(self.window, colunas_adicionais)
                        self.colunas_selecionadas = dialog.show()
                        
                        if self.colunas_selecionadas:
                            self.log(f"✅ Colunas selecionadas: {', '.join(self.colunas_selecionadas)}")
                        else:
                            self.log("ℹ️ Nenhuma coluna adicional selecionada")
                    
                    # Carregar o DataFrame completo para uso posterior
                    _, _, self.df_csv = self.csv_reader.read_csv(arquivo_csv)
            else:
                self.log("ℹ️ CSV não possui colunas adicionais para seleção")
                self.colunas_selecionadas = []
                
        except Exception as e:
            self.log(f"⚠️ Erro ao verificar colunas do CSV: {e}")
            self.colunas_selecionadas = []

    def selecionar_template(self) -> None:
        arquivo = filedialog.askopenfilename(
            title="Selecionar template DOCX", 
            filetypes=[("Word Documents", "*.docx"), ("Todos os arquivos", "*.*")]
        )
        if arquivo:
            self.template_entry.delete(0, tk.END)
            self.template_entry.insert(0, arquivo)
            # Atualizar diretório automático se estiver ativo
            if self.auto_directory_var.get():
                self._update_auto_directory()

    def selecionar_pasta(self) -> None:
        pasta = filedialog.askdirectory(title="Selecionar pasta de saída")
        if pasta:
            self.pasta_entry.delete(0, tk.END)
            self.pasta_entry.insert(0, pasta)
            # Desmarcar auto diretório se usuário selecionar manualmente
            self.auto_directory_var.set(False)
            self.pasta_entry.config(state='normal')

    def limpar_campos(self) -> None:
        """Limpa todos os campos da interface"""
        for entry in [self.csv_entry, self.template_entry, self.pasta_entry]:
            entry.delete(0, tk.END)
        
        for entry in self.campos_entries.values():
            entry.delete(0, tk.END)
        
        self._clear_log()
        self.progress['value'] = 0
        
        # Limpar seleção de colunas
        self.colunas_selecionadas = []
        self.df_csv = None
        
        # Restaurar auto directory
        self.auto_directory_var.set(True)
        self.pasta_entry.config(state='disabled')
        self._set_default_template()

    def _clear_log(self) -> None:
        """Limpa o log de execução"""
        self.log_text.config(state='normal')
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state='disabled')

    def log(self, mensagem: str) -> None:
        """Adiciona mensagem ao log"""
        self.log_text.config(state='normal')
        self.log_text.insert(tk.END, mensagem + "\n")
        self.log_text.see(tk.END)
        self.log_text.config(state='disabled')
        self.window.update()

    def _get_fixed_data(self) -> Dict[str, str]:
        """Obtém os dados dos campos fixos"""
        dados = {}
        for campo_key, entry in self.campos_entries.items():
            valor = entry.get().strip()
            if not valor:
                raise ValueError(f"O campo '{campo_key}' é obrigatório!")
            dados[campo_key] = valor
        return dados

    def _validate_inputs(self, csv_path: str) -> bool:
        """Valida os inputs necessários (apenas CSV é obrigatório)"""
        if not csv_path:
            messagebox.showerror("Erro", "Selecione um arquivo CSV!")
            return False
        
        if not Path(csv_path).exists():
            messagebox.showerror("Erro", "Arquivo CSV não encontrado!")
            return False
        
        return True

    def _obter_dados_csv_por_nome(self, nome_caso_teste: str) -> Dict[str, str]:
        """Obtém os dados do CSV para um caso de teste específico"""
        if self.df_csv is None or not self.colunas_selecionadas:
            return {}
        
        try:
            # Encontrar a linha correspondente ao nome do caso de teste
            # Assumindo que a primeira coluna contém os nomes
            coluna_nome = self.df_csv.columns[0]
            linha = self.df_csv[self.df_csv[coluna_nome] == nome_caso_teste]
            
            if not linha.empty:
                dados = {}
                for coluna in self.colunas_selecionadas:
                    if coluna in linha.columns:
                        valor = linha[coluna].iloc[0]
                        # Converter para string e tratar valores NaN
                        if pd.isna(valor):
                            dados[coluna] = ""
                        else:
                            dados[coluna] = str(valor)
                    else:
                        dados[coluna] = ""
                return dados
        
        except Exception as e:
            self.log(f"⚠️ Erro ao obter dados do CSV para '{nome_caso_teste}': {e}")
        
        return {}

    def _get_output_directory(self, template_path: str) -> str:
        """Determina o diretório de saída baseado nas configurações"""
        if self.auto_directory_var.get() and template_path:
            # Usar nome do template para criar diretório automático
            template_name = Path(template_path).stem
            return f"evidencias_{template_name}"
        else:
            # Usar diretório especificado pelo usuário ou padrão
            output_folder = self.pasta_entry.get().strip()
            return output_folder or 'evidencias_geradas'

    def processar_documentos(self) -> None:
        """Processa os documentos em lote - versão robusta que nunca falha"""
        try:
            # Obter dados fixos
            try:
                dados_fixos = self._get_fixed_data()
            except ValueError as e:
                messagebox.showerror("Erro", str(e))
                self.gerar_btn.config(state='normal')
                return
            
            csv_path = self.csv_entry.get()
            template_path = self.template_entry.get().strip()
            
            if not self._validate_inputs(csv_path):
                self.gerar_btn.config(state='normal')
                return

            # Determinar pasta de saída
            output_folder = self._get_output_directory(template_path)
            
            # Garantir que temos um template válido
            template_path = self._garantir_template_valido(template_path)
            
            # Criar pasta de saída
            try:
                Path(output_folder).mkdir(exist_ok=True)
                self.log(f"📁 Pasta de saída: {output_folder}")
            except Exception as e:
                self.log(f"⚠️ Aviso: Não foi possível criar a pasta '{output_folder}': {e}")
                self.log("📁 Usando pasta atual para salvar os documentos...")
                output_folder = '.'
            
            self.log("📖 Lendo arquivo CSV...")
            casos_teste, colunas_csv, self.df_csv = self.csv_reader.read_csv(csv_path)
            
            if not casos_teste:
                messagebox.showerror("Erro", "Não foi possível ler os casos de teste do CSV")
                self.gerar_btn.config(state='normal')
                return
            
            # Se não tivermos colunas selecionadas mas o CSV tiver colunas adicionais, 
            # perguntar novamente (pode acontecer se o usuário cancelou anteriormente)
            if not self.colunas_selecionadas and colunas_csv and len(colunas_csv) > 0:
                self._verificar_colunas_csv(csv_path)
            
            # Determinar modo de operação
            use_default_template = True
            if template_path and Path(template_path).exists():
                use_default_template = False
                self.log("📁 Usando template personalizado...")
            else:
                self.log("📝 Gerando documentos com template padrão...")
            
            self._process_test_cases(casos_teste, dados_fixos, template_path, output_folder, use_default_template)
            
        except Exception as e:
            self.log(f"❌ Erro inesperado: {e}")
            messagebox.showerror("Erro", f"Erro inesperado: {e}")
            self.gerar_btn.config(state='normal')

    def _garantir_template_valido(self, template_path: str) -> str:
        """Garante que temos um template válido, criando automaticamente se necessário"""
        if not template_path or not Path(template_path).exists():
            self.log("📝 Nenhum template válido encontrado, criando automaticamente...")
            if self._criar_template_exemplo_automatico():
                new_template_path = 'template_evidencias.docx'
                self.template_entry.delete(0, tk.END)
                self.template_entry.insert(0, new_template_path)
                self.log("✅ Template padrão criado e configurado automaticamente")
                return new_template_path
            else:
                self.log("⚠️ Não foi possível criar template, usando geração padrão...")
                return ""
        return template_path

    def _process_test_cases(self, casos_teste: List[str], dados_fixos: Dict[str, str], 
                           template_path: str, output_folder: str, use_default_template: bool) -> None:
        """Processa cada caso de teste individualmente"""
        self.log(f"📊 Encontrados {len(casos_teste)} casos de teste\n")
        self.progress['maximum'] = len(casos_teste)
        
        sucessos = 0
        erros = []
        arquivos_gerados = set()
        
        for i, caso_teste in enumerate(casos_teste, 1):
            try:
                self.progress['value'] = i
                self.log(f"🔄 Processando: {caso_teste}")
                
                # Obter dados do CSV para este caso de teste
                dados_csv = self._obter_dados_csv_por_nome(caso_teste)
                
                if self._generate_single_document(caso_teste, dados_fixos, template_path, 
                                                output_folder, arquivos_gerados, use_default_template,
                                                dados_csv):
                    sucessos += 1
                else:
                    erros.append((caso_teste, "Erro na geração"))
                    
            except Exception as e:
                self.log(f"❌ Erro no caso '{caso_teste}': {e}\n")
                erros.append((caso_teste, str(e)))
            
            time.sleep(0.05)  # Pequena pausa para não sobrecarregar
            
        self._show_final_results(sucessos, len(erros), len(casos_teste), 
                               output_folder, arquivos_gerados)
        self.gerar_btn.config(state='normal')

    def _generate_single_document(self, caso_teste: str, dados_fixos: Dict[str, str], 
                                template_path: str, output_folder: str, 
                                arquivos_gerados: set, use_default_template: bool,
                                dados_csv: Dict[str, str] = None) -> bool:
        """Gera um único documento - AGORA PRESERVA TEMPLATE ORIGINAL E ADICIONA DADOS APÓS"""
        try:
            dados_completos = dados_fixos.copy()
            dados_completos['Caso de Teste'] = caso_teste
            
            # Usar template se fornecido e existir, caso contrário criar documento padrão
            if not use_default_template:
                try:
                    doc = Document(template_path)
                    # AGORA: Adiciona dados APÓS o conteúdo do template original
                    self.doc_processor.fill_template(doc, dados_completos, self.campos_config,
                                                   self.colunas_selecionadas, dados_csv)
                except Exception as e:
                    self.log(f"⚠️ Erro ao usar template personalizado: {e}. Usando template padrão...")
                    doc = self.default_doc_generator.create_default_document(
                        dados_completos, self.campos_config, self.colunas_selecionadas, dados_csv)
            else:
                # Criar documento padrão com todos os dados
                doc = self.default_doc_generator.create_default_document(
                    dados_completos, self.campos_config, self.colunas_selecionadas, dados_csv)
            
            # Nome do arquivo usa apenas o nome do Caso de Teste
            nome_base = self.doc_processor.clean_filename(caso_teste)
            nome_arquivo = self._generate_unique_filename(f"{nome_base}.docx", arquivos_gerados)
            
            caminho_completo = Path(output_folder) / nome_arquivo
            
            # Tentar salvar o documento
            try:
                doc.save(caminho_completo)
                self.log(f"✅ Salvo: {nome_arquivo}")
                return True
            except Exception as e:
                # Fallback: tentar salvar com nome diferente
                try:
                    nome_alternativo = f"Evidencia_{datetime.now().strftime('%H%M%S')}.docx"
                    caminho_alternativo = Path(output_folder) / nome_alternativo
                    doc.save(caminho_alternativo)
                    self.log(f"✅ Salvo (nome alternativo): {nome_alternativo}")
                    arquivos_gerados.add(nome_alternativo)
                    return True
                except Exception as e2:
                    self.log(f"❌ Erro ao salvar documento: {e2}")
                    return False
            
        except Exception as e:
            self.log(f"❌ Erro crítico ao gerar documento: {e}")
            return False

    def _generate_unique_filename(self, filename: str, existing_files: set) -> str:
        """Gera um nome de arquivo único"""
        contador = 1
        nome_original = filename
        
        while filename in existing_files:
            base, ext = os.path.splitext(nome_original)
            filename = f"{base}_{contador}{ext}"
            contador += 1
        
        existing_files.add(filename)
        return filename

    def _abrir_pasta(self, caminho_pasta: str) -> None:
        """Abre a pasta no explorador de arquivos do sistema operacional"""
        try:
            caminho_absoluto = Path(caminho_pasta).absolute()
            
            if platform.system() == "Windows":
                os.startfile(caminho_absoluto)
            elif platform.system() == "Darwin":  # macOS
                os.system(f'open "{caminho_absoluto}"')
            else:  # Linux e outros
                os.system(f'xdg-open "{caminho_absoluto}"')
                
            self.log(f"📁 Pasta aberta: {caminho_absoluto}")
        except Exception as e:
            self.log(f"⚠️ Não foi possível abrir a pasta automaticamente: {e}")
            self.log(f"📁 Pasta manual: {caminho_absoluto}")

    def _show_final_results(self, sucessos: int, erros: int, total: int, 
                           output_folder: str, arquivos_gerados: set) -> None:
        """Exibe os resultados finais do processamento"""
        self.log("\n" + "=" * 50)
        self.log("🎉 PROCESSO CONCLUÍDO!")
        self.log("=" * 50)
        self.log(f"📊 Total processado: {total}")
        self.log(f"✅ Sucessos: {sucessos}")
        self.log(f"❌ Erros: {erros}")
        self.log(f"📁 Pasta: {Path(output_folder).absolute()}")
        
        if self.colunas_selecionadas:
            self.log(f"📋 Colunas incluídas: {', '.join(self.colunas_selecionadas)}")
        
        if sucessos > 0:
            self.log(f"\n📋 Arquivos gerados:")
            for arquivo in sorted(arquivos_gerados)[:10]:  # Mostra apenas os 10 primeiros
                self.log(f"• {arquivo}")
            if len(arquivos_gerados) > 10:
                self.log(f"• ... e mais {len(arquivos_gerados) - 10} arquivos")
        
        # Abrir pasta automaticamente quando há sucessos
        if sucessos > 0:
            self.log("\n📁 Abrindo pasta de resultados...")
            self._abrir_pasta(output_folder)
            
            messagebox.showinfo("Sucesso", 
                              f"✅ {sucessos} documentos gerados com sucesso!\n" +
                              (f"📋 Colunas incluídas: {', '.join(self.colunas_selecionadas)}\n" if self.colunas_selecionadas else "") +
                              f"📁 Pasta aberta automaticamente: {Path(output_folder).absolute()}")
        else:
            messagebox.showwarning("Concluído com avisos", 
                                 f"Processo concluído com {erros} erro(s).\n"
                                 f"✅ {sucessos} documentos gerados com sucesso.\n" +
                                 (f"📋 Colunas incluídas: {', '.join(self.colunas_selecionadas)}\n" if self.colunas_selecionadas else "") +
                                 f"📁 Pasta: {Path(output_folder).absolute()}")

    def iniciar_processamento(self) -> None:
        """Inicia o processamento em thread separada"""
        self.gerar_btn.config(state='disabled')
        self._clear_log()
        self.progress['value'] = 0
        
        thread = threading.Thread(target=self.processar_documentos, daemon=True)
        thread.start()


# Função de compatibilidade para manter a interface antiga
def create_template_generator(parent, settings=None):
    """Função de fábrica para criar o módulo"""
    return TemplateGeneratorModule(parent, settings)


# Teste local do módulo
if __name__ == "__main__":
    root = tk.Tk()
    root.title("Teste Template Generator")
    root.geometry("800x600")
    
    app = TemplateGeneratorModule(root)
    app.show()
    
    root.mainloop()