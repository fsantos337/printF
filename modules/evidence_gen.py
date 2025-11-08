import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, colorchooser, ttk
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pyautogui
from pynput import mouse, keyboard
from PIL import Image, ImageTk, ImageDraw, ImageFont, ImageFilter
from datetime import datetime
import math
import re
import glob
import json
import uuid
import shutil
import subprocess

# Importar sistema de estilos
try:
    from modules.styles import LiquidGlassStyle
    STYLES_AVAILABLE = True
except ImportError:
    try:
        from styles import LiquidGlassStyle
        STYLES_AVAILABLE = True
    except ImportError:
        STYLES_AVAILABLE = False
        print("⚠️ Estilos Liquid Glass não disponíveis, usando fallback")

class EvidenceGeneratorModule:
    """Módulo completo de geração de documentos de evidências"""
    
    def __init__(self, parent=None, settings=None):
        self.parent = parent
        self.root = None
        self.settings = settings or {}
        self.output_dir = os.getcwd()
        self.prints = []
        self.doc = None
        self.using_template = False
        self.template_path = None
        self.current_index = 0
        self.evidence_dir = None
        self.metadata_path = None
        self.metadata = {"evidencias": [], "proximo_id": 1}
        
        self.gravando = False
        self.listener_mouse = None
        self.listener_keyboard = None
        self.popup = None
        self.processamento_cancelado = False
        self.saved_file_path = None
        
        # Atributos para o editor
        self.elements = []
        self.undo_stack = []
        self.temp_element = None
        self.original_img = None
        self.editing_img = None
        self.display_img = None
        self.current_tk_img = None
        self.canvas = None
        self.canvas_img = None
        self.scale_factor = 1.0
        self.comment_entry = None
        self.current_img_label = None
        self.pos_label = None
        self.color_chooser_window = None
        
        self.using_liquid_glass = STYLES_AVAILABLE and self.settings.get('theme', 'liquid_glass') == 'liquid_glass'
        self.style_manager = LiquidGlassStyle if STYLES_AVAILABLE and self.using_liquid_glass else None
        
        print(f"🎨 EvidenceGenerator - Liquid Glass: {self.using_liquid_glass}")

    def _apply_styles(self, window):
        """Aplica estilos à janela"""
        if self.using_liquid_glass and self.style_manager:
            try:
                self.style_manager.apply_window_style(window)
                return True
            except Exception as e:
                print(f"⚠️ Erro ao aplicar estilos: {e}")
                window.configure(bg='#f5f5f5')
                return False
        else:
            window.configure(bg='#f5f5f5')
            return True

    def _create_styled_frame(self, parent, **kwargs):
        """Cria frame com estilos aplicados"""
        if self.using_liquid_glass and self.style_manager:
            try:
                return self.style_manager.create_glass_frame(parent, **kwargs)
            except Exception as e:
                print(f"⚠️ Erro ao criar frame: {e}")
                return tk.Frame(parent, bg='#f5f5f5', **kwargs)
        else:
            return tk.Frame(parent, bg='#f5f5f5', **kwargs)

    def _create_styled_button(self, parent, text, command, style_type="glass", **kwargs):
        """Cria botão com estilos aplicados"""
        if self.using_liquid_glass and self.style_manager:
            try:
                if style_type == "accent":
                    return self.style_manager.create_accent_button(parent, text, command, **kwargs)
                else:
                    return self.style_manager.create_glass_button(parent, text, command, **kwargs)
            except Exception as e:
                print(f"⚠️ Erro ao criar botão: {e}")
                return self._create_fallback_button(parent, text, command, style_type, **kwargs)
        else:
            return self._create_fallback_button(parent, text, command, style_type, **kwargs)

    def _create_fallback_button(self, parent, text, command, style_type="glass", **kwargs):
        """Cria botão fallback"""
        btn = tk.Button(parent, text=text, command=command, 
                      bg='#3498db' if style_type == "accent" else '#ecf0f1',
                      fg='white' if style_type == "accent" else '#2c3e50',
                      font=("Arial", 10, "bold" if style_type == "accent" else "normal"),
                      relief="flat",
                      cursor="hand2",
                      **kwargs)
        
        if style_type == "accent":
            btn.bind("<Enter>", lambda e: btn.config(bg='#2980b9'))
            btn.bind("<Leave>", lambda e: btn.config(bg='#3498db'))
        else:
            btn.bind("<Enter>", lambda e: btn.config(bg='#d5dbdb'))
            btn.bind("<Leave>", lambda e: btn.config(bg='#ecf0f1'))
        
        return btn

    def _create_styled_label(self, parent, text, style_type="glass", **kwargs):
        """Cria label com estilos aplicados"""
        if self.using_liquid_glass and self.style_manager:
            try:
                if style_type == "title":
                    return self.style_manager.create_title_label(parent, text, **kwargs)
                else:
                    label = ttk.Label(parent, text=text, style="Glass.TLabel", **kwargs)
                    return label
            except Exception as e:
                print(f"⚠️ Erro ao criar label: {e}")
                return self._create_fallback_label(parent, text, style_type, **kwargs)
        else:
            return self._create_fallback_label(parent, text, style_type, **kwargs)

    def _create_fallback_label(self, parent, text, style_type="glass", **kwargs):
        """Cria label fallback"""
        bg_color = '#f5f5f5'
        font_config = ("Arial", 14, "bold") if style_type == "title" else ("Arial", 10)
        return tk.Label(parent, text=text, bg=bg_color, fg='#2c3e50', 
                      font=font_config, **kwargs)

    def _create_styled_entry(self, parent, **kwargs):
        """Cria entry com estilos aplicados"""
        if self.using_liquid_glass and self.style_manager:
            try:
                return self.style_manager.create_glass_entry(parent, **kwargs)
            except Exception as e:
                print(f"⚠️ Erro ao criar entry: {e}")
                return tk.Entry(parent, bg='white', fg='#2c3e50', 
                              relief="solid", bd=1, **kwargs)
        else:
            return tk.Entry(parent, bg='white', fg='#2c3e50', 
                          relief="solid", bd=1, **kwargs)

    def _salvar_metadata(self):
        """Salva os metadados no arquivo JSON"""
        if self.metadata_path:
            with open(self.metadata_path, 'w', encoding='utf-8') as f:
                json.dump(self.metadata, f, indent=2, ensure_ascii=False)

    def carregar_evidencias(self, dir_path):
        """Carrega as evidências baseadas nos metadados - SUPORTA MÚLTIPLOS FORMATOS"""
        self.metadata_path = os.path.join(dir_path, "evidencias_metadata.json")
        
        FORMATOS_SUPORTADOS = ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.tif']
        
        if os.path.exists(self.metadata_path):
            try:
                with open(self.metadata_path, 'r', encoding='utf-8') as f:
                    self.metadata = json.load(f)
            except:
                self.metadata = {"evidencias": [], "proximo_id": 1}
        else:
            self.metadata = {"evidencias": [], "proximo_id": 1}
            
            for arquivo in os.listdir(dir_path):
                _, ext = os.path.splitext(arquivo)
                if ext.lower() in FORMATOS_SUPORTADOS:
                    caminho_completo = os.path.join(dir_path, arquivo)
                    timestamp = datetime.fromtimestamp(os.path.getmtime(caminho_completo))
                    
                    self.metadata["evidencias"].append({
                        "id": self.metadata["proximo_id"],
                        "arquivo": arquivo,
                        "comentario": "",
                        "timestamp": timestamp.strftime('%Y-%m-%d %H:%M:%S'),
                        "excluida": False
                    })
                    self.metadata["proximo_id"] += 1
            
            self._salvar_metadata()
        
        evidencias_ativas = []
        for evidencia in self.metadata["evidencias"]:
            if not evidencia.get("excluida", False):
                caminho = os.path.join(dir_path, evidencia["arquivo"])
                if os.path.exists(caminho):
                    _, ext = os.path.splitext(evidencia["arquivo"])
                    if ext.lower() in FORMATOS_SUPORTADOS:
                        evidencias_ativas.append(caminho)
        
        return evidencias_ativas

    def recarregar_evidencias(self):
        """Recarrega a lista de evidências"""
        if self.evidence_dir:
            self.prints = self.carregar_evidencias(self.evidence_dir)
            return True
        return False

    def obter_comentario(self, nome_arquivo):
        """Obtém o comentário salvo nos metadados"""
        for evidencia in self.metadata["evidencias"]:
            if evidencia["arquivo"] == nome_arquivo:
                return evidencia.get("comentario", "")
        return ""
    
    def show(self):
        """Mostra a interface do módulo"""
        if not self.root:
            self._create_interface()
        else:
            self.root.deiconify()
            self.root.lift()
            self.root.focus_set()

    def _create_interface(self):
        """Cria la interface do módulo"""
        self.root = tk.Toplevel(self.parent)
        self.root.title("PrintF - Gerador de Documentos de Evidências")
        self.root.geometry("350x350")
        self.root.resizable(True, False)
        
        self._apply_styles(self.root)
        
        if self.parent:
            self.root.transient(self.parent)
            self.root.grab_set()
        
        main_frame = self._create_styled_frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=30)
        
        self._create_styled_label(main_frame, text="Gerador de Documentos de Evidências", 
                                 style_type="title").pack(pady=20)
        
        desc_text = "Este módulo permite gerar documentos DOCX\na partir de evidências capturadas."
        desc_label = self._create_styled_label(main_frame, text=desc_text)
        desc_label.pack(pady=10)
        
        def iniciar():
            if self.mostrar_janela_configuracao():
                pass
        
        btn_frame = self._create_styled_frame(main_frame)
        btn_frame.pack(pady=20)
        
        self._create_styled_button(btn_frame, text="Iniciar Gerador", 
                                  command=iniciar, style_type="accent", width=20).pack(pady=10)
        
        self._create_styled_button(btn_frame, text="Voltar ao Menu Principal", 
                                  command=self.hide, style_type="glass", width=20).pack(pady=5)

    def hide(self):
        """Esconde a interface do módulo de forma segura"""
        if self.root:
            try:
                if hasattr(self, 'gravando') and self.gravando:
                    self.finalizar()
                
                if hasattr(self, 'listener_keyboard') and self.listener_keyboard:
                    try:
                        self.listener_keyboard.stop()
                    except:
                        pass
                    self.listener_keyboard = None
                
                if hasattr(self, 'listener_mouse') and self.listener_mouse:
                    try:
                        self.listener_mouse.stop()
                    except:
                        pass
                    self.listener_mouse = None
                
                if hasattr(self, 'popup') and self.popup and self.popup.winfo_exists():
                    try:
                        self.popup.destroy()
                    except:
                        pass
                    self.popup = None
                
                try:
                    self.root.grab_release()
                except:
                    pass
                    
                self.root.withdraw()
                
                if self.parent and self.parent.winfo_exists():
                    try:
                        self.parent.deiconify()
                        self.parent.lift()
                        self.parent.focus_force()
                    except:
                        pass
                    
            except Exception as e:
                print(f"Erro ao esconder módulo: {e}")
                try:
                    self.root.destroy()
                    self.root = None
                except:
                    pass

    def finalizar(self):
        """Método para finalizar gravação (para compatibilidade)"""
        self.gravando = False

    def mostrar_janela_configuracao(self):
        """Janela de configuração - SIMPLIFICADA SEM LISTA"""
        config_window = tk.Toplevel(self.root)
        config_window.title("Configuração de Arquivo")
        config_window.geometry("650x515")
        config_window.resizable(True, True)
        config_window.minsize(650,515)
        
        self._apply_styles(config_window)
        
        if self.parent:
            config_window.transient(self.root)
            config_window.grab_set()
        
        main_frame = self._create_styled_frame(config_window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=50, pady=50)
        
        self._create_styled_label(main_frame, text="PrintF - Configuração de Arquivo", 
                                 style_type="title").pack(pady=(0, 10))
        
        instrucoes_text = "Configure as opções abaixo e clique em 'GERAR DOCUMENTO' para iniciar"
        instrucoes_label = self._create_styled_label(main_frame, text=instrucoes_text)
        instrucoes_label.pack(pady=(0, 20))
        if not self.using_liquid_glass:
            instrucoes_label.config(font=("Arial", 9), fg='#7f8c8d')
        
        template_label = self._create_styled_label(main_frame, text="Selecione o template DOCX: *")
        template_label.pack(anchor="w", pady=(10, 5))
        if not self.using_liquid_glass:
            template_label.config(font=("Arial", 10, "bold"))
        
        template_frame = self._create_styled_frame(main_frame)
        template_frame.pack(fill=tk.X, pady=5)
        
        self.template_var = tk.StringVar()
        template_entry = self._create_styled_entry(template_frame, textvariable=self.template_var, width=50)
        template_entry.pack(side=tk.LEFT, padx=(0, 5), fill=tk.X, expand=True)
        
        def selecionar_template():
            template_path = filedialog.askopenfilename(
                title="Selecione o template DOCX",
                filetypes=[("Documentos Word", "*.docx")]
            )
            if template_path:
                self.template_var.set(template_path)
        
        btn_template = self._create_styled_button(template_frame, text="Procurar", 
                                                  command=selecionar_template, style_type="glass")
        btn_template.pack(side=tk.RIGHT)
        
        dir_label = self._create_styled_label(main_frame, text="Selecione o diretório onde estão as evidências: *")
        dir_label.pack(anchor="w", pady=(10, 5))
        if not self.using_liquid_glass:
            dir_label.config(font=("Arial", 10, "bold"))
        
        formatos_info = self._create_styled_label(main_frame, text="📷 Formatos suportados: PNG, JPG, JPEG, BMP, GIF, TIFF")
        formatos_info.pack(anchor="w", pady=(0, 5))
        if not self.using_liquid_glass:
            formatos_info.config(font=("Arial", 8), fg='#7f8c8d')
        
        dir_frame = self._create_styled_frame(main_frame)
        dir_frame.pack(fill=tk.X, pady=5)
        
        self.dir_var = tk.StringVar()
        dir_entry = self._create_styled_entry(dir_frame, textvariable=self.dir_var, width=50)
        dir_entry.pack(side=tk.LEFT, padx=(0, 5), fill=tk.X, expand=True)
        
        def selecionar_diretorio():
            dir_path = filedialog.askdirectory(title="Selecione o diretório onde estão as evidências")
            if dir_path:
                self.dir_var.set(dir_path)
                try:
                    arquivos = self.carregar_evidencias(dir_path)
                    if arquivos:
                        messagebox.showinfo("Arquivos Encontrados", 
                                          f"✅ {len(arquivos)} arquivo(s) de imagem encontrado(s)\n\n"
                                          f"Clique em 'GERAR DOCUMENTO' para continuar.")
                    else:
                        messagebox.showwarning("Nenhum Arquivo", 
                                             "⚠️ Nenhum arquivo de imagem encontrado neste diretório.\n\n"
                                             "Formatos suportados: PNG, JPG, JPEG, BMP, GIF, TIFF")
                except Exception as e:
                    messagebox.showerror("Erro", f"Erro ao verificar diretório:\n{e}")
        
        btn_dir = self._create_styled_button(dir_frame, text="Procurar", 
                                            command=selecionar_diretorio, style_type="glass")
        btn_dir.pack(side=tk.RIGHT)
        
        info_frame = self._create_styled_frame(main_frame)
        info_frame.pack(fill=tk.X, pady=20)
        
        info_text = "💡 Dica: Os arquivos serão processados em ordem cronológica"
        info_label = self._create_styled_label(info_frame, text=info_text)
        info_label.pack()
        if not self.using_liquid_glass:
            info_label.config(font=("Arial", 9), fg='#7f8c8d', wraplength=600)
        
        separator_frame = self._create_styled_frame(main_frame)
        separator_frame.pack(fill=tk.X, pady=15)
        
        if self.using_liquid_glass and self.style_manager:
            separator = ttk.Separator(separator_frame, orient='horizontal', style="Glass.TSeparator")
        else:
            separator = ttk.Separator(separator_frame, orient='horizontal')
        separator.pack(fill=tk.X)
        
        btn_container = self._create_styled_frame(main_frame)
        btn_container.pack(side=tk.BOTTOM, fill=tk.X, pady=(10, 0))
        
        btn_frame = self._create_styled_frame(btn_container)
        btn_frame.pack(anchor="center")
        
        def iniciar_geracao():
            if not self.template_var.get() or not self.dir_var.get():
                messagebox.showerror("Erro", "Por favor, selecione o template e o diretório de evidências.")
                return
            
            if not os.path.exists(self.template_var.get()):
                messagebox.showerror("Erro", "O arquivo de template selecionado não existe.")
                return
            
            if not os.path.exists(self.dir_var.get()):
                messagebox.showerror("Erro", "O diretório de evidências selecionado não existe.")
                return
            
            png_files = self.carregar_evidencias(self.dir_var.get())
            if not png_files:
                messagebox.showerror("Erro", "Nenhuma evidência de imagem encontrada no diretório selecionado.\n\nFormatos suportados: PNG, JPG, JPEG, BMP, GIF, TIFF")
                return
            
            self.template_path = self.template_var.get()
            self.output_dir = self.dir_var.get()
            self.evidence_dir = self.dir_var.get()
            self.prints = png_files
            self.current_index = 0
            
            config_window.destroy()            
            self.iniciar_processamento()
        
        btn_gerar = self._create_styled_button(btn_frame, text="✅ GERAR DOCUMENTO", 
                                              command=iniciar_geracao, style_type="accent", width=25)
        btn_gerar.pack(side=tk.LEFT, padx=10, pady=5)
        
        btn_cancelar = self._create_styled_button(btn_frame, text="❌ CANCELAR", 
                                                  command=config_window.destroy, style_type="glass", width=25)
        btn_cancelar.pack(side=tk.LEFT, padx=10, pady=5)
        
        config_window.update_idletasks()
        
        self.root.wait_window(config_window)
        return self.template_path is not None and self.output_dir is not None and self.prints

    def iniciar_processamento(self):
        """Inicia o processamento das evidências - AGORA USA NAVEGAÇÃO"""
        os.makedirs(self.output_dir, exist_ok=True)

        try:
            if os.path.exists(self.template_path):
                self.doc = Document(self.template_path)
                self.using_template = True
            else:
                self.doc = Document()
                self.using_template = False
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar template: {str(e)}")
            self.doc = Document()
            self.using_template = False
        
        self.mostrar_janela_navegacao()

    def mostrar_janela_navegacao(self):
        """Janela principal de navegação pelas evidências"""
        if self.popup and self.popup.winfo_exists():
            self.popup.destroy()

        self.popup = tk.Toplevel(self.root)
        self.popup.title("Navegação de Evidências")
        self.popup.geometry("1200x800")
        self.popup.resizable(True, True)
        
        self._apply_styles(self.popup)
        
        self.popup.transient(self.root)
        
        self.popup.grid_columnconfigure(0, weight=1)
        self.popup.grid_rowconfigure(0, weight=1)
        
        img_frame = self._create_styled_frame(self.popup)
        img_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        img_frame.grid_rowconfigure(0, weight=1)
        img_frame.grid_columnconfigure(0, weight=1)
        
        self.current_img_label = tk.Label(img_frame, bg="white")
        self.current_img_label.grid(row=0, column=0, sticky="nsew")
        
        comment_frame = self._create_styled_frame(self.popup)
        comment_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=(0, 5))
                
        self._create_styled_label(comment_frame, text="Comentário:").pack(anchor="w")
        
        comment_entry_frame = self._create_styled_frame(comment_frame)
        comment_entry_frame.pack(fill=tk.X, pady=2)
        
        self.comment_entry = tk.Entry(comment_entry_frame, font=("Arial", 10))
        self.comment_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.comment_entry.bind("<FocusOut>", lambda e: self.salvar_comentario())
        
        buttons_main_frame = self._create_styled_frame(self.popup)
        buttons_main_frame.grid(row=2, column=0, sticky="nsew", padx=10, pady=5)
        
        nav_frame = self._create_styled_frame(buttons_main_frame)
        nav_frame.pack(expand=True, pady=2)
        
        self._create_styled_button(nav_frame, text="⏮️ Primeira", command=self.primeira_evidencia, 
                                 style_type="glass").pack(side=tk.LEFT, padx=2)
        self._create_styled_button(nav_frame, text="◀️ Anterior", command=self.anterior_evidencia,
                                 style_type="glass").pack(side=tk.LEFT, padx=2)
        
        self.pos_label = tk.Label(nav_frame, text="", font=("Arial", 12, "bold"))
        self.pos_label.pack(side=tk.LEFT, padx=15)
        
        self._create_styled_button(nav_frame, text="▶️ Próxima", command=self.proxima_evidencia,
                                 style_type="glass").pack(side=tk.LEFT, padx=2)
        self._create_styled_button(nav_frame, text="⏭️ Última", command=self.ultima_evidencia,
                                 style_type="glass").pack(side=tk.LEFT, padx=2)
        
        self._create_styled_button(nav_frame, text="🔢 Ir para...", command=self.ir_para_especifica,
                                 style_type="glass").pack(side=tk.LEFT, padx=2)
        
        action_frame = self._create_styled_frame(buttons_main_frame)
        action_frame.pack(expand=True, pady=2)
        
        self._create_styled_button(action_frame, text="✏️ Editar Print", command=self.editar_evidencia_atual,
                                 style_type="glass").pack(side=tk.LEFT, padx=5)
        self._create_styled_button(action_frame, text="🗑️ Excluir Print", command=self.excluir_evidencia_atual,
                                 style_type="glass").pack(side=tk.LEFT, padx=5)
        
        control_frame = self._create_styled_frame(self.popup)
        control_frame.grid(row=3, column=0, sticky="ew", padx=10, pady=5)
        
        control_buttons_frame = self._create_styled_frame(control_frame)
        control_buttons_frame.pack(expand=True)
        
        self._create_styled_button(control_buttons_frame, text="❌ Cancelar", command=self.cancelar_processamento,
                                 style_type="glass").pack(side=tk.LEFT, padx=5)
        
        self._create_styled_button(control_buttons_frame, text="✅ Gerar Evidência", command=self.finalizar_processamento,
                                 style_type="accent").pack(side=tk.LEFT, padx=5)
        
        self.current_index = 0
        self.atualizar_exibicao()
        
        self.popup.protocol("WM_DELETE_WINDOW", self.cancelar_processamento)

    def atualizar_exibicao(self):
        """Atualiza a exibição da evidência atual"""
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        
        try:
            img = Image.open(caminho_print)
            
            self.popup.update()
            available_width = self.popup.winfo_width() - 40
            available_height = self.popup.winfo_height() - 250
            
            img.thumbnail((available_width, available_height))
            self.current_img_tk = ImageTk.PhotoImage(img)
            self.current_img_label.config(image=self.current_img_tk)
            
            self.pos_label.config(text=f"Evidência {self.current_index + 1} de {len(self.prints)}")
            
            nome_arquivo = os.path.basename(caminho_print)
            comentario = self.obter_comentario(nome_arquivo)
            self.comment_entry.delete(0, tk.END)
            self.comment_entry.insert(0, comentario)
            
        except Exception as e:
            print(f"Erro ao carregar imagem: {e}")

    def salvar_comentario(self):
        """Salva o comentário da evidência atual"""
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        nome_arquivo = os.path.basename(caminho_print)
        comentario = self.comment_entry.get()
        
        for evidencia in self.metadata["evidencias"]:
            if evidencia["arquivo"] == nome_arquivo:
                evidencia["comentario"] = comentario
                break
                
        self._salvar_metadata()

    def primeira_evidencia(self):
        self.salvar_comentario()
        self.current_index = 0
        self.atualizar_exibicao()

    def anterior_evidencia(self):
        self.salvar_comentario()
        if self.current_index > 0:
            self.current_index -= 1
            self.atualizar_exibicao()

    def proxima_evidencia(self):
        self.salvar_comentario()
        if self.current_index < len(self.prints) - 1:
            self.current_index += 1
            self.atualizar_exibicao()

    def ultima_evidencia(self):
        self.salvar_comentario()
        self.current_index = len(self.prints) - 1
        self.atualizar_exibicao()

    def ir_para_especifica(self):
        self.salvar_comentario()
        if not self.prints:
            return
            
        numero = simpledialog.askinteger("Navegar", 
                                       f"Digite o número da evidência (1-{len(self.prints)}):",
                                       minvalue=1, maxvalue=len(self.prints))
        if numero:
            self.current_index = numero - 1
            self.atualizar_exibicao()

    def editar_evidencia_atual(self):
        """Abre o editor para a evidência atual e atualiza a exibição após edição"""
        self.salvar_comentario()
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        
        # CORREÇÃO: Criar uma variável para controlar se a edição foi salva
        self.edicao_salva = False
        
        # Abrir editor modal (aguardar até ser fechado)
        editor = self.abrir_editor(caminho_print, self.popup)
        
        # CORREÇÃO: Esperar o editor ser fechado antes de continuar
        if editor:
            # Focar na janela do editor e aguardar
            editor.focus_set()
            editor.grab_set()
            self.popup.wait_window(editor)
        
        # CORREÇÃO: Forçar a recarga completa da imagem após a edição
        # Limpar todas as referências de imagem
        if hasattr(self, 'current_img_tk'):
            del self.current_img_tk
            self.current_img_tk = None
        
        # Forçar o garbage collection
        import gc
        gc.collect()
        
        # Atualizar a exibição para mostrar a imagem editada
        self.atualizar_exibicao()
        
        # Feedback visual
        print(f"✅ Evidência atualizada: {os.path.basename(caminho_print)}")

    def excluir_evidencia_atual(self):
        self.salvar_comentario()
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        nome_arquivo = os.path.basename(caminho_print)
        
        if messagebox.askyesno("Confirmar Exclusão", 
                             "Tem certeza que deseja excluir este print?"):
            try:
                os.remove(caminho_print)
                
                for evidencia in self.metadata["evidencias"]:
                    if evidencia["arquivo"] == nome_arquivo:
                        evidencia["excluida"] = True
                        break
                
                self._salvar_metadata()
                
                self.recarregar_evidencias()
                
                if not self.prints:
                    messagebox.showinfo("Info", "Todas as evidências foram processadas.")
                    self.finalizar_processamento()
                    return
                
                if self.current_index >= len(self.prints):
                    self.current_index = len(self.prints) - 1
                
                self.atualizar_exibicao()
                messagebox.showinfo("Sucesso", "Evidência excluída!")
                
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao excluir: {str(e)}")

    def finalizar_processamento(self):
        """Processa todas as evidências e gera o DOCX"""
        self.salvar_comentario()
        
        try:
            doc_path = self.gerar_documento()
            
            pasta_para_abrir = os.path.dirname(doc_path)
            
            resposta = messagebox.askyesno(
                "Sucesso", 
                f"Documento gerado com sucesso em:\n{doc_path}\n\nDeseja abrir a pasta onde o documento foi salvo?",
                parent=self.popup
            )
            
            if resposta:
                if not self._abrir_pasta(pasta_para_abrir):
                    messagebox.showinfo(
                        "Abrir Pasta", 
                        f"Pasta do documento:\n{pasta_para_abrir}",
                        parent=self.popup
                    )
                    
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar documento: {e}", parent=self.popup)
        
        if self.popup and self.popup.winfo_exists():
            self.popup.destroy()
            self.popup = None

    def _abrir_pasta(self, caminho_pasta):
        """Abre a pasta no explorador de arquivos do sistema"""
        try:
            if os.name == 'nt':
                os.startfile(caminho_pasta)
            elif os.name == 'posix':
                if sys.platform == 'darwin':
                    subprocess.run(['open', caminho_pasta])
                else:
                    subprocess.run(['xdg-open', caminho_pasta])
            return True
        except Exception as e:
            print(f"Erro ao abrir pasta: {e}")
            return False

    def cancelar_processamento(self):
        self.salvar_comentario()
        if messagebox.askyesno("Confirmar", "Deseja cancelar o processamento?"):
            if self.popup:
                self.popup.destroy()
                self.popup = None

    def gerar_documento(self):
        """Gera o documento DOCX com as evidências e retorna o caminho do documento"""
        doc_path = None
        try:
            print("📄 Iniciando geração do documento DOCX...")
            
            if self.template_path and os.path.exists(self.template_path):
                self.doc = Document(self.template_path)
                self.using_template = True
                print(f"✅ Template carregado: {self.template_path}")
            else:
                self.doc = Document()
                self.using_template = False
                print("ℹ️ Criando documento vazio (sem template)")
            
            if not self.using_template:
                titulo = self.doc.add_heading('Evidências Capturadas', 0)
                titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if not self.using_template:
                data_hora = self.doc.add_paragraph()
                data_hora.add_run(f"Data e hora da geração: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}").italic = True
                data_hora.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            for i, print_path in enumerate(self.prints, 1):
                print(f"📷 Adicionando evidência {i}: {print_path}")
                
                # 🔥 CORREÇÃO: Adicionar separador ANTES da primeira evidência
                if i == 1:
                    self.doc.add_paragraph("―" * 36).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                self.doc.add_paragraph().add_run(f"Evidência {i}").bold = True
                
                nome_arquivo = os.path.basename(print_path)
                comentario = self.obter_comentario(nome_arquivo)
                if comentario:
                    comentario_para = self.doc.add_paragraph()
                    comentario_para.add_run(f"Comentário: {comentario}").italic = True
                
                try:
                    paragraph = self.doc.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run()
                    
                    if os.path.exists(print_path):
                        run.add_picture(print_path, width=Inches(6.0))
                        print(f"✅ Imagem {i} adicionada com sucesso")
                    else:
                        print(f"⚠️ Arquivo não encontrado: {print_path}")
                        self.doc.add_paragraph(f"[Arquivo de imagem não encontrado: {print_path}]")
                        
                except Exception as e:
                    print(f"❌ Erro ao adicionar imagem {print_path}: {e}")
                    self.doc.add_paragraph(f"[Erro ao carregar imagem: {print_path}]")
                
                self.doc.add_paragraph("―" * 36).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            template_filename = os.path.basename(self.template_path)
            template_name = os.path.splitext(template_filename)[0]
            
            template_name = self._limpar_nome_arquivo(template_name)
            
            doc_filename = f"{template_name}_{timestamp}.docx"
            doc_path = os.path.join(self.output_dir, doc_filename)
            
            os.makedirs(os.path.dirname(doc_path), exist_ok=True)
            
            if len(doc_path) > 255:
                short_name = f"Evidencias_{timestamp}.docx"
                doc_path = os.path.join(self.output_dir, short_name)
                print(f"⚠️ Caminho muito longo, usando nome reduzido: {short_name}")
            
            self.doc.save(doc_path)
            print(f"✅ Documento salvo em: {doc_path}")
            
            return doc_path
            
        except Exception as e:
            print(f"❌ Erro ao gerar documento: {e}")
            import traceback
            traceback.print_exc()
            raise

    def _limpar_nome_arquivo(self, nome):
        """Remove caracteres inválidos para nomes de arquivo no Windows, mantendo caracteres PT-BR"""
        caracteres_invalidos = r'[\\/*?:"<>|]'
        nome_limpo = re.sub(caracteres_invalidos, '_', nome)
        
        nome_limpo = re.sub(r'[^\w\s\-\.\(\)áàâãéèêíïóôõöúçñÁÀÂÃÉÈÊÍÏÓÔÕÖÚÇÑ]', '', nome_limpo)
        
        if len(nome_limpo) > 100:
            nome_limpo = nome_limpo[:100]
            
        return nome_limpo.strip()

    def abrir_editor(self, caminho_print, parent):
        """Abre editor de imagens para a evidência - RETORNA A JANELA DO EDITOR"""
        editor = tk.Toplevel(parent)
        editor.title("Editor de Evidência")
        editor.geometry("1200x800")
        
        self._apply_styles(editor)
        
        main_frame = self._create_styled_frame(editor)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        tools_frame = self._create_styled_frame(main_frame)
        tools_frame.pack(side=tk.TOP, fill=tk.X, pady=5)
        
        canvas_frame = self._create_styled_frame(main_frame)
        canvas_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        self.original_img = Image.open(caminho_print).convert("RGBA")
        img_w, img_h = self.original_img.size
        
        max_w, max_h = 1000, 700
        scale = min(max_w / img_w, max_h / img_h)
        self.scale_factor = scale
        disp_w, disp_h = int(img_w * scale), int(img_h * scale)
        
        self.editing_img = self.original_img.copy()
        self.display_img = self.editing_img.resize((disp_w, disp_h), Image.LANCZOS)

        self.current_tk_img = ImageTk.PhotoImage(self.display_img)
        self.elements = []
        self.undo_stack = []
        self.temp_element = None
        
        self.canvas = tk.Canvas(canvas_frame, width=disp_w, height=disp_h, cursor="cross", bg="gray")
        self.canvas.pack(padx=5, pady=5)
        self.canvas_img = self.canvas.create_image(0, 0, anchor="nw", image=self.current_tk_img)
        
        tool_var = tk.StringVar(value="rectangle")
        color_var = tk.StringVar(value="#FF0000")
        width_var = tk.IntVar(value=3)
        
        self._create_styled_label(tools_frame, text="Ferramenta:").pack(side=tk.LEFT, padx=5)
        
        icon_frame = self._create_styled_frame(tools_frame)
        icon_frame.pack(side=tk.LEFT, padx=5)
        
        tool_icons = {
            "rectangle": "⬜",
            "circle": "🔴",
            "arrow": "👉",
            "text": "🆎",
            "blur": "🌀"
        }

        def criar_botao_ferramenta(parent, texto, valor, variavel):
            btn = tk.Radiobutton(parent, text=texto, font=("Arial", 12), 
                               variable=variavel, value=valor, indicatoron=0, 
                               width=3, height=2, relief=tk.RAISED,
                               cursor="hand2")
            return btn

        for tool_value, icon in tool_icons.items():
            btn = criar_botao_ferramenta(icon_frame, icon, tool_value, tool_var)
            btn.pack(side=tk.LEFT, padx=2)

        for widget in icon_frame.winfo_children():
            if isinstance(widget, tk.Radiobutton) and widget.cget("value") == "rectangle":
                widget.config(relief=tk.SUNKEN, bg="#e3f2fd")
                break

        def update_button_appearance(*args):
            selected_tool = tool_var.get()
            for widget in icon_frame.winfo_children():
                if isinstance(widget, tk.Radiobutton):
                    if widget.cget("value") == selected_tool:
                        widget.config(relief=tk.SUNKEN, bg="#e3f2fd")
                    else:
                        widget.config(relief=tk.RAISED, bg="SystemButtonFace")

        tool_var.trace("w", update_button_appearance)
        
        color_frame = self._create_styled_frame(tools_frame)
        color_frame.pack(side=tk.LEFT, padx=20)
        
        self._create_styled_label(color_frame, text="Cor:").pack(side=tk.LEFT)
        
        colors = ["#FF0000", "#00FF00", "#FFFF00", "#000000", "#FFFFFF"]
        color_buttons_frame = self._create_styled_frame(color_frame)
        color_buttons_frame.pack(side=tk.LEFT, padx=5)
        
        for color in colors:
            btn = tk.Button(color_buttons_frame, bg=color, width=2, height=1, 
                           command=lambda c=color: self.set_color(color_var, c, color_preview))
            btn.pack(side=tk.LEFT, padx=1)
        
        custom_btn = self._create_styled_button(color_frame, text="Personalizada", 
                              command=lambda: self.choose_custom_color(editor, color_var, color_preview))
        custom_btn.pack(side=tk.LEFT, padx=5)
        
        color_preview = tk.Frame(color_frame, width=30, height=20, bg=color_var.get())
        color_preview.pack(side=tk.LEFT, padx=5)
        
        width_frame = self._create_styled_frame(tools_frame)
        width_frame.pack(side=tk.LEFT, padx=20)
        
        self._create_styled_label(width_frame, text="Espessura:").pack(side=tk.LEFT)
        tk.Scale(width_frame, from_=1, to=10, variable=width_var, orient=tk.HORIZONTAL, 
                length=100, showvalue=1).pack(side=tk.LEFT, padx=5)
        
        def undo_action():
            if self.elements:
                removed_element = self.elements.pop()
                self.undo_stack.append(removed_element)
                refresh_display()
        
        undo_btn = self._create_styled_button(tools_frame, text="↩️ Desfazer (Ctrl+Z)", command=undo_action)
        undo_btn.pack(side=tk.LEFT, padx=20)
        
        start_xy = {"x": None, "y": None}
        
        def refresh_display():
            self.canvas.delete("all")
            self.canvas.create_image(0, 0, anchor="nw", image=self.current_tk_img)
            
            for element in self.elements:
                elem_type, coords, color, width, text = element
                scaled_coords = [int(c * self.scale_factor) for c in coords]
                
                if elem_type == "circle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_oval(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "rectangle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "arrow":
                    x1, y1, x2, y2 = scaled_coords
                    self.draw_arrow_on_canvas(x1, y1, x2, y2, color, width)
                elif elem_type == "text":
                    x, y = scaled_coords
                    self.canvas.create_text(x, y, text=text, fill=color, font=("Arial", 12), anchor="nw")
                elif elem_type == "blur":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline="#FF00FF", width=2, dash=(5,5))
            
            if self.temp_element:
                elem_type, coords, color, width, text = self.temp_element
                scaled_coords = [int(c * self.scale_factor) for c in coords]
                
                if elem_type == "circle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_oval(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "rectangle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "arrow":
                    x1, y1, x2, y2 = scaled_coords
                    self.draw_arrow_on_canvas(x1, y1, x2, y2, color, width)
                elif elem_type == "blur":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline="#FF00FF", width=2, dash=(5,5))
        
        def draw_arrow_on_canvas(x1, y1, x2, y2, color, width):
            self.canvas.create_line(x1, y1, x2, y2, fill=color, width=width)
            
            angle = math.atan2(y2 - y1, x2 - x1)
            
            arrow_size = 15
            x3 = x2 - arrow_size * math.cos(angle - math.pi/6)
            y3 = y2 - arrow_size * math.sin(angle - math.pi/6)
            x4 = x2 - arrow_size * math.cos(angle + math.pi/6)
            y4 = y2 - arrow_size * math.sin(angle + math.pi/6)
            
            self.canvas.create_polygon(x2, y2, x3, y3, x4, y4, fill=color, outline=color)
        
        def on_button_press(event):
            start_xy["x"], start_xy["y"] = event.x, event.y
        
        def on_motion(event):
            if start_xy["x"] is not None:
                sx, sy = start_xy["x"], start_xy["y"]
                ex, ey = event.x, event.y
                
                ix1, iy1 = int(sx / self.scale_factor), int(sy / self.scale_factor)
                ix2, iy2 = int(ex / self.scale_factor), int(ey / self.scale_factor)
                
                tool = tool_var.get()
                color = color_var.get()
                width = width_var.get()
                
                if tool == "circle":
                    radius = int(((ix2 - ix1)**2 + (iy2 - iy1)**2)**0.5)
                    self.temp_element = ("circle", [ix1-radius, iy1-radius, ix1+radius, iy1+radius], color, width, "")
                elif tool == "rectangle":
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.temp_element = ("rectangle", [x1_norm, y1_norm, x2_norm, y2_norm], color, width, "")
                elif tool == "arrow":
                    self.temp_element = ("arrow", [ix1, iy1, ix2, iy2], color, width, "")
                elif tool == "blur":
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.temp_element = ("blur", [x1_norm, y1_norm, x2_norm, y2_norm], "#FF00FF", 2, "")
                
                refresh_display()
        
        def on_button_release(event):
            if start_xy["x"] is not None:
                sx, sy = start_xy["x"], start_xy["y"]
                ex, ey = event.x, event.y
                
                ix1, iy1 = int(sx / self.scale_factor), int(sy / self.scale_factor)
                ix2, iy2 = int(ex / self.scale_factor), int(ey / self.scale_factor)
                
                tool = tool_var.get()
                color = color_var.get()
                width = width_var.get()
                
                self.undo_stack.clear()
                
                if tool == "circle":
                    radius = int(((ix2 - ix1)**2 + (iy2 - iy1)**2)**0.5)
                    self.elements.append(("circle", [ix1-radius, iy1-radius, ix1+radius, iy1+radius], color, width, ""))
                
                elif tool == "rectangle":
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.elements.append(("rectangle", [x1_norm, y1_norm, x2_norm, y2_norm], color, width, ""))
                
                elif tool == "arrow":
                    self.elements.append(("arrow", [ix1, iy1, ix2, iy2], color, width, ""))
                
                elif tool == "blur":
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.elements.append(("blur", [x1_norm, y1_norm, x2_norm, y2_norm], "", 0, ""))
                
                elif tool == "text":
                    text = simpledialog.askstring("Texto", "Digite o texto:", parent=editor)
                    if text:
                        self.elements.append(("text", [ix1, iy1], color, width, text))
                        refresh_display()
                
                self.temp_element = None
                refresh_display()
            
            start_xy["x"], start_xy["y"] = None, None
        
        def on_key_press(event):
            if event.keysym == 'z' and (event.state & 0x4):
                undo_action()

        editor.bind('<Control-z>', on_key_press)
        editor.bind('<Control-Z>', on_key_press)
        
        self.canvas.bind("<ButtonPress-1>", on_button_press)
        self.canvas.bind("<B1-Motion>", on_motion)
        self.canvas.bind("<ButtonRelease-1>", on_button_release)
        
        refresh_display()
        
        button_frame = self._create_styled_frame(canvas_frame)
        button_frame.pack(pady=10)
        
        def salvar_edicao():
            if hasattr(self, 'color_chooser_window') and self.color_chooser_window:
                try:
                    self.color_chooser_window.destroy()
                except:
                    pass
            
            draw = ImageDraw.Draw(self.editing_img)
            
            for element in self.elements:
                elem_type, coords, color, width, text = element
                
                if elem_type == "circle":
                    x1, y1, x2, y2 = coords
                    draw.ellipse((x1, y1, x2, y2), outline=color, width=width)
                
                elif elem_type == "rectangle":
                    x1, y1, x2, y2 = coords
                    x1_norm = min(x1, x2)
                    y1_norm = min(y1, y2)
                    x2_norm = max(x1, x2)
                    y2_norm = max(y1, y2)
                    draw.rectangle((x1_norm, y1_norm, x2_norm, y2_norm), outline=color, width=width)
                
                elif elem_type == "arrow":
                    x1, y1, x2, y2 = coords
                    draw.line((x1, y1, x2, y2), fill=color, width=width)
                    
                    angle = math.atan2(y2 - y1, x2 - x1)
                    
                    arrow_size = 20
                    x3 = x2 - arrow_size * math.cos(angle - math.pi/6)
                    y3 = y2 - arrow_size * math.sin(angle - math.pi/6)
                    x4 = x2 - arrow_size * math.cos(angle + math.pi/6)
                    y4 = y2 - arrow_size * math.sin(angle + math.pi/6)
                    
                    draw.polygon([(x2, y2), (x3, y3), (x4, y4)], fill=color)
                
                elif elem_type == "text":
                    x, y = coords
                    try:
                        font = ImageFont.truetype("arial.ttf", 16)
                    except:
                        font = ImageFont.load_default()
                    
                    draw.text((x, y), text, fill=color, font=font)
                
                elif elem_type == "blur":
                    x1, y1, x2, y2 = coords
                    region = self.editing_img.crop((x1, y1, x2, y2))
                    blurred_region = region.filter(ImageFilter.GaussianBlur(15))
                    self.editing_img.paste(blurred_region, (x1, y1, x2, y2))
            
            self.editing_img.convert("RGB").save(caminho_print, "PNG")
            
            # CORREÇÃO: Marcar que a edição foi salva
            self.edicao_salva = True
            
            messagebox.showinfo("Edição", "Evidência atualizada com sucesso!")
            editor.destroy()

        def fechar_editor():
            if hasattr(self, 'color_chooser_window') and self.color_chooser_window:
                try:
                    self.color_chooser_window.destroy()
                except:
                    pass
            
            # CORREÇÃO: Se não foi salvo, marcar como não salvo
            if not hasattr(self, 'edicao_salva'):
                self.edicao_salva = False
                
            editor.destroy()

        editor.protocol("WM_DELETE_WINDOW", fechar_editor)
        
        self._create_styled_button(button_frame, text="💾 Salvar e Fechar", command=salvar_edicao, 
                                 style_type="accent").pack()

        editor.transient(parent)
        
        # CORREÇÃO: Retornar a janela do editor para controle modal
        return editor

    def set_color(self, color_var, color, preview_widget):
        color_var.set(color)
        preview_widget.config(bg=color)
    
    def choose_custom_color(self, parent, color_var, preview_widget):
        if hasattr(self, 'color_chooser_window') and self.color_chooser_window:
            try:
                self.color_chooser_window.destroy()
            except:
                pass
        
        color = colorchooser.askcolor(title="Escolha uma cor", initialcolor=color_var.get(), parent=parent)
        if color[1]:
            color_var.set(color[1])
            preview_widget.config(bg=color[1])
    
    def draw_arrow_on_canvas(self, x1, y1, x2, y2, color, width):
        self.canvas.create_line(x1, y1, x2, y2, fill=color, width=width)
        
        angle = math.atan2(y2 - y1, x2 - x1)
        
        arrow_size = 15
        x3 = x2 - arrow_size * math.cos(angle - math.pi/6)
        y3 = y2 - arrow_size * math.sin(angle - math.pi/6)
        x4 = x2 - arrow_size * math.cos(angle + math.pi/6)
        y4 = y2 - arrow_size * math.sin(angle + math.pi/6)
        
        self.canvas.create_polygon(x2, y2, x3, y3, x4, y4, fill=color, outline=color)


# Modo de execução independente
if __name__ == "__main__":
    root = tk.Tk()
    root.title("PrintF - Gerador de Evidências")
    root.geometry("350x350")
    root.resizable(False, False)
    
    try:
        root.configure(bg='#f5f5f5')
    except:
        pass
    
    root.eval('tk::PlaceWindow . center')
    
    main_frame = tk.Frame(root, bg='#f5f5f5', padx=30, pady=30)
    main_frame.pack(fill=tk.BOTH, expand=True)
    
    title_label = tk.Label(main_frame, text="PrintF - Gerador de Evidências", 
                         font=("Arial", 16, "bold"), bg='#f5f5f5', fg='#2c3e50')
    title_label.pack(pady=20)
    
    def iniciar_gerador():
        gerador = EvidenceGeneratorModule(root)
        gerador.show()
    
    start_btn = tk.Button(main_frame, text="Iniciar Gerador de Evidências", 
                         command=iniciar_gerador, width=25,
                         bg='#3498db', fg='white', font=("Arial", 12, "bold"),
                         relief="flat", cursor="hand2")
    start_btn.pack(pady=10)
    
    start_btn.bind("<Enter>", lambda e: start_btn.config(bg='#2980b9'))
    start_btn.bind("<Leave>", lambda e: start_btn.config(bg='#3498db'))
    
    exit_btn = tk.Button(main_frame, text="Sair", command=root.quit, width=15,
                        bg='#e74c3c', fg='white', font=("Arial", 10),
                        relief="flat", cursor="hand2")
    exit_btn.pack(pady=10)
    
    exit_btn.bind("<Enter>", lambda e: exit_btn.config(bg='#c0392b'))
    exit_btn.bind("<Leave>", lambda e: exit_btn.config(bg='#e74c3c'))
    
    root.mainloop()