import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, colorchooser, ttk
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pyautogui
from pynput import mouse, keyboard
from PIL import Image, ImageTk, ImageDraw, ImageFont, ImageGrab, ImageFilter  # 🔥 ADICIONADO ImageFilter
from datetime import datetime
import math
import re
import screeninfo
import glob
import json
import uuid
import time
import ctypes
from ctypes import wintypes, byref
import tkinter.font as tkfont
import shutil
import subprocess  # 🔥 ADICIONADO PARA ABRIR PASTA

# 🔥 CONTROLE AUTOMÁTICO DA BARRA DE TAREFAS
try:
    import win32gui
    import win32con
    import win32api
    WIN32_AVAILABLE = True
except ImportError:
    win32gui = None
    win32con = None
    win32api = None
    WIN32_AVAILABLE = False
    
# 🔥 ADICIONAR MSS PARA CAPTURA MULTI-MONITOR
try:
    import mss
    MSS_AVAILABLE = True
except ImportError:
    mss = None
    MSS_AVAILABLE = False

# ------------------ Gravador e Docx ------------------
class CaptureModule:
    def __init__(self, parent=None, settings=None):
        """Inicializa o módulo de captura para trabalhar com a main.py"""
        self.parent = parent
        self.settings = settings or {}
        self.root = None  # Será definido quando mostrar a interface
        
        # 🔥 ADICIONADO: Gerenciador de estilo
        self.style_manager = None
        self.using_liquid_glass = False
        self._setup_styles()
        
        self.gravando = False
        self.pausado = False
        self.output_dir = os.getcwd()
        self.listener_mouse = None
        self.prints = []
        self.doc = None
        self.evidencia_count = 0
        self.demanda = ""
        self.tipo_demanda = ""
        self.chamado = ""
        self.titulo = ""
        self.using_template = False
        self.template_path = None
        self.color_chooser_window = None
        self.current_index = 0
        self.evidence_dir = None
        self.metadata_path = None
        self.metadata = {"evidencias": [], "proximo_id": 1}
        self.popup = None
        self.current_img_label = None
        self.current_img_tk = None
        self.comment_entry = None
        self.manter_evidencias = None
        self.modo_captura = "ocultar"  # Valores: "manter", "ocultar"
        
        # 🔥 NOVOS ATRIBUTOS PARA PASTA AUTOMÁTICA
        self.pasta_automatica = False
        self.pasta_automatica_path = None
        
        # 🔥 ATRIBUTOS PARA NAVEGAÇÃO E EDIÇÃO
        self.elements = []  # Lista de elementos desenhados
        self.undo_stack = []  # PILHA PARA DESFAZER AÇÕES
        self.temp_element = None
        self.original_img = None
        self.editing_img = None
        self.display_img = None
        self.current_tk_img = None
        self.canvas = None
        self.canvas_img = None
        self.scale_factor = 1.0
        
        # Listener de teclado para atalhos
        self.listener_keyboard = None

    def _setup_styles(self):
        """Configura estilos visuais baseados no tema selecionado"""
        try:
            # 🔥 CORREÇÃO: Tentar importar de múltiplas formas
            try:
                # Tentar importar do módulo styles diretamente
                from styles import LiquidGlassStyle
                self.style_manager = LiquidGlassStyle
                self.using_liquid_glass = True
                print("✅ Estilo Liquid Glass carregado do módulo styles!")
                
            except ImportError:
                try:
                    # Tentar importar do diretório modules
                    from modules.styles import LiquidGlassStyle
                    self.style_manager = LiquidGlassStyle
                    self.using_liquid_glass = True
                    print("✅ Estilo Liquid Glass carregado do módulo modules.styles!")
                    
                except ImportError:
                    # Tentar importar relativo
                    import importlib.util
                    spec = importlib.util.spec_from_file_location("styles", "styles.py")
                    if spec and spec.loader:
                        styles_module = importlib.util.module_from_spec(spec)
                        spec.loader.exec_module(styles_module)
                        self.style_manager = styles_module.LiquidGlassStyle
                        self.using_liquid_glass = True
                        print("✅ Estilo Liquid Glass carregado de styles.py!")
                    else:
                        raise ImportError("Não foi possível carregar styles.py")
            
            # Verificar se o tema está habilitado nas configurações
            theme_to_use = self.settings.get('theme', 'liquid_glass')
            if theme_to_use == 'liquid_glass' and self.style_manager:
                self.using_liquid_glass = True
                print("✅ Estilo Liquid Glass configurado no módulo de captura!")
            else:
                self.using_liquid_glass = False
                print(f"ℹ️ Usando estilo padrão no módulo de captura (tema: {theme_to_use})")
            
        except ImportError as e:
            # Fallback para estilo padrão
            print(f"⚠️ Liquid Glass não disponível no módulo de captura: {e}")
            self.using_liquid_glass = False
        except Exception as e:
            print(f"⚠️ Erro ao configurar Liquid Glass no módulo de captura: {e}")
            self.using_liquid_glass = False

    def _apply_style_to_window(self, window):
        """Aplica o estilo Liquid Glass a uma janela se disponível"""
        if self.using_liquid_glass and self.style_manager:
            try:
                self.style_manager.apply_window_style(window)
                return True
            except Exception as e:
                print(f"⚠️ Erro ao aplicar estilo à janela: {e}")
                self.using_liquid_glass = False
        return False

    def _create_styled_frame(self, parent, **kwargs):
        """Cria um frame com estilo apropriado"""
        if self.using_liquid_glass and self.style_manager:
            try:
                return self.style_manager.create_glass_frame(parent, **kwargs)
            except Exception as e:
                print(f"⚠️ Erro ao criar frame estilizado: {e}")
                self.using_liquid_glass = False
        
        # Fallback para frame padrão
        frame = tk.Frame(parent, **kwargs)
        if not self.using_liquid_glass:
            frame.configure(bg='#f5f5f5')
        return frame

    def _create_styled_button(self, parent, text, command, style_type="glass", **kwargs):
        """Cria um botão com estilo apropriado"""
        if self.using_liquid_glass and self.style_manager:
            try:
                if style_type == "accent":
                    return self.style_manager.create_accent_button(parent, text, command, **kwargs)
                else:
                    return self.style_manager.create_glass_button(parent, text, command, **kwargs)
            except Exception as e:
                print(f"⚠️ Erro ao criar botão estilizado: {e}")
                self.using_liquid_glass = False
        
        # Fallback para botão padrão
        btn = tk.Button(parent, text=text, command=command, **kwargs)
        if style_type == "accent":
            btn.configure(bg="#3498db", fg="white", font=("Arial", 11, "bold"), relief="flat")
        else:
            btn.configure(bg="#ecf0f1", fg="#2c3e50", font=("Arial", 10), relief="flat")
        return btn

    def _create_styled_label(self, parent, text, style_type="glass", **kwargs):
        """Cria um label com estilo apropriado"""
        if self.using_liquid_glass and self.style_manager:
            try:
                if style_type == "title":
                    return self.style_manager.create_title_label(parent, text, **kwargs)
                else:
                    return ttk.Label(parent, text=text, style="Glass.TLabel", **kwargs)
            except Exception as e:
                print(f"⚠️ Erro ao criar label estilizado: {e}")
                self.using_liquid_glass = False
        
        # Fallback para label padrão
        label = tk.Label(parent, text=text, **kwargs)
        if style_type == "title":
            label.configure(font=("Arial", 16, "bold"), fg="#2c3e50", bg='#f5f5f5')
        else:
            label.configure(font=("Arial", 10), fg="#2c3e50", bg='#f5f5f5')
        return label

    def _create_styled_entry(self, parent, **kwargs):
        """Cria um entry com estilo apropriado"""
        if self.using_liquid_glass and self.style_manager:
            try:
                return self.style_manager.create_glass_entry(parent, **kwargs)
            except Exception as e:
                print(f"⚠️ Erro ao criar entry estilizado: {e}")
                self.using_liquid_glass = False
        
        # Fallback para entry padrão
        return tk.Entry(parent, **kwargs)

    def show(self):
        """Mostra a interface do módulo de captura"""
        if not self.root:
            self._create_interface()
        else:
            self.root.deiconify()
            self.root.lift()
            self.root.focus_set()

    def hide(self):
        """Esconde a interface do módulo de forma segura"""
        if self.root:
            try:
                # Parar qualquer gravação em andamento
                if self.gravando:
                    self.finalizar()
                
                # Parar listeners
                if hasattr(self, 'listener_keyboard') and self.listener_keyboard:
                    try:
                        self.listener_keyboard.stop()
                    except:
                        pass
                    self.listener_keyboard = None
                
                if self.listener_mouse:
                    try:
                        self.listener_mouse.stop()
                    except:
                        pass
                    self.listener_mouse = None
                
                # Fechar janelas secundárias
                if self.popup and self.popup.winfo_exists():
                    try:
                        self.popup.destroy()
                    except:
                        pass
                    self.popup = None
                
                # Liberar grabs
                try:
                    self.root.grab_release()
                except:
                    pass
                    
                # Esconder a janela
                self.root.withdraw()
                
                # 🔥 CORREÇÃO EXTRA: Voltar o foco para a janela principal
                if self.parent and self.parent.winfo_exists():
                    try:
                        self.parent.deiconify()
                        self.parent.lift()
                        self.parent.focus_force()
                    except:
                        pass
                    
            except Exception as e:
                print(f"Erro ao esconder módulo: {e}")
                # Fallback: destruir completamente se houver problemas
                try:
                    self.root.destroy()
                    self.root = None
                except:
                    pass

    def _create_interface(self):
        """Cria a interface do módulo de captura"""
        self.root = tk.Toplevel(self.parent)
        self.root.title("PrintF - Capturar Evidências")
        self.root.geometry("500x400")
        
        # 🔥 APLICAR ESTILO À JANELA
        self._apply_style_to_window(self.root)
        
        # 🔥 CORREÇÃO: Usar protocolo correto para fechar
        self.root.protocol("WM_DELETE_WINDOW", self.hide)
        self.root.resizable(False, False)

        # 🔥 CRIAR FRAME PRINCIPAL COM ESTILO
        main_frame = self._create_styled_frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # Interface com estilos aplicados
        self._create_styled_label(main_frame, text="📷 PrintF - Capturar Evidências", 
                                style_type="title").pack(pady=20)
        
        self._create_styled_button(main_frame, text="▶ Iniciar Gravação (F8)", 
                                 command=self.iniciar, style_type="accent").pack(pady=8, fill=tk.X)
        
        self._create_styled_button(main_frame, text="⏸ Pausar Gravação (F6)", 
                                 command=self.pausar, style_type="glass").pack(pady=8, fill=tk.X)
        
        self._create_styled_button(main_frame, text="▶ Retomar Gravação (F7)", 
                                 command=self.retomar, style_type="glass").pack(pady=8, fill=tk.X)
        
        self._create_styled_button(main_frame, text="⏹ Finalizar Gravação (F9)", 
                                 command=self.finalizar, style_type="glass").pack(pady=8, fill=tk.X)
        
        # 🔥 BOTÃO VOLTAR COM ESTILO DE ERRO (vermelho)
        if self.using_liquid_glass and self.style_manager:
            try:
                voltar_btn = ttk.Button(main_frame, text="⬅ Voltar ao Menu", 
                                      command=self.hide,
                                      style="Back.TButton")
                voltar_btn.pack(pady=15, fill=tk.X)
            except:
                # Fallback se o estilo Error não estiver disponível
                self._create_styled_button(main_frame, text="⬅ Voltar ao Menu", 
                                         command=self.hide, style_type="glass").pack(pady=15, fill=tk.X)
        else:
            voltar_btn = tk.Button(main_frame, text="⬅ Voltar ao Menu", 
                                 command=self.hide,
                                 bg="#e7b13c", fg="white", font=("Arial", 11), relief="flat")
            voltar_btn.pack(pady=15, fill=tk.X)

        # Configurar atalhos
        self._setup_shortcuts()

    def _setup_shortcuts(self):
        """Configura atalhos de teclado"""
        def on_press(key):
            try:
                if key == keyboard.Key.f6:
                    self.root.after(0, self.pausar)
                elif key == keyboard.Key.f7:
                    self.root.after(0, self.retomar)
                elif key == keyboard.Key.f8:
                    self.root.after(0, self.iniciar)
                elif key == keyboard.Key.f9:
                    self.root.after(0, self.finalizar)
            except Exception as e:
                print(f"Erro no atalho: {e}")

        # 🔥 CORREÇÃO: Parar listener anterior se existir
        if hasattr(self, 'listener_keyboard') and self.listener_keyboard:
            try:
                self.listener_keyboard.stop()
            except:
                pass
                
        self.listener_keyboard = keyboard.Listener(on_press=on_press, suppress=False)
        self.listener_keyboard.start()

    def _salvar_metadata(self):
        """Salva os metadados no arquivo JSON"""
        if self.metadata_path:
            with open(self.metadata_path, 'w', encoding='utf-8') as f:
                json.dump(self.metadata, f, indent=2, ensure_ascii=False)

    def carregar_evidencias(self, dir_path):
        """Carrega as evidências baseadas nos metadados"""
        self.metadata_path = os.path.join(dir_path, "evidencias_metadata.json")
        
        if os.path.exists(self.metadata_path):
            try:
                with open(self.metadata_path, 'r', encoding='utf-8') as f:
                    self.metadata = json.load(f)
            except:
                self.metadata = {"evidencias": [], "proximo_id": 1}
        
        # Carrega evidências ativas (não excluídas)
        evidencias_ativas = []
        for evidencia in self.metadata["evidencias"]:
            if not evidencia.get("excluida", False):
                caminho = os.path.join(dir_path, evidencia["arquivo"])
                if os.path.exists(caminho):
                    evidencias_ativas.append(caminho)
        
        return evidencias_ativas

    def recarregar_evidencias(self):
        """Recarrega a lista de evidências"""
        if self.evidence_dir:
            self.prints = self.carregar_evidencias(self.evidence_dir)
            return True
        return False

    # 🔥 MÉTODOS DE CAPTURA SIMPLIFICADOS E OTIMIZADOS
    def capture_inteligente(self, x, y):
        """
        Captura a tela baseado no modo selecionado pelo usuário
        """
        if self.modo_captura == "manter":
            # Modo manter: captura tela COMPLETA (incluindo barra de tarefas)
            return self.capture_tela_completa_mss(x, y)
        else:
            # Modo ocultar: captura apenas área de trabalho (sem barra)
            return self.capture_work_area_pyautogui(x, y)

    def capture_tela_completa_mss(self, x, y):
        """
        Captura a tela completa INCLUINDO a barra de tarefas.
        Funciona no primário e secundário, mesmo com coordenadas negativas.
        """
        try:
            # 🔥 ESTRATÉGIA 1: Win32 API para captura precisa de monitor específico
            if WIN32_AVAILABLE:
                try:
                    # Encontrar o monitor que contém o ponto (x, y)
                    monitor_handle = win32api.MonitorFromPoint((x, y), win32con.MONITOR_DEFAULTTONEAREST)
                    monitor_info = win32gui.GetMonitorInfo(monitor_handle)
                    
                    # Área completa do monitor (inclui barra)
                    monitor_area = monitor_info["Monitor"]  # (left, top, right, bottom)
                    
                    # Capturar usando MSS para melhor compatibilidade com múltiplos monitores
                    if MSS_AVAILABLE:
                        with mss.mss() as sct:
                            monitor_mss = {
                                "left": monitor_area[0],
                                "top": monitor_area[1], 
                                "width": monitor_area[2] - monitor_area[0],
                                "height": monitor_area[3] - monitor_area[1]
                            }
                            screenshot = sct.grab(monitor_mss)
                            img = Image.frombytes("RGB", screenshot.size, screenshot.bgra, "raw", "BGRX")
                            
                            rel_x = x - monitor_area[0]
                            rel_y = y - monitor_area[1]
                            
                            metodo_utilizado = f"Win32 + MSS Monitor Completo {monitor_area}"
                            print(f"✅ CAPTURA WIN32+MSS - Monitor {monitor_area} | Coord: ({rel_x},{rel_y})")
                            
                            return img, (rel_x, rel_y), metodo_utilizado
                    else:
                        # Fallback para ImageGrab se MSS não disponível
                        screenshot = ImageGrab.grab(bbox=monitor_area)
                        rel_x = x - monitor_area[0]
                        rel_y = y - monitor_area[1]
                        
                        metodo_utilizado = f"Win32 Monitor Completo {monitor_area}"
                        print(f"✅ CAPTURA WIN32 - Monitor {monitor_area} | Coord: ({rel_x},{rel_y})")
                        
                        return screenshot, (rel_x, rel_y), metodo_utilizado
                        
                except Exception as e:
                    print(f"⚠️  Win32 falhou (capturando com alternativa): {e}")

            # 🔥 ESTRATÉGIA 2: MSS como alternativa principal
            if MSS_AVAILABLE:
                try:
                    with mss.mss() as sct:
                        monitor_encontrado = None
                        
                        # Procurar em todos os monitores (exceto o virtual)
                        for monitor in sct.monitors[1:]:
                            if (monitor["left"] <= x < monitor["left"] + monitor["width"] and
                                monitor["top"] <= y < monitor["top"] + monitor["height"]):
                                monitor_encontrado = monitor
                                break

                        # Fallback para primeiro monitor se não encontrou
                        if not monitor_encontrado:
                            monitor_encontrado = sct.monitors[1] if len(sct.monitors) > 1 else sct.monitors[0]
                            print(f"⚠️  Monitor não encontrado para coordenadas ({x},{y}), usando monitor {monitor_encontrado} como fallback")

                        # Capturar a tela completa do monitor encontrado
                        screenshot = sct.grab(monitor_encontrado)
                        img = Image.frombytes("RGB", screenshot.size, screenshot.bgra, "raw", "BGRX")

                        # Calcular coordenadas relativas ao monitor
                        rel_x = x - monitor_encontrado["left"]
                        rel_y = y - monitor_encontrado["top"]

                        metodo_utilizado = f"MSS Monitor Completo {monitor_encontrado['width']}x{monitor_encontrado['height']}"
                        print(f"✅ CAPTURA MSS - Monitor {monitor_encontrado} | Coord: ({rel_x},{rel_y})")

                        return img, (rel_x, rel_y), metodo_utilizado
                        
                except Exception as e:
                    print(f"⚠️  MSS falhou (capturando com alternativa): {e}")

            # 🔥 ESTRATÉGIA 3: Fallback com ImageGrab
            try:
                screenshot = ImageGrab.grab()
                metodo_utilizado = "Fallback - ImageGrab (tela completa)"
                print(f"⚠️  Usando fallback ImageGrab para coordenadas ({x},{y})")
                return screenshot, (x, y), metodo_utilizado
                
            except Exception as e:
                print(f"⚠️  ImageGrab falhou: {e}")

            # 🔥 ESTRATÉGIA 4: Fallback final com pyautogui
            try:
                screenshot = pyautogui.screenshot()
                metodo_utilizado = "Fallback - pyautogui (apenas primário)"
                print(f"⚠️  Usando fallback pyautogui para coordenadas ({x},{y})")
                return screenshot, (x, y), metodo_utilizado
                
            except Exception as e:
                print(f"❌ Todos os métodos de captura falharam: {e}")
                raise

        except Exception as e:
            print(f"❌ Falha crítica na captura completa: {e}")
            # Último recurso - retorna imagem preta ou levanta exceção
            try:
                # Tenta criar uma imagem preta como fallback extremo
                img = Image.new('RGB', (100, 100), color='black')
                return img, (0, 0), f"Fallback Extremo - Erro: {str(e)}"
            except:
                raise Exception(f"Falha completa na captura de tela: {str(e)}")

    def capture_work_area_pyautogui(self, x, y):
        """
        Captura apenas a área de trabalho (SEM barra de tarefas) usando pyautogui
        Este método é usado no modo "ocultar"
        """
        try:
            # Captura com pyautogui (já captura sem a barra automaticamente)
            screenshot = pyautogui.screenshot()
            
            metodo_utilizado = "PyAutoGUI - Área de Trabalho (sem barra)"
            print(f"✅ CAPTURA PYAUTOGUI - Área de Trabalho | Coord: ({x},{y})")
            
            return screenshot, (x, y), metodo_utilizado
            
        except Exception as e:
            print(f"❌ Falha na captura com pyautogui: {e}")
            # Fallback extremo
            screenshot = pyautogui.screenshot()
            return screenshot, (x, y), f"Fallback - Erro: {str(e)}"

    def estimativa_segura_barra_tarefas(self, altura_tela):
        """
        Estimativa conservadora da altura da barra de tarefas
        """
        if altura_tela >= 2160:  # 4K
            return 100
        elif altura_tela >= 1440:  # QHD
            return 80
        elif altura_tela >= 1080:  # Full HD
            return 70
        else:  # Resoluções menores
            return 60
        
    # 🔥 NOVA FUNÇÃO: APLICAR TIMESTAMP MODERNO COM FUNDO
    def aplicar_timestamp_moderno(self, caminho_imagem, evidencia_meta):
        """Aplica o timestamp com fundo semi-transparente e texto branco"""
        img = Image.open(caminho_imagem).convert("RGBA")
        draw = ImageDraw.Draw(img)
        
        # Calcular posição em pixels
        img_width, img_height = img.size
        pos_x = int(evidencia_meta["timestamp_posicao"]["x"] * img_width)
        pos_y = int(evidencia_meta["timestamp_posicao"]["y"] * img_height)
        
        # Configurações do texto
        texto = evidencia_meta["timestamp_texto"]
        texto_cor = evidencia_meta["timestamp_cor"]  # Branco
        fundo_cor = evidencia_meta.get("timestamp_fundo", "#000000B2")  # Preto 70%
        tamanho = evidencia_meta["timestamp_tamanho"]
        
        # Converter cor de fundo para RGBA
        if fundo_cor.startswith("#") and len(fundo_cor) == 9:  # Formato #RRGGBBAA
            r = int(fundo_cor[1:3], 16)
            g = int(fundo_cor[3:5], 16)
            b = int(fundo_cor[5:7], 16)
            a = int(fundo_cor[7:9], 16)
            fundo_rgba = (r, g, b, a)
        else:
            fundo_rgba = (0, 0, 0, 178)  # Fallback: preto 70%
        
        # Usar fonte
        try:
            font = ImageFont.truetype("arial.ttf", tamanho)
        except:
            font = ImageFont.load_default()
        
        # 🔥 CALCULAR TAMANHO DO TEXTO PARA CRIAR FUNDO
        bbox = draw.textbbox((0, 0), texto, font=font)
        texto_largura = bbox[2] - bbox[0]
        texto_altura = bbox[3] - bbox[1]
        
        # 🔥 DEFINIR PADDING E CANTOS ARREDONDADOS
        padding = 10
        borda_radius = 8
        
        # Coordenadas do fundo
        fundo_x1 = pos_x - padding
        fundo_y1 = pos_y - padding
        fundo_x2 = pos_x + texto_largura + padding
        fundo_y2 = pos_y + texto_altura + padding
        
        # 🔥 DESENHAR FUNDO COM CANTOS ARREDONDADOS
        # Criar máscara para cantos arredondados
        mask = Image.new("L", (fundo_x2 - fundo_x1, fundo_y2 - fundo_y1), 0)
        mask_draw = ImageDraw.Draw(mask)
        mask_draw.rounded_rectangle(
            [0, 0, fundo_x2 - fundo_x1, fundo_y2 - fundo_y1],
            radius=borda_radius,
            fill=255
        )
        
        # Aplicar fundo semi-transparente
        fundo_img = Image.new("RGBA", (fundo_x2 - fundo_x1, fundo_y2 - fundo_y1), fundo_rgba)
        img.paste(fundo_img, (fundo_x1, fundo_y1), mask)
        
        # 🔥 DESENHAR TEXTO BRANCO (SEM BORDAS PRETAS)
        draw.text((pos_x, pos_y), texto, fill=texto_cor, font=font)
        
        # Salvar a imagem
        img.save(caminho_imagem)

    # 🔥 NOVA FUNÇÃO: CRIAR PASTA AUTOMÁTICA NO DIRETÓRIO DO TEMPLATE
    def _criar_pasta_automatica(self):
        """Cria uma pasta automática no mesmo diretório do template DOCX com o nome do documento"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # 🔥 OBTER NOME DO TEMPLATE SEM EXTENSÃO
        template_filename = os.path.basename(self.template_path)
        template_name = os.path.splitext(template_filename)[0]
        
        # 🔥 CORREÇÃO: LIMPAR NOME DO TEMPLATE PARA EVITAR CARACTERES INVÁLIDOS
        template_name = self._limpar_nome_arquivo(template_name)
        
        # 🔥 CRIAR NOME DA PASTA COM NOME DO DOCX + TIMESTAMP
        nome_pasta = f"{template_name}_{timestamp}"
        
        # 🔥 CRIAR PASTA NO MESMO DIRETÓRIO DO TEMPLATE
        template_dir = os.path.dirname(self.template_path)
        pasta_automatica = os.path.join(template_dir, nome_pasta)
        
        # Criar a pasta
        os.makedirs(pasta_automatica, exist_ok=True)
        print(f"📁 Pasta automática criada: {pasta_automatica}")
        
        return pasta_automatica

    # 🔥 CORREÇÃO: FUNÇÃO PARA LIMPAR NOME DE ARQUIVO MANTENDO CARACTERES PT-BR
    def _limpar_nome_arquivo(self, nome):
        """Remove caracteres inválidos para nomes de arquivo no Windows, mantendo caracteres PT-BR"""
        # Caracteres inválidos no Windows: \ / : * ? " < > |
        caracteres_invalidos = r'[\\/*?:"<>|]'
        nome_limpo = re.sub(caracteres_invalidos, '_', nome)
        
        # 🔥 CORREÇÃO: Permitir caracteres acentuados e especiais do português
        # Manter letras acentuadas, ç, ñ, e outros caracteres comuns no PT-BR
        # Esta regex mantém: letras (incluindo acentuadas), números, espaços, hífens, underscores, pontos e parênteses
        nome_limpo = re.sub(r'[^\w\s\-\.\(\)áàâãéèêíïóôõöúçñÁÀÂÃÉÈÊÍÏÓÔÕÖÚÇÑ]', '', nome_limpo)
        
        # 🔥 CORREÇÃO ADICIONAL: LIMITAR TAMANHO DO NOME PARA EVITAR CAMINHOS MUITO LONGOS
        if len(nome_limpo) > 100:  # Aumentado para 100 caracteres
            nome_limpo = nome_limpo[:100]
            
        return nome_limpo.strip()

    # ---------- Nova janela de configuração ----------
    def mostrar_janela_configuracao(self):
        config_window = tk.Toplevel(self.root)
        config_window.title("Configuração de Gravação")
        config_window.geometry("600x600")
        config_window.resizable(False, False)
        
        # 🔥 APLICAR ESTILO À JANELA
        self._apply_style_to_window(config_window)
        
        # 🔥 CORREÇÃO: Usar transient mas SEM grab_set
        config_window.transient(self.root)
        
        # 🔥 CRIAR FRAME PRINCIPAL COM ESTILO
        main_frame = self._create_styled_frame(config_window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        self._create_styled_label(main_frame, text="PrintF - Configuração de Gravação", 
                                style_type="title").pack(pady=10)
        
        # Seleção de template
        self._create_styled_label(main_frame, text="Selecione o template DOCX:").pack(anchor="w", pady=(10, 5))
        
        template_frame = self._create_styled_frame(main_frame)
        template_frame.pack(fill=tk.X, pady=5)
        
        self.template_var = tk.StringVar()
        template_entry = self._create_styled_entry(template_frame, textvariable=self.template_var, width=40)
        template_entry.pack(side=tk.LEFT, padx=(0, 5), fill=tk.X, expand=True)
        
        def selecionar_template():
            template_path = filedialog.askopenfilename(
                title="Selecione o template DOCX",
                filetypes=[("Documentos Word", "*.docx")]
            )
            if template_path:
                self.template_var.set(template_path)
        
        self._create_styled_button(template_frame, text="Procurar", command=selecionar_template).pack(side=tk.RIGHT)
        
        # 🔥 MODIFICADO: Seleção de diretório de destino (OPCIONAL)
        self._create_styled_label(main_frame, text="Selecione o diretório de destino (opcional):").pack(anchor="w", pady=(20, 5))
        
        # 🔥 ADICIONADO: Label informativo sobre pasta automática
        if self.using_liquid_glass and self.style_manager:
            info_label = ttk.Label(
                main_frame, 
                text="Se não selecionar um destino, será criada uma pasta automaticamente no mesmo diretório do template com o nome do documento.", 
                style="Subtitle.TLabel",
                justify=tk.LEFT
            )
        else:
            info_label = tk.Label(
                main_frame, 
                text="Se não selecionar um destino, será criada uma pasta automaticamente no mesmo diretório do template com o nome do documento.", 
                font=("Arial", 9), 
                foreground="gray",
                justify=tk.LEFT,
                bg='#f5f5f5'
            )
        info_label.pack(anchor="w", pady=(0, 10))
        
        dir_frame = self._create_styled_frame(main_frame)
        dir_frame.pack(fill=tk.X, pady=5)
        
        self.dir_var = tk.StringVar()
        dir_entry = self._create_styled_entry(dir_frame, textvariable=self.dir_var, width=40)
        dir_entry.pack(side=tk.LEFT, padx=(0, 5), fill=tk.X, expand=True)
        
        def selecionar_diretorio():
            dir_path = filedialog.askdirectory(title="Selecione o diretório para salvar")
            if dir_path:
                self.dir_var.set(dir_path)
        
        self._create_styled_button(dir_frame, text="Procurar", command=selecionar_diretorio).pack(side=tk.RIGHT)
        
        # 🔥 NOVO: Seleção do modo de captura (APENAS 2 OPÇÕES)
        self._create_styled_label(main_frame, text="Modo de Captura da Barra de Tarefas:", 
                                style_type="title").pack(anchor="w", pady=(20, 10))
        
        # Variável para os RadioButtons
        self.modo_captura_var = tk.StringVar(value="manter")  # Valor padrão
        
        # Frame para os RadioButtons
        modo_frame = self._create_styled_frame(main_frame)
        modo_frame.pack(fill=tk.X, pady=5)
        
        # RadioButton 1: Manter barra completa
        if self.using_liquid_glass and self.style_manager:
            rb1 = ttk.Radiobutton(
                modo_frame, 
                text="Manter barra de tarefas (data/hora visível na barra do Windows)",
                variable=self.modo_captura_var, 
                value="manter",
                style="Glass.TRadiobutton"
            )
        else:
            rb1 = tk.Radiobutton(
                modo_frame, 
                text="Manter barra de tarefas (data/hora visível na barra do Windows)",
                variable=self.modo_captura_var, 
                value="manter",
                bg='#f5f5f5'
            )
        rb1.pack(anchor="w", pady=2)
        
        # RadioButton 2: Ocultar barra
        if self.using_liquid_glass and self.style_manager:
            rb2 = ttk.Radiobutton(
                modo_frame, 
                text="Ocultar barra de tarefas (data/hora será adicionada na imagem)",
                variable=self.modo_captura_var, 
                value="ocultar",
                style="Glass.TRadiobutton"
            )
        else:
            rb2 = tk.Radiobutton(
                modo_frame, 
                text="Ocultar barra de tarefas (data/hora será adicionada na imagem)",
                variable=self.modo_captura_var, 
                value="ocultar",
                bg='#f5f5f5'
            )
        rb2.pack(anchor="w", pady=2)
        
        # Checkbox para manter evidências
        self._create_styled_label(main_frame, text="Opções de saída:", style_type="title").pack(anchor="w", pady=(20, 10))
        
        # Variável para o checkbox - valor padrão True (marcado)
        self.manter_evidencias_var = tk.BooleanVar(value=True)
        
        # Checkbox
        checkbox_frame = self._create_styled_frame(main_frame)
        checkbox_frame.pack(fill=tk.X, pady=5)
        
        if self.using_liquid_glass and self.style_manager:
            manter_checkbox = ttk.Checkbutton(
                checkbox_frame, 
                text="Manter arquivos de evidência (prints) após gerar o DOCX",
                variable=self.manter_evidencias_var,
                style="Glass.TCheckbutton"
            )
        else:
            manter_checkbox = tk.Checkbutton(
                checkbox_frame, 
                text="Manter arquivos de evidência (prints) após gerar o DOCX",
                variable=self.manter_evidencias_var,
                bg='#f5f5f5'
            )
        manter_checkbox.pack(anchor="w")
        
        # Label informativa
        if self.using_liquid_glass and self.style_manager:
            info_label = ttk.Label(
                main_frame, 
                text="Se desmarcado, os arquivos de print serão excluídos após a geração do DOCX.\nSe foi criada uma pasta automática, ela também será excluída.", 
                style="Subtitle.TLabel",
                justify=tk.LEFT
            )
        else:
            info_label = tk.Label(
                main_frame, 
                text="Se desmarcado, os arquivos de print serão excluídos após a geração do DOCX.\nSe foi criada uma pasta automática, ela também será excluída.", 
                font=("Arial", 9), 
                foreground="gray",
                justify=tk.LEFT,
                bg='#f5f5f5'
            )
        info_label.pack(anchor="w", pady=(5, 15))
        
        # Frame para os botões na parte inferior
        btn_frame = self._create_styled_frame(main_frame)
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(20, 0))
        
        def iniciar_com_config():
            if not self.template_var.get():
                messagebox.showerror("Erro", "Por favor, selecione o template DOCX.")
                return
            
            template_path = self.template_var.get()
            if not os.path.exists(template_path):
                messagebox.showerror("Erro", "O arquivo de template selecionado não existe.")
                return
            
            # 🔥 CORREÇÃO: DEFINIR template_path ANTES de usar
            self.template_path = template_path
            
            # 🔥 VERIFICAÇÃO DO DIRETÓRIO DE DESTINO
            dir_selecionado = self.dir_var.get().strip()
            
            if dir_selecionado:
                # Usuário selecionou um diretório manualmente
                if os.path.exists(dir_selecionado):
                    try:
                        # Verificar se existe algum arquivo na pasta raiz
                        for item in os.listdir(dir_selecionado):
                            item_path = os.path.join(dir_selecionado, item)
                            # Ignorar ocultos e verificar apenas arquivos (não pastas)
                            if not item.startswith('.') and os.path.isfile(item_path):
                                messagebox.showerror(
                                    "Arquivos na Pasta", 
                                    f"A pasta selecionada contém arquivos.\n\n"
                                    f"Para evitar misturar evidências, a pasta deve estar vazia "
                                    f"ou conter apenas outras pastas.\n\n"                                
                                )
                                return
                                
                    except PermissionError:
                        messagebox.showerror("Erro de Permissão", "Sem permissão para acessar a pasta selecionada.")
                        return
                    except Exception as e:
                        messagebox.showerror("Erro", f"Erro ao verificar a pasta: {str(e)}")
                        return
                
                self.output_dir = dir_selecionado
                self.evidence_dir = dir_selecionado
                self.pasta_automatica = False
                self.pasta_automatica_path = None
            else:
                # 🔥 CRIAR PASTA AUTOMÁTICA NO DIRETÓRIO DO TEMPLATE
                self.pasta_automatica_path = self._criar_pasta_automatica()
                self.output_dir = self.pasta_automatica_path
                self.evidence_dir = self.pasta_automatica_path
                self.pasta_automatica = True
            
            # 🔥 Armazena a escolha do modo de captura
            self.modo_captura = self.modo_captura_var.get()
            
            # 🔥 VERIFICAÇÃO ADICIONAL: Limpar qualquer estado residual
            self.gravando = False
            self.pausado = False
            self.prints = []
            
            # Armazenar a escolha do usuário
            self.manter_evidencias = self.manter_evidencias_var.get()
            
            # 🔥 CORREÇÃO: Fechar corretamente
            config_window.destroy()  # Usar destroy em vez de apenas fechar
              
            self.iniciar_gravacao()
        
        # Centralizar os botões horizontalmente
        button_container = self._create_styled_frame(btn_frame)
        button_container.pack(expand=True)
        
        self._create_styled_button(button_container, text="Iniciar Gravação", 
                                 command=iniciar_com_config, style_type="accent").pack(side=tk.LEFT, padx=10)
        
        def cancelar_config():
            config_window.destroy()
            
        self._create_styled_button(button_container, text="Cancelar", 
                                 command=cancelar_config, style_type="glass").pack(side=tk.LEFT, padx=10)
        
        # Forçar atualização da interface e ajustar tamanho se necessário
        config_window.update_idletasks()
        
        # Se a janela for muito grande para a tela, ajustar
        screen_width = config_window.winfo_screenwidth()
        screen_height = config_window.winfo_screenheight()
        
        if config_window.winfo_height() > screen_height:
            config_window.geometry(f"600x{screen_height-100}")
        
        # 🔥 CORREÇÃO: Não usar wait_window que pode travar
        return self.template_path is not None

    def iniciar(self):
        """Inicia o processo de configuração da gravação"""
        # 🔥 CORREÇÃO: Resetar estado ANTES de iniciar
        self.gravando = False
        self.pausado = False
        self.prints = []
        self.evidencia_count = 0
        self.pasta_automatica = False
        self.pasta_automatica_path = None
        
        # Mostrar janela de configuração
        if self.mostrar_janela_configuracao():
            print("✅ Configuração concluída, iniciando gravação...")
        else:
            print("❌ Configuração cancelada pelo usuário")

    def pausar(self):
        if self.gravando and not self.pausado:
            self.pausado = True
            messagebox.showinfo("Gravação Pausada", "Gravação pausada. Clique em Retomar para continuar.")
        else:
            messagebox.showwarning("Aviso", "Gravação não está ativa ou já está pausada.")

    def retomar(self):
        if self.gravando and self.pausado:
            self.pausado = False
            messagebox.showinfo("Gravação Retomada", "Gravação retomada. Continue clicando para capturar telas.")
        else:
            messagebox.showwarning("Aviso", "Gravação não está pausada.")

    def finalizar(self):
        if not self.gravando:
            messagebox.showwarning("Aviso", "Nenhuma gravação em andamento.")
            return

        # Parar listener do mouse
        if self.listener_mouse:
            self.listener_mouse.stop()
            self.listener_mouse = None

        self.gravando = False
        self.pausado = False

        # Fechar popup se estiver aberto
        if self.popup and self.popup.winfo_exists():
            self.popup.destroy()
            self.popup = None

        # Gerar documento ou mostrar navegação
        if self.prints:
            self.mostrar_janela_navegacao()
        else:
            messagebox.showwarning("Aviso", "Nenhuma evidência capturada.")

    def iniciar_gravacao(self):
        """Inicia a gravação após configuração"""
        # Criar diretório de evidências se não existir
        if not os.path.exists(self.evidence_dir):
            os.makedirs(self.evidence_dir)

        # Inicializar metadados
        self.metadata_path = os.path.join(self.evidence_dir, "evidencias_metadata.json")
        if os.path.exists(self.metadata_path):
            try:
                with open(self.metadata_path, 'r', encoding='utf-8') as f:
                    self.metadata = json.load(f)
            except:
                self.metadata = {"evidencias": [], "proximo_id": 1}
        else:
            self.metadata = {"evidencias": [], "proximo_id": 1}

        # Carregar template
        try:
            if os.path.exists(self.template_path):
                self.doc = Document(self.template_path)
                self.using_template = True
                print("Template carregado com sucesso!")
            else:
                self.doc = Document()
                self.using_template = False
                print("Template não encontrado. Criando documento vazio.")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar template: {str(e)}")
            self.doc = Document()
            self.using_template = False

        # Iniciar gravação
        self.gravando = True
        self.pausado = False

        # Configurar listener do mouse
        def on_click(x, y, button, pressed):
            if pressed and button == mouse.Button.left and self.gravando and not self.pausado:
                self.capturar_tela(x, y)

        if self.listener_mouse:
            self.listener_mouse.stop()

        self.listener_mouse = mouse.Listener(on_click=on_click)
        self.listener_mouse.start()

        # Mostrar feedback
        self.mostrar_janela_feedback()
        
        # 🔥 MENSAGEM ATUALIZADA COM INFORMAÇÃO SOBRE PASTA AUTOMÁTICA
        mensagem = "✅ Gravação iniciada com sucesso!\n\n"
        mensagem += "Clique com o botão esquerdo do mouse para capturar telas.\n\n"
        mensagem += "Atalhos disponíveis:\n"
        mensagem += "• F6: Pausar gravação\n"
        mensagem += "• F7: Retomar gravação\n" 
        mensagem += "• F9: Finalizar gravação\n\n"
        
        if self.pasta_automatica:
            mensagem += f"📁 Pasta automática criada:\n{self.pasta_automatica_path}"
        
        messagebox.showinfo("Gravação Iniciada", mensagem)

    def capturar_tela(self, x, y):
        """Captura a tela e salva a evidência"""
        try:
            # 🔥 CAPTURA INTELIGENTE BASEADA NO MODO SELECIONADO
            screenshot, (rel_x, rel_y), metodo_utilizado = self.capture_inteligente(x, y)
            
            # Gerar nome único para o arquivo
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"evidencia_{self.metadata['proximo_id']:04d}_{timestamp}.png"
            filepath = os.path.join(self.evidence_dir, filename)
            
            # Salvar a imagem
            screenshot.save(filepath, "PNG")
            
            # 🔥 ADICIONAR METADADOS DA EVIDÊNCIA
            evidencia_meta = {
                "id": self.metadata['proximo_id'],
                "arquivo": filename,
                "timestamp": timestamp,
                "coordenadas": {"x": x, "y": y},
                "coordenadas_relativas": {"x": rel_x, "y": rel_y},
                "metodo_captura": metodo_utilizado,
                "modo_captura": self.modo_captura,
                "comentario": "",
                "excluida": False,
                "timestamp_texto": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                "timestamp_cor": "#FFFFFF",  # Branco
                "timestamp_tamanho": 16,
                "timestamp_posicao": {"x": 0.02, "y": 0.02},  # Canto superior esquerdo
                "timestamp_fundo": "#000000B2"  # Preto 70%
            }
            
            # 🔥 APLICAR TIMESTAMP MODERNO
            self.aplicar_timestamp_moderno(filepath, evidencia_meta)
            
            # Atualizar metadados
            self.metadata["evidencias"].append(evidencia_meta)
            self.metadata["proximo_id"] += 1
            self._salvar_metadata()
            
            # Adicionar à lista de prints
            self.prints.append(filepath)
            self.evidencia_count += 1
            
            # Atualizar feedback
            if self.popup and self.popup.winfo_exists():
                try:
                    self.current_index = len(self.prints) - 1
                    self.atualizar_popup()
                except Exception as e:
                    print(f"Erro ao atualizar popup: {e}")
            
            print(f"✅ Captura {self.evidencia_count} salva: {filename}")
            
        except Exception as e:
            print(f"❌ Erro ao capturar tela: {e}")
            messagebox.showerror("Erro", f"Erro ao capturar tela: {e}")

    def mostrar_janela_feedback(self):
        """Mostra janela de feedback durante a gravação"""
        if self.popup and self.popup.winfo_exists():
            self.popup.destroy()
            
        self.popup = tk.Toplevel(self.root)
        self.popup.title("PrintF - Gravando...")
        self.popup.geometry("400x300")
        self.popup.resizable(False, False)
        
        # 🔥 APLICAR ESTILO À JANELA
        self._apply_style_to_window(self.popup)
        
        # 🔥 CORREÇÃO: Usar transient mas SEM grab_set
        self.popup.transient(self.root)
        
        # 🔥 CRIAR FRAME PRINCIPAL COM ESTILO
        main_frame = self._create_styled_frame(self.popup)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        self._create_styled_label(main_frame, text="📷 Gravando Evidências", 
                                style_type="title").pack(pady=10)
        
        # Status
        status_frame = self._create_styled_frame(main_frame)
        status_frame.pack(fill=tk.X, pady=10)
        
        self._create_styled_label(status_frame, text="Status:", 
                                style_type="glass").pack(side=tk.LEFT)
        
        self.status_label = self._create_styled_label(status_frame, text="▶ Gravando", 
                                                    style_type="glass")
        self.status_label.pack(side=tk.RIGHT)
        
        # Contador
        count_frame = self._create_styled_frame(main_frame)
        count_frame.pack(fill=tk.X, pady=10)
        
        self._create_styled_label(count_frame, text="Evidências capturadas:", 
                                style_type="glass").pack(side=tk.LEFT)
        
        self.count_label = self._create_styled_label(count_frame, text="0", 
                                                   style_type="glass")
        self.count_label.pack(side=tk.RIGHT)
        
        # Imagem atual
        img_frame = self._create_styled_frame(main_frame)
        img_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        self._create_styled_label(img_frame, text="Última captura:", 
                                style_type="glass").pack(anchor="w")
        
        # Container para a imagem (com tamanho fixo)
        img_container = self._create_styled_frame(img_frame)
        img_container.pack(fill=tk.BOTH, expand=True, pady=5)
        img_container.configure(height=150)
        
        self.current_img_label = tk.Label(img_container, bg="white", relief="solid", bd=1)
        self.current_img_label.pack(fill=tk.BOTH, expand=True)
        
        # Comentário
        comment_frame = self._create_styled_frame(main_frame)
        comment_frame.pack(fill=tk.X, pady=10)
        
        self._create_styled_label(comment_frame, text="Comentário (opcional):", 
                                style_type="glass").pack(anchor="w")
        
        self.comment_entry = self._create_styled_entry(comment_frame)
        self.comment_entry.pack(fill=tk.X, pady=5)
        self.comment_entry.bind("<Return>", lambda e: self.adicionar_comentario())
        
        self._create_styled_button(comment_frame, text="Adicionar Comentário", 
                                 command=self.adicionar_comentario, style_type="glass").pack(pady=5)
        
        # Botão finalizar
        btn_frame = self._create_styled_frame(main_frame)
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(10, 0))
        
        self._create_styled_button(btn_frame, text="⏹ Finalizar Gravação", 
                                 command=self.finalizar, style_type="accent").pack()
        
        self.atualizar_popup()

    def atualizar_popup(self):
        """Atualiza o popup de feedback"""
        if not self.popup or not self.popup.winfo_exists():
            return
            
        try:
            # Atualizar status
            status = "⏸ Pausada" if self.pausado else "▶ Gravando"
            self.status_label.config(text=status)
            
            # Atualizar contador
            self.count_label.config(text=str(self.evidencia_count))
            
            # Atualizar imagem se houver capturas
            if self.prints and self.current_index < len(self.prints):
                img_path = self.prints[self.current_index]
                
                # Carregar e redimensionar imagem
                img = Image.open(img_path)
                img.thumbnail((300, 150), Image.Resampling.LANCZOS)
                
                self.current_img_tk = ImageTk.PhotoImage(img)
                self.current_img_label.config(image=self.current_img_tk)
                
        except Exception as e:
            print(f"Erro ao atualizar popup: {e}")

    def adicionar_comentario(self):
        """Adiciona comentário à evidência atual"""
        if not self.prints or self.current_index >= len(self.prints):
            messagebox.showwarning("Aviso", "Nenhuma evidência selecionada.")
            return
            
        comentario = self.comment_entry.get().strip()
        if not comentario:
            messagebox.showwarning("Aviso", "Digite um comentário.")
            return
            
        try:
            # Atualizar metadados
            evidencia_id = self.metadata["evidencias"][self.current_index]["id"]
            for evidencia in self.metadata["evidencias"]:
                if evidencia["id"] == evidencia_id:
                    evidencia["comentario"] = comentario
                    break
                    
            self._salvar_metadata()
            
            # Limpar campo
            self.comment_entry.delete(0, tk.END)
            
            messagebox.showinfo("Sucesso", "Comentário adicionado com sucesso!")
            
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao adicionar comentário: {e}")

    # 🔥 ADICIONADO: MÉTODOS DE NAVEGAÇÃO E EDIÇÃO
    def mostrar_janela_navegacao(self):
        """Janela principal de navegação pelas evidências"""
        if self.popup and self.popup.winfo_exists():
            self.popup.destroy()

        self.popup = tk.Toplevel(self.root)
        self.popup.title("Navegação de Evidências")
        self.popup.geometry("1200x800")
        self.popup.resizable(True, True)
        
        # 🔥 APLICAR ESTILO À JANELA
        self._apply_style_to_window(self.popup)
        
        # 🔥 CORREÇÃO: Usar transient mas SEM grab_set
        self.popup.transient(self.root)
        
        # Configurar grid para melhor organização
        self.popup.grid_columnconfigure(0, weight=1)
        self.popup.grid_rowconfigure(0, weight=1)  # A área da imagem expande
        
        # Frame da imagem (maior para melhor visualização)
        img_frame = self._create_styled_frame(self.popup)
        img_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        img_frame.grid_rowconfigure(0, weight=1)
        img_frame.grid_columnconfigure(0, weight=1)
        
        self.current_img_label = tk.Label(img_frame, bg="white")
        self.current_img_label.grid(row=0, column=0, sticky="nsew")
        
        # Frame do comentário (abaixo da imagem)
        comment_frame = self._create_styled_frame(self.popup)
        comment_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=(0, 5))  # Reduzido espaçamento
                
        self._create_styled_label(comment_frame, text="Comentário:").pack(anchor="w")
        
        # Criar um frame para o campo de entrada
        comment_entry_frame = self._create_styled_frame(comment_frame)
        comment_entry_frame.pack(fill=tk.X, pady=2)  # Reduzido espaçamento
        
        # Campo de comentário
        self.comment_entry = tk.Entry(comment_entry_frame, font=("Arial", 10))
        self.comment_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.comment_entry.bind("<FocusOut>", lambda e: self.salvar_comentario())
        
        # Frame principal para os botões de navegação e ação
        buttons_main_frame = self._create_styled_frame(self.popup)
        buttons_main_frame.grid(row=2, column=0, sticky="nsew", padx=10, pady=5)
        
        # Frame para centralizar os botões de navegação
        nav_frame = self._create_styled_frame(buttons_main_frame)
        nav_frame.pack(expand=True, pady=2)  # Centralizado e com pouco espaçamento
        
        # Botões de navegação (centralizados)
        self._create_styled_button(nav_frame, text="⏮️ Primeira", command=self.primeira_evidencia, 
                                 style_type="glass").pack(side=tk.LEFT, padx=2)
        self._create_styled_button(nav_frame, text="◀️ Anterior", command=self.anterior_evidencia,
                                 style_type="glass").pack(side=tk.LEFT, padx=2)
        
        # Indicador de posição
        self.pos_label = tk.Label(nav_frame, text="", font=("Arial", 12, "bold"))
        self.pos_label.pack(side=tk.LEFT, padx=15)
        
        self._create_styled_button(nav_frame, text="▶️ Próxima", command=self.proxima_evidencia,
                                 style_type="glass").pack(side=tk.LEFT, padx=2)
        self._create_styled_button(nav_frame, text="⏭️ Última", command=self.ultima_evidencia,
                                 style_type="glass").pack(side=tk.LEFT, padx=2)
        
        # Pular para específica
        self._create_styled_button(nav_frame, text="🔢 Ir para...", command=self.ir_para_especifica,
                                 style_type="glass").pack(side=tk.LEFT, padx=2)
        
        # Botões de ação no mesmo nível (Editar e Excluir Print)
        action_frame = self._create_styled_frame(buttons_main_frame)
        action_frame.pack(expand=True, pady=2)
        
        self._create_styled_button(action_frame, text="✏️ Editar Print", command=self.editar_evidencia_atual,
                                 style_type="glass").pack(side=tk.LEFT, padx=5)
        self._create_styled_button(action_frame, text="🗑️ Excluir Print", command=self.excluir_evidencia_atual,
                                 style_type="glass").pack(side=tk.LEFT, padx=5)
        
        # Frame de controle (parte inferior)
        control_frame = self._create_styled_frame(self.popup)
        control_frame.grid(row=3, column=0, sticky="ew", padx=10, pady=5)
        
        # Frame para centralizar os botões de controle
        control_buttons_frame = self._create_styled_frame(control_frame)
        control_buttons_frame.pack(expand=True)
        
        # Botões na ordem solicitada: Cancelar primeiro, depois Gerar Evidência
        self._create_styled_button(control_buttons_frame, text="❌ Cancelar", command=self.cancelar_processamento,
                                 style_type="error").pack(side=tk.LEFT, padx=5)
        
        self._create_styled_button(control_buttons_frame, text="✅ Gerar Evidência", command=self.finalizar_processamento,
                                 style_type="accent").pack(side=tk.LEFT, padx=5)
        
        # Carregar primeira evidência
        self.current_index = 0
        self.atualizar_exibicao()
        
        self.popup.protocol("WM_DELETE_WINDOW", self.cancelar_processamento)

    def atualizar_exibicao(self):
        """Atualiza a exibição da evidência atual"""
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        
        try:
            # Carrega e exibe a imagem com tamanho maior
            img = Image.open(caminho_print)
            
            # Obter o tamanho da área disponível para a imagem
            self.popup.update()
            available_width = self.popup.winfo_width() - 40  # Margens
            available_height = self.popup.winfo_height() - 250  # Espaço para controles
            
            # Ajustar a imagem para caber na área disponível
            img.thumbnail((available_width, available_height))
            self.current_img_tk = ImageTk.PhotoImage(img)
            self.current_img_label.config(image=self.current_img_tk)
            
            # Atualiza indicador de posição
            self.pos_label.config(text=f"Evidência {self.current_index + 1} de {len(self.prints)}")
            
            # Carrega comentário salvo
            nome_arquivo = os.path.basename(caminho_print)
            comentario = self.obter_comentario(nome_arquivo)
            self.comment_entry.delete(0, tk.END)
            self.comment_entry.insert(0, comentario)
            
        except Exception as e:
            print(f"Erro ao carregar imagem: {e}")

    def obter_comentario(self, nome_arquivo):
        """Obtém o comentário salvo nos metadados"""
        for evidencia in self.metadata["evidencias"]:
            if evidencia["arquivo"] == nome_arquivo:
                return evidencia.get("comentario", "")
        return ""

    def salvar_comentario(self):
        """Salva o comentário da evidência atual"""
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        nome_arquivo = os.path.basename(caminho_print)
        comentario = self.comment_entry.get()
        
        # Atualiza metadados
        for evidencia in self.metadata["evidencias"]:
            if evidencia["arquivo"] == nome_arquivo:
                evidencia["comentario"] = comentario
                break
                
        self._salvar_metadata()        

    # Métodos de navegação
    def primeira_evidencia(self):
        self.salvar_comentario()  # Salva automaticamente antes de navegar
        self.current_index = 0
        self.atualizar_exibicao()

    def anterior_evidencia(self):
        self.salvar_comentario()  # Salva automaticamente antes de navegar
        if self.current_index > 0:
            self.current_index -= 1
            self.atualizar_exibicao()

    def proxima_evidencia(self):
        self.salvar_comentario()  # Salva automaticamente antes de navegar
        if self.current_index < len(self.prints) - 1:
            self.current_index += 1
            self.atualizar_exibicao()

    def ultima_evidencia(self):
        self.salvar_comentario()  # Salva automaticamente antes de navegar
        self.current_index = len(self.prints) - 1
        self.atualizar_exibicao()

    def ir_para_especifica(self):
        self.salvar_comentario()  # Salva automaticamente antes de navegar
        if not self.prints:
            return
            
        numero = simpledialog.askinteger("Navegar", 
                                       f"Digite o número da evidência (1-{len(self.prints)}):",
                                       minvalue=1, maxvalue=len(self.prints))
        if numero:
            self.current_index = numero - 1
            self.atualizar_exibicao()

    def editar_evidencia_atual(self):
        self.salvar_comentario()  # Salva automaticamente antes de navegar
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        self.abrir_editor(caminho_print, self.popup)
        # Recarrega a imagem após edição
        self.atualizar_exibicao()

    def excluir_evidencia_atual(self):
        self.salvar_comentario()  # Salva automaticamente antes de navegar
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        nome_arquivo = os.path.basename(caminho_print)
        
        if messagebox.askyesno("Confirmar Exclusão", 
                             "Tem certeza que deseja excluir este print?"):
            try:
                # Remove arquivo físico
                os.remove(caminho_print)
                
                # Marca como excluída nos metadados
                for evidencia in self.metadata["evidencias"]:
                    if evidencia["arquivo"] == nome_arquivo:
                        evidencia["excluida"] = True
                        break
                
                self._salvar_metadata()
                
                # Recarrega a lista de evidências
                self.recarregar_evidencias()
                
                if not self.prints:
                    messagebox.showinfo("Info", "Todas as evidências foram processadas.")
                    self.finalizar_processamento()
                    return
                
                # Ajusta o índice se necessário
                if self.current_index >= len(self.prints):
                    self.current_index = len(self.prints) - 1
                
                self.atualizar_exibicao()
                messagebox.showinfo("Sucesso", "Evidência excluída!")
                
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao excluir: {str(e)}")

    def finalizar_processamento(self):
        """Processa todas as evidências e gera o DOCX"""
        self.salvar_comentario()  # Salva automaticamente antes de navegar
        
        # Gerar documento
        try:
            doc_path = self.gerar_documento()
            
            # 🔥 ADICIONADO: ABRIR PASTA APÓS GERAR DOCUMENTO
            pasta_para_abrir = os.path.dirname(doc_path)
            
            resposta = messagebox.askyesno(
                "Sucesso", 
                f"Documento gerado com sucesso em:\n{doc_path}\n\nDeseja abrir a pasta onde o documento foi salvo?",
                parent=self.popup
            )
            
            if resposta:
                if not self._abrir_pasta(pasta_para_abrir):
                    messagebox.showinfo(
                        "Abrir Pasta", 
                        f"Pasta do documento:\n{pasta_para_abrir}",
                        parent=self.popup
                    )
                    
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar documento: {e}", parent=self.popup)
        
        # Fechar janela de navegação
        if self.popup and self.popup.winfo_exists():
            self.popup.destroy()
            self.popup = None

    def _abrir_pasta(self, caminho_pasta):
        """Abre a pasta no explorador de arquivos do sistema"""
        try:
            if os.name == 'nt':  # Windows
                os.startfile(caminho_pasta)
            elif os.name == 'posix':  # Linux ou macOS
                if sys.platform == 'darwin':  # macOS
                    subprocess.run(['open', caminho_pasta])
                else:  # Linux
                    subprocess.run(['xdg-open', caminho_pasta])
            return True
        except Exception as e:
            print(f"Erro ao abrir pasta: {e}")
            return False

    def cancelar_processamento(self):
        self.salvar_comentario()  # Salva automaticamente ao fechar
        if messagebox.askyesno("Confirmar", "Deseja cancelar o processamento?"):
            if self.popup:
                self.popup.destroy()
                self.popup = None

    # ---------- Editor de prints ----------
    def abrir_editor(self, caminho_print, parent):
        editor = tk.Toplevel(parent)
        editor.title("Editor de Evidência")
        editor.geometry("1200x800")
        
        # 🔥 APLICAR ESTILO À JANELA
        self._apply_style_to_window(editor)
        
        # Frame principal
        main_frame = self._create_styled_frame(editor)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Frame para ferramentas e opções
        tools_frame = self._create_styled_frame(main_frame)
        tools_frame.pack(side=tk.TOP, fill=tk.X, pady=5)
        
        # Frame para a área de desenho
        canvas_frame = self._create_styled_frame(main_frame)
        canvas_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        # Carrega a imagem original
        self.original_img = Image.open(caminho_print).convert("RGBA")
        img_w, img_h = self.original_img.size
        
        # Calcula o fator de escala para exibição
        max_w, max_h = 1000, 700
        scale = min(max_w / img_w, max_h / img_h)
        self.scale_factor = scale
        disp_w, disp_h = int(img_w * scale), int(img_h * scale)
        
        # Cria cópia da imagem para edição
        self.editing_img = self.original_img.copy()
        self.display_img = self.editing_img.resize((disp_w, disp_h), Image.LANCZOS)

        # Variáveis para controle
        self.current_tk_img = ImageTk.PhotoImage(self.display_img)
        self.elements = []  # Lista de elementos desenhados
        self.undo_stack = []  # PILHA PARA DESFAZER AÇÕES - NOVO
        self.temp_element = None
        
        # Canvas para a imagem
        self.canvas = tk.Canvas(canvas_frame, width=disp_w, height=disp_h, cursor="cross", bg="gray")
        self.canvas.pack(padx=5, pady=5)
        self.canvas_img = self.canvas.create_image(0, 0, anchor="nw", image=self.current_tk_img)
        
        # Variáveis de controle - COR PADRÃO VERMELHA
        tool_var = tk.StringVar(value="rectangle")  # RETÂNGULO COMO PADRÃO
        color_var = tk.StringVar(value="#FF0000")   # VERMELHO COMO PADRÃO
        width_var = tk.IntVar(value=3)
        
        # Ferramentas - SUBSTITUINDO RADIOBUTTONS POR ÍCONES EMOJI
        self._create_styled_label(tools_frame, text="Ferramenta:").pack(side=tk.LEFT, padx=5)
        
        # Frame para os botões de ícone
        icon_frame = self._create_styled_frame(tools_frame)
        icon_frame.pack(side=tk.LEFT, padx=5)
        
        # 🔥 ADICIONADO: Ferramenta de mosaico
        tool_icons = {
            "rectangle": "⬜",   # Retângulo
            "circle": "🔴",      # Círculo  
            "arrow": "👉",       # Seta - Mão apontando
            "text": "🆎",        # Texto - Botão AB
            "blur": "🌀"         # Mosaico - NOVO
        }

        # Função para criar botões com estilo consistente
        def criar_botao_ferramenta(parent, texto, valor, variavel):
            btn = tk.Radiobutton(parent, text=texto, font=("Arial", 12), 
                               variable=variavel, value=valor, indicatoron=0, 
                               width=3, height=2, relief=tk.RAISED,
                               cursor="hand2")
            return btn

        # Cria os botões para cada ferramenta
        for tool_value, icon in tool_icons.items():
            btn = criar_botao_ferramenta(icon_frame, icon, tool_value, tool_var)
            btn.pack(side=tk.LEFT, padx=2)

        # Destacar o botão do retângulo (selecionado por padrão)
        for widget in icon_frame.winfo_children():
            if isinstance(widget, tk.Radiobutton) and widget.cget("value") == "rectangle":
                widget.config(relief=tk.SUNKEN, bg="#e3f2fd")  # Azul claro para selecionado
                break

        # Função para atualizar a aparência dos botões
        def update_button_appearance(*args):
            selected_tool = tool_var.get()
            for widget in icon_frame.winfo_children():
                if isinstance(widget, tk.Radiobutton):
                    if widget.cget("value") == selected_tool:
                        widget.config(relief=tk.SUNKEN, bg="#e3f2fd")  # Selecionado
                    else:
                        widget.config(relief=tk.RAISED, bg="SystemButtonFace")  # Normal

        tool_var.trace("w", update_button_appearance)
        
        # Controles de cor and espessura - APENAS CORES ESSENCIAIS
        color_frame = self._create_styled_frame(tools_frame)
        color_frame.pack(side=tk.LEFT, padx=20)
        
        self._create_styled_label(color_frame, text="Cor:").pack(side=tk.LEFT)
        
        # Paleta de cores reduzida (apenas as essenciais)
        colors = ["#FF0000", "#00FF00", "#FFFF00", "#000000", "#FFFFFF"]
        color_buttons_frame = self._create_styled_frame(color_frame)
        color_buttons_frame.pack(side=tk.LEFT, padx=5)
        
        for color in colors:
            btn = tk.Button(color_buttons_frame, bg=color, width=2, height=1, 
                           command=lambda c=color: self.set_color(color_var, c, color_preview))
            btn.pack(side=tk.LEFT, padx=1)
        
        # Botão para cor personalizada
        custom_btn = self._create_styled_button(color_frame, text="Personalizada", 
                              command=lambda: self.choose_custom_color(editor, color_var, color_preview))
        custom_btn.pack(side=tk.LEFT, padx=5)
        
        # Preview de cor
        color_preview = tk.Frame(color_frame, width=30, height=20, bg=color_var.get())
        color_preview.pack(side=tk.LEFT, padx=5)
        
        # Controle de espessura
        width_frame = self._create_styled_frame(tools_frame)
        width_frame.pack(side=tk.LEFT, padx=20)
        
        self._create_styled_label(width_frame, text="Espessura:").pack(side=tk.LEFT)
        tk.Scale(width_frame, from_=1, to=10, variable=width_var, orient=tk.HORIZONTAL, 
                length=100, showvalue=1).pack(side=tk.LEFT, padx=5)
        
        # BOTÃO DESFAZER - NOVO
        def undo_action():
            if self.elements:  # Se houver elementos para desfazer
                # Remove o último elemento e adiciona à pilha de desfazer
                removed_element = self.elements.pop()
                self.undo_stack.append(removed_element)
                refresh_display()
        
        undo_btn = self._create_styled_button(tools_frame, text="↩️ Desfazer (Ctrl+Z)", command=undo_action)
        undo_btn.pack(side=tk.LEFT, padx=20)
        
        # Variáveis para desenho
        start_xy = {"x": None, "y": None}
        
        def refresh_display():
            # Redesenha todos os elementos
            self.canvas.delete("all")
            self.canvas.create_image(0, 0, anchor="nw", image=self.current_tk_img)
            
            for element in self.elements:
                elem_type, coords, color, width, text = element
                scaled_coords = [int(c * self.scale_factor) for c in coords]
                
                if elem_type == "circle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_oval(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "rectangle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "arrow":
                    x1, y1, x2, y2 = scaled_coords
                    self.draw_arrow_on_canvas(x1, y1, x2, y2, color, width)
                elif elem_type == "text":
                    x, y = scaled_coords
                    self.canvas.create_text(x, y, text=text, fill=color, font=("Arial", 12), anchor="nw")
                # 🔥 ADICIONADO: Caso para mosaico
                elif elem_type == "blur":
                    x1, y1, x2, y2 = scaled_coords
                    # Desenhar um retângulo tracejado para representar a área de blur
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline="#FF00FF", width=2, dash=(5,5))
            
            # Desenha elemento temporário durante a criação
            if self.temp_element:
                elem_type, coords, color, width, text = self.temp_element
                scaled_coords = [int(c * self.scale_factor) for c in coords]
                
                if elem_type == "circle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_oval(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "rectangle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "arrow":
                    x1, y1, x2, y2 = scaled_coords
                    self.draw_arrow_on_canvas(x1, y1, x2, y2, color, width)
                # 🔥 ADICIONADO: Caso temporário para mosaico
                elif elem_type == "blur":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline="#FF00FF", width=2, dash=(5,5))
        
        def draw_arrow_on_canvas(x1, y1, x2, y2, color, width):
            # Desenha a linha da seta
            self.canvas.create_line(x1, y1, x2, y2, fill=color, width=width)
            
            # Calcula o ângulo da seta
            angle = math.atan2(y2 - y1, x2 - x1)
            
            # Desenha a ponta da seta (triângulo)
            arrow_size = 15
            x3 = x2 - arrow_size * math.cos(angle - math.pi/6)
            y3 = y2 - arrow_size * math.sin(angle - math.pi/6)
            x4 = x2 - arrow_size * math.cos(angle + math.pi/6)
            y4 = y2 - arrow_size * math.sin(angle + math.pi/6)
            
            self.canvas.create_polygon(x2, y2, x3, y3, x4, y4, fill=color, outline=color)
        
        def on_button_press(event):
            start_xy["x"], start_xy["y"] = event.x, event.y
        
        def on_motion(event):
            if start_xy["x"] is not None:
                # Desenho em tempo real
                sx, sy = start_xy["x"], start_xy["y"]
                ex, ey = event.x, event.y
                
                # Converte para coordenadas da imagem original
                ix1, iy1 = int(sx / self.scale_factor), int(sy / self.scale_factor)
                ix2, iy2 = int(ex / self.scale_factor), int(ey / self.scale_factor)
                
                tool = tool_var.get()
                color = color_var.get()
                width = width_var.get()
                
                if tool == "circle":
                    radius = int(((ix2 - ix1)**2 + (iy2 - iy1)**2)**0.5)
                    self.temp_element = ("circle", [ix1-radius, iy1-radius, ix1+radius, iy1+radius], color, width, "")
                elif tool == "rectangle":
                    # Garante que x2 >= x1 and y2 >= y1
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.temp_element = ("rectangle", [x1_norm, y1_norm, x2_norm, y2_norm], color, width, "")
                elif tool == "arrow":
                    self.temp_element = ("arrow", [ix1, iy1, ix2, iy2], color, width, "")
                # 🔥 ADICIONADO: Caso para mosaico
                elif tool == "blur":
                    # Para blur, também usamos coordenadas normalizadas
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.temp_element = ("blur", [x1_norm, y1_norm, x2_norm, y2_norm], "#FF00FF", 2, "")
                
                refresh_display()
        
        def on_button_release(event):
            if start_xy["x"] is not None:
                # Converte coordenadas da tela para coordenadas da imagem original
                sx, sy = start_xy["x"], start_xy["y"]
                ex, ey = event.x, event.y
                
                ix1, iy1 = int(sx / self.scale_factor), int(sy / self.scale_factor)
                ix2, iy2 = int(ex / self.scale_factor), int(ey / self.scale_factor)
                
                tool = tool_var.get()
                color = color_var.get()
                width = width_var.get()
                
                # Limpa a pilha de desfazer quando uma nova ação é realizada - NOVO
                self.undo_stack.clear()
                
                if tool == "circle":
                    radius = int(((ix2 - ix1)**2 + (iy2 - iy1)**2)**0.5)
                    self.elements.append(("circle", [ix1-radius, iy1-radius, ix1+radius, iy1+radius], color, width, ""))
                
                elif tool == "rectangle":
                    # Garante que x2 >= x1 and y2 >= y1
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.elements.append(("rectangle", [x1_norm, y1_norm, x2_norm, y2_norm], color, width, ""))
                
                elif tool == "arrow":
                    self.elements.append(("arrow", [ix1, iy1, ix2, iy2], color, width, ""))
                
                # 🔥 ADICIONADO: Caso para mosaico
                elif tool == "blur":
                    # Garante que x2 >= x1 and y2 >= y1
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.elements.append(("blur", [x1_norm, y1_norm, x2_norm, y2_norm], "", 0, ""))
                
                elif tool == "text":
                    # Para texto, pede o conteúdo e adiciona na posição clicada
                    text = simpledialog.askstring("Texto", "Digite o texto:", parent=editor)
                    if text:
                        self.elements.append(("text", [ix1, iy1], color, width, text))
                        # Atualiza a visualização para mostrar o texto imediatamente
                        refresh_display()
                
                self.temp_element = None
                refresh_display()
            
            start_xy["x"], start_xy["y"] = None, None
        
        # BIND DO CTRL+Z (atalho global dentro do editor)
        def on_key_press(event):
            if event.keysym == 'z' and (event.state & 0x4):  # Ctrl+Z
                undo_action()

        editor.bind('<Control-z>', on_key_press)
        editor.bind('<Control-Z>', on_key_press)
        
        # Bind events
        self.canvas.bind("<ButtonPress-1>", on_button_press)
        self.canvas.bind("<B1-Motion>", on_motion)
        self.canvas.bind("<ButtonRelease-1>", on_button_release)
        
        # Atualiza a visualização inicial
        refresh_display()
        
        # Frame para o botão Salvar (AGORA MAIS PRÓXIMO DA IMAGEM)
        button_frame = self._create_styled_frame(canvas_frame)
        button_frame.pack(pady=10)  # Reduzido o padding para ficar mais próximo
        
        def salvar_edicao():
            # Fecha a janela de seleção de cor personalizada se estiver aberta
            if hasattr(self, 'color_chooser_window') and self.color_chooser_window:
                try:
                    self.color_chooser_window.destroy()
                except:
                    pass
            
            # Aplica todos os elementos à imagem
            draw = ImageDraw.Draw(self.editing_img)
            
            for element in self.elements:
                elem_type, coords, color, width, text = element
                
                if elem_type == "circle":
                    x1, y1, x2, y2 = coords
                    draw.ellipse((x1, y1, x2, y2), outline=color, width=width)
                
                elif elem_type == "rectangle":
                    x1, y1, x2, y2 = coords
                    # Garante que as coordenadas estão normalizadas
                    x1_norm = min(x1, x2)
                    y1_norm = min(y1, y2)
                    x2_norm = max(x1, x2)
                    y2_norm = max(y1, y2)
                    draw.rectangle((x1_norm, y1_norm, x2_norm, y2_norm), outline=color, width=width)
                
                elif elem_type == "arrow":
                    x1, y1, x2, y2 = coords
                    draw.line((x1, y1, x2, y2), fill=color, width=width)
                    
                    # Calcula o ângulo da seta
                    angle = math.atan2(y2 - y1, x2 - x1)
                    
                    # Desenha a ponta da seta (triângulo)
                    arrow_size = 20
                    x3 = x2 - arrow_size * math.cos(angle - math.pi/6)
                    y3 = y2 - arrow_size * math.sin(angle - math.pi/6)
                    x4 = x2 - arrow_size * math.cos(angle + math.pi/6)
                    y4 = y2 - arrow_size * math.sin(angle + math.pi/6)
                    
                    draw.polygon([(x2, y2), (x3, y3), (x4, y4)], fill=color)
                
                elif elem_type == "text":
                    x, y = coords
                    try:
                        font = ImageFont.truetype("arial.ttf", 16)
                    except:
                        font = ImageFont.load_default()
                    
                    # Desenha texto diretamente
                    draw.text((x, y), text, fill=color, font=font)
                
                # 🔥 ADICIONADO: Aplicar mosaico
                elif elem_type == "blur":
                    x1, y1, x2, y2 = coords
                    # Aplicar efeito de blur na área selecionada
                    region = self.editing_img.crop((x1, y1, x2, y2))
                    # Aumentar o valor do radius para um blur mais forte
                    blurred_region = region.filter(ImageFilter.GaussianBlur(15))
                    self.editing_img.paste(blurred_region, (x1, y1, x2, y2))
            
            self.editing_img.convert("RGB").save(caminho_print, "PNG")
            messagebox.showinfo("Edição", "Evidência atualizada com sucesso!")
            editor.destroy()

        # Função para fechar o editor e garantir que la janela de cor seja fechada
        def fechar_editor():
            # Fecha a janela de seleção de cor personalizada se estiver aberta
            if hasattr(self, 'color_chooser_window') and self.color_chooser_window:
                try:
                    self.color_chooser_window.destroy()
                except:
                    pass
            editor.destroy()

        # Configurar o protocolo de fechamento da janela
        editor.protocol("WM_DELETE_WINDOW", fechar_editor)
        
        # Botão Salvar e Fechar (MESMO PADRÃO DOS OUTROS BOTÕES)
        self._create_styled_button(button_frame, text="💾 Salvar e Fechar", command=salvar_edicao, 
                                 style_type="accent").pack()

        editor.transient(parent)

    def set_color(self, color_var, color, preview_widget):
        color_var.set(color)
        preview_widget.config(bg=color)
    
    def choose_custom_color(self, parent, color_var, preview_widget):
        # Fecha qualquer janela de cor anterior que possa estar aberta
        if hasattr(self, 'color_chooser_window') and self.color_chooser_window:
            try:
                self.color_chooser_window.destroy()
            except:
                pass
        
        # Abre a nova janela de seleção de cor
        color = colorchooser.askcolor(title="Escolha uma cor", initialcolor=color_var.get(), parent=parent)
        if color[1]:
            color_var.set(color[1])
            preview_widget.config(bg=color[1])
    
    def draw_arrow_on_canvas(self, x1, y1, x2, y2, color, width):
        # Desenha a linha da seta
        self.canvas.create_line(x1, y1, x2, y2, fill=color, width=width)
        
        # Calcula o ângulo da seta
        angle = math.atan2(y2 - y1, x2 - x1)
        
        # Desenha the ponta da seta (triângulo)
        arrow_size = 15
        x3 = x2 - arrow_size * math.cos(angle - math.pi/6)
        y3 = y2 - arrow_size * math.sin(angle - math.pi/6)
        x4 = x2 - arrow_size * math.cos(angle + math.pi/6)
        y4 = y2 - arrow_size * math.sin(angle + math.pi/6)
        
        self.canvas.create_polygon(x2, y2, x3, y3, x4, y4, fill=color, outline=color)

    def gerar_documento(self):
        """Gera o documento DOCX com as evidências e retorna o caminho do documento"""
        doc_path = None
        try:
            print("🔄 Iniciando geração do documento DOCX...")
            
            # 🔥 CORREÇÃO: Criar novo documento em vez de reutilizar o existente
            if self.template_path and os.path.exists(self.template_path):
                self.doc = Document(self.template_path)
                self.using_template = True
                print(f"✅ Template carregado: {self.template_path}")
            else:
                self.doc = Document()
                self.using_template = False
                print("ℹ️ Criando documento vazio (sem template)")
            
            # Adicionar título se não estiver usando template
            if not self.using_template:
                titulo = self.doc.add_heading('Evidências Capturadas', 0)
                titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Adicionar data e hora
            if not self.using_template:
                data_hora = self.doc.add_paragraph()
                data_hora.add_run(f"Data e hora da geração: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}").italic = True
                data_hora.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Adicionar evidências
            for i, print_path in enumerate(self.prints, 1):
                print(f"📷 Adicionando evidência {i}: {print_path}")
                
                # Adicionar título da evidência
                self.doc.add_paragraph().add_run(f"Evidência {i}").bold = True
                
                # Adicionar comentário se existir
                nome_arquivo = os.path.basename(print_path)
                comentario = self.obter_comentario(nome_arquivo)
                if comentario:
                    comentario_para = self.doc.add_paragraph()
                    comentario_para.add_run(f"Comentário: {comentario}").italic = True
                
                # Adicionar imagem
                try:
                    paragraph = self.doc.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run()
                    
                    # 🔥 CORREÇÃO: Verificar se o arquivo existe antes de adicionar
                    if os.path.exists(print_path):
                        run.add_picture(print_path, width=Inches(6.0))
                        print(f"✅ Imagem {i} adicionada com sucesso")
                    else:
                        print(f"⚠️ Arquivo não encontrado: {print_path}")
                        self.doc.add_paragraph(f"[Arquivo de imagem não encontrado: {print_path}]")
                        
                except Exception as e:
                    print(f"❌ Erro ao adicionar imagem {print_path}: {e}")
                    self.doc.add_paragraph(f"[Erro ao carregar imagem: {print_path}]")
                
                # Adicionar separador
                self.doc.add_paragraph("―" * 50).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 🔥 CORREÇÃO: USAR NOME DO TEMPLATE PARA O DOCUMENTO
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Obter nome do template sem extensão
            template_filename = os.path.basename(self.template_path)
            template_name = os.path.splitext(template_filename)[0]
            
            # 🔥 CORREÇÃO: LIMPAR NOME DO TEMPLATE PARA EVITAR PROBLEMAS NO WINDOWS
            template_name = self._limpar_nome_arquivo(template_name)
            
            # Criar nome do documento usando o nome do template
            doc_filename = f"{template_name}_{timestamp}.docx"
            doc_path = os.path.join(self.output_dir, doc_filename)
            
            # 🔥 CORREÇÃO: Verificar se o diretório existe antes de salvar
            os.makedirs(os.path.dirname(doc_path), exist_ok=True)
            
            # 🔥 CORREÇÃO ADICIONAL: VERIFICAR SE O CAMINHO É VÁLIDO
            if len(doc_path) > 255:
                # Se o caminho for muito longo, criar um nome mais curto
                short_name = f"Evidencias_{timestamp}.docx"
                doc_path = os.path.join(self.output_dir, short_name)
                print(f"⚠️ Caminho muito longo, usando nome reduzido: {short_name}")
            
            self.doc.save(doc_path)
            print(f"✅ Documento salvo em: {doc_path}")
            
            # 🔥 EXCLUSÃO CONDICIONAL DAS EVIDÊNCIAS E PASTA AUTOMÁTICA
            if not self.manter_evidencias:
                print("🗑️ Excluindo arquivos de evidência conforme solicitado...")
                for print_path in self.prints:
                    try:
                        if os.path.exists(print_path):
                            os.remove(print_path)
                            print(f"🗑️ Excluído: {print_path}")
                    except Exception as e:
                        print(f"⚠️ Erro ao excluir {print_path}: {e}")
                
                # Também excluir o arquivo de metadados
                try:
                    if self.metadata_path and os.path.exists(self.metadata_path):
                        os.remove(self.metadata_path)
                        print(f"🗑️ Excluído: {self.metadata_path}")
                except Exception as e:
                    print(f"⚠️ Erro ao excluir metadados: {e}")
                
                # 🔥 SE FOI CRIADA UMA PASTA AUTOMÁTICA, EXCLUIR A PASTA INTEIRA
                if self.pasta_automatica and self.pasta_automatica_path:
                    try:
                        # Mover o DOCX para o diretório do template antes de excluir a pasta
                        template_dir = os.path.dirname(self.template_path)
                        novo_caminho_doc = os.path.join(template_dir, doc_filename)
                        
                        # 🔥 CORREÇÃO: VERIFICAR SE O ARQUIVO JÁ EXISTE
                        if os.path.exists(novo_caminho_doc):
                            # Adicionar um sufixo único se o arquivo já existir
                            base_name = os.path.splitext(doc_filename)[0]
                            extension = os.path.splitext(doc_filename)[1]
                            counter = 1
                            while os.path.exists(novo_caminho_doc):
                                novo_caminho_doc = os.path.join(template_dir, f"{base_name}_{counter}{extension}")
                                counter += 1
                        
                        # Mover o documento para o diretório do template
                        shutil.move(doc_path, novo_caminho_doc)
                        print(f"📄 Documento movido para: {novo_caminho_doc}")
                        
                        # Excluir a pasta automática (agora vazia)
                        shutil.rmtree(self.pasta_automatica_path)
                        print(f"🗑️ Pasta automática excluída: {self.pasta_automatica_path}")
                        
                        # Atualizar o caminho do documento
                        doc_path = novo_caminho_doc
                        
                    except Exception as e:
                        print(f"⚠️ Erro ao processar pasta automática: {e}")
            
            print(f"✅ Documento gerado com sucesso: {doc_path}")
            return doc_path
            
        except Exception as e:
            print(f"❌ Erro ao gerar documento: {e}")
            # 🔥 CORREÇÃO: Mostrar detalhes do erro
            import traceback
            traceback.print_exc()
            raise

    def close(self):
        """Fecha o módulo de forma segura"""
        try:
            # Parar gravação se estiver ativa
            if self.gravando:
                self.finalizar()
            
            # Parar listeners
            if self.listener_mouse:
                try:
                    self.listener_mouse.stop()
                except:
                    pass
                self.listener_mouse = None
                
            if hasattr(self, 'listener_keyboard') and self.listener_keyboard:
                try:
                    self.listener_keyboard.stop()
                except:
                    pass
                self.listener_keyboard = None
            
            # Fechar janelas
            if self.popup and self.popup.winfo_exists():
                try:
                    self.popup.destroy()
                except:
                    pass
                self.popup = None
                
            if self.root and self.root.winfo_exists():
                try:
                    self.root.destroy()
                except:
                    pass
                self.root = None
                
        except Exception as e:
            print(f"Erro ao fechar módulo de captura: {e}")

# Função de compatibilidade para manter a interface existente
def main():
    """Função principal para execução standalone"""
    root = tk.Tk()
    root.withdraw()  # Esconder a janela principal
    
    app = CaptureModule(parent=root)
    app.show()
    
    root.mainloop()

if __name__ == "__main__":
    main()