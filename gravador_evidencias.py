import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, colorchooser, ttk
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pyautogui
from pynput import mouse, keyboard
from PIL import Image, ImageTk, ImageDraw, ImageFont, ImageGrab 
from datetime import datetime
import math
import re
import screeninfo
import glob
import json
import uuid
import time
import ctypes
from ctypes import wintypes, byref
import tkinter.font as tkfont

# 🔥 CONTROLE AUTOMÁTICO DA BARRA DE TAREFAS
try:
    import win32gui
    import win32con
    import win32api
    WIN32_AVAILABLE = True
except ImportError:
    win32gui = None
    win32con = None
    win32api = None
    WIN32_AVAILABLE = False
    
# 🔥 ADICIONAR MSS PARA CAPTURA MULTI-MONITOR
try:
    import mss
    MSS_AVAILABLE = True
except ImportError:
    mss = None
    MSS_AVAILABLE = False

# ------------------ Gravador e Docx ------------------
class GravadorDocx:
    def __init__(self):
        self.gravando = False
        self.pausado = False
        self.output_dir = os.getcwd()
        self.listener_mouse = None
        self.prints = []
        self.doc = None
        self.evidencia_count = 0
        self.demanda = ""
        self.tipo_demanda = ""
        self.chamado = ""
        self.titulo = ""
        self.using_template = False
        self.template_path = None
        self.color_chooser_window = None
        self.current_index = 0
        self.evidence_dir = None
        self.metadata_path = None
        self.metadata = {"evidencias": [], "proximo_id": 1}
        self.popup = None
        self.current_img_label = None
        self.current_img_tk = None
        self.comment_entry = None
        self.manter_evidencias = None
        self.modo_captura = "manter"  # Valores: "manter", "ocultar"
        self.TIMESTAMP_TAMANHO_PADRAO = 24
        self.TIMESTAMP_POSICAO_PADRAO_X = 0.85  # Mais para a direita
        self.TIMESTAMP_POSICAO_PADRAO_Y = 0.92  # Mais para baixo

    def _salvar_metadata(self):
        """Salva os metadados no arquivo JSON"""
        if self.metadata_path:
            with open(self.metadata_path, 'w', encoding='utf-8') as f:
                json.dump(self.metadata, f, indent=2, ensure_ascii=False)

    def carregar_evidencias(self, dir_path):
        """Carrega as evidências baseadas nos metadados"""
        self.metadata_path = os.path.join(dir_path, "evidencias_metadata.json")
        
        if os.path.exists(self.metadata_path):
            try:
                with open(self.metadata_path, 'r', encoding='utf-8') as f:
                    self.metadata = json.load(f)
            except:
                self.metadata = {"evidencias": [], "proximo_id": 1}
        
        # Carrega evidências ativas (não excluídas)
        evidencias_ativas = []
        for evidencia in self.metadata["evidencias"]:
            if not evidencia.get("excluida", False):
                caminho = os.path.join(dir_path, evidencia["arquivo"])
                if os.path.exists(caminho):
                    evidencias_ativas.append(caminho)
        
        return evidencias_ativas

    def recarregar_evidencias(self):
        """Recarrega a lista de evidências"""
        if self.evidence_dir:
            self.prints = self.carregar_evidencias(self.evidence_dir)
            return True
        return False

    # 🔥 MÉTODOS DE CAPTURA SIMPLIFICADOS E OTIMIZADOS
    def capture_inteligente(self, x, y):
        """
        Captura a tela baseado no modo selecionado pelo usuário
        """
        if self.modo_captura == "manter":
            # Modo manter: captura tela COMPLETA (incluindo barra de tarefas)
            return self.capture_tela_completa_mss(x, y)
        else:
            # Modo ocultar: captura apenas área de trabalho (sem barra)
            return self.capture_work_area_pyautogui(x, y)

    def capture_tela_completa_mss(self, x, y):
        """
        Captura a tela completa INCLUINDO a barra de tarefas.
        Funciona no primário e secundário, mesmo com coordenadas negativas.
        """
        try:
            # 🔥 ESTRATÉGIA 1: Win32 API para captura precisa de monitor específico
            if WIN32_AVAILABLE:
                try:
                    # Encontrar o monitor que contém o ponto (x, y)
                    monitor_handle = win32api.MonitorFromPoint((x, y), win32con.MONITOR_DEFAULTTONEAREST)
                    monitor_info = win32gui.GetMonitorInfo(monitor_handle)
                    
                    # Área completa do monitor (inclui barra)
                    monitor_area = monitor_info["Monitor"]  # (left, top, right, bottom)
                    
                    # Capturar usando MSS para melhor compatibilidade com múltiplos monitores
                    if MSS_AVAILABLE:
                        with mss.mss() as sct:
                            monitor_mss = {
                                "left": monitor_area[0],
                                "top": monitor_area[1], 
                                "width": monitor_area[2] - monitor_area[0],
                                "height": monitor_area[3] - monitor_area[1]
                            }
                            screenshot = sct.grab(monitor_mss)
                            img = Image.frombytes("RGB", screenshot.size, screenshot.bgra, "raw", "BGRX")
                            
                            rel_x = x - monitor_area[0]
                            rel_y = y - monitor_area[1]
                            
                            metodo_utilizado = f"Win32 + MSS Monitor Completo {monitor_area}"
                            print(f"✅ CAPTURA WIN32+MSS - Monitor {monitor_area} | Coord: ({rel_x},{rel_y})")
                            
                            return img, (rel_x, rel_y), metodo_utilizado
                    else:
                        # Fallback para ImageGrab se MSS não disponível
                        screenshot = ImageGrab.grab(bbox=monitor_area)
                        rel_x = x - monitor_area[0]
                        rel_y = y - monitor_area[1]
                        
                        metodo_utilizado = f"Win32 Monitor Completo {monitor_area}"
                        print(f"✅ CAPTURA WIN32 - Monitor {monitor_area} | Coord: ({rel_x},{rel_y})")
                        
                        return screenshot, (rel_x, rel_y), metodo_utilizado
                        
                except Exception as e:
                    print(f"⚠️  Win32 falhou (capturando com alternativa): {e}")

            # 🔥 ESTRATÉGIA 2: MSS como alternativa principal
            if MSS_AVAILABLE:
                try:
                    with mss.mss() as sct:
                        monitor_encontrado = None
                        
                        # Procurar em todos os monitores (exceto o virtual)
                        for monitor in sct.monitors[1:]:
                            if (monitor["left"] <= x < monitor["left"] + monitor["width"] and
                                monitor["top"] <= y < monitor["top"] + monitor["height"]):
                                monitor_encontrado = monitor
                                break

                        # Fallback para primeiro monitor se não encontrou
                        if not monitor_encontrado:
                            monitor_encontrado = sct.monitors[1] if len(sct.monitors) > 1 else sct.monitors[0]
                            print(f"⚠️  Monitor não encontrado para coordenadas ({x},{y}), usando monitor {monitor_encontrado} como fallback")

                        # Capturar a tela completa do monitor encontrado
                        screenshot = sct.grab(monitor_encontrado)
                        img = Image.frombytes("RGB", screenshot.size, screenshot.bgra, "raw", "BGRX")

                        # Calcular coordenadas relativas ao monitor
                        rel_x = x - monitor_encontrado["left"]
                        rel_y = y - monitor_encontrado["top"]

                        metodo_utilizado = f"MSS Monitor Completo {monitor_encontrado['width']}x{monitor_encontrado['height']}"
                        print(f"✅ CAPTURA MSS - Monitor {monitor_encontrado} | Coord: ({rel_x},{rel_y})")

                        return img, (rel_x, rel_y), metodo_utilizado
                        
                except Exception as e:
                    print(f"⚠️  MSS falhou (capturando com alternativa): {e}")

            # 🔥 ESTRATÉGIA 3: Fallback com ImageGrab
            try:
                screenshot = ImageGrab.grab()
                metodo_utilizado = "Fallback - ImageGrab (tela completa)"
                print(f"⚠️  Usando fallback ImageGrab para coordenadas ({x},{y})")
                return screenshot, (x, y), metodo_utilizado
                
            except Exception as e:
                print(f"⚠️  ImageGrab falhou: {e}")

            # 🔥 ESTRATÉGIA 4: Fallback final com pyautogui
            try:
                screenshot = pyautogui.screenshot()
                metodo_utilizado = "Fallback - pyautogui (apenas primário)"
                print(f"⚠️  Usando fallback pyautogui para coordenadas ({x},{y})")
                return screenshot, (x, y), metodo_utilizado
                
            except Exception as e:
                print(f"❌ Todos os métodos de captura falharam: {e}")
                raise

        except Exception as e:
            print(f"❌ Falha crítica na captura completa: {e}")
            # Último recurso - retorna imagem preta ou levanta exceção
            try:
                # Tenta criar uma imagem preta como fallback extremo
                img = Image.new('RGB', (100, 100), color='black')
                return img, (0, 0), f"Fallback Extremo - Erro: {str(e)}"
            except:
                raise Exception(f"Falha completa na captura de tela: {str(e)}")

    def capture_work_area_pyautogui(self, x, y):
        """
        Captura apenas a área de trabalho (SEM barra de tarefas) em QUALQUER monitor
        """
        try:
            # 🔥 ESTRATÉGIA 1: Win32 API + MSS para multi-monitor
            if WIN32_AVAILABLE and MSS_AVAILABLE:
                try:
                    # Encontrar o monitor que contém o ponto (x, y)
                    monitor_handle = win32api.MonitorFromPoint((x, y), win32con.MONITOR_DEFAULTTONEAREST)
                    monitor_info = win32gui.GetMonitorInfo(monitor_handle)
                    
                    # 🔥 USAR WORK AREA (área sem barra de tarefas)
                    work_area = monitor_info["Work"]  # (left, top, right, bottom)
                    
                    with mss.mss() as sct:
                        # Configurar captura específica para o monitor encontrado
                        monitor_mss = {
                            "left": work_area[0],
                            "top": work_area[1], 
                            "width": work_area[2] - work_area[0],
                            "height": work_area[3] - work_area[1]
                        }
                        
                        screenshot = sct.grab(monitor_mss)
                        img = Image.frombytes("RGB", screenshot.size, screenshot.bgra, "raw", "BGRX")
                        
                        # Calcular coordenadas relativas à work area do monitor
                        rel_x = x - work_area[0]
                        rel_y = y - work_area[1]
                        
                        metodo_utilizado = f"Win32+MSS Work Area Monitor {work_area}"
                        print(f"✅ CAPTURA SEM BARRA - Monitor Work Area {work_area} | Coord: ({rel_x},{rel_y})")
                        
                        return img, (rel_x, rel_y), metodo_utilizado
                        
                except Exception as e:
                    print(f"⚠️  Win32+MSS falhou (capturando com alternativa): {e}")

            # 🔥 ESTRATÉGIA 2: MSS puro para multi-monitor (sem Win32)
            if MSS_AVAILABLE:
                try:
                    with mss.mss() as sct:
                        monitor_encontrado = None
                        
                        # Procurar em todos os monitores (exceto o virtual)
                        for i, monitor in enumerate(sct.monitors[1:], 1):
                            if (monitor["left"] <= x < monitor["left"] + monitor["width"] and
                                monitor["top"] <= y < monitor["top"] + monitor["height"]):
                                monitor_encontrado = monitor
                                break

                        # Fallback para primeiro monitor se não encontrou
                        if not monitor_encontrado:
                            monitor_encontrado = sct.monitors[1] if len(sct.monitors) > 1 else sct.monitors[0]
                            print(f"⚠️  Monitor não encontrado para coordenadas ({x},{y}), usando fallback")

                        # 🔥 ESTIMAR WORK AREA (recortar barra de tarefas)
                        barra_altura = self.estimativa_segura_barra_tarefas(monitor_encontrado["height"])
                        work_area = {
                            "left": monitor_encontrado["left"],
                            "top": monitor_encontrado["top"], 
                            "width": monitor_encontrado["width"],
                            "height": monitor_encontrado["height"] - barra_altura
                        }

                        # Capturar a work area do monitor
                        screenshot = sct.grab(work_area)
                        img = Image.frombytes("RGB", screenshot.size, screenshot.bgra, "raw", "BGRX")

                        # Calcular coordenadas relativas ao monitor
                        rel_x = x - work_area["left"]
                        rel_y = y - work_area["top"]

                        metodo_utilizado = f"MSS Work Area Monitor {work_area['width']}x{work_area['height']}"
                        print(f"✅ CAPTURA SEM BARRA - MSS Work Area | Coord: ({rel_x},{rel_y})")

                        return img, (rel_x, rel_y), metodo_utilizado
                        
                except Exception as e:
                    print(f"⚠️  MSS falhou (capturando com alternativa): {e}")

            # 🔥 ESTRATÉGIA 3: ScreenInfo + ImageGrab para multi-monitor
            try:
                monitors = screeninfo.get_monitors()
                target_monitor = None
                
                # Encontrar o monitor que contém as coordenadas (x, y)
                for monitor in monitors:
                    if (monitor.x <= x < monitor.x + monitor.width and
                        monitor.y <= y < monitor.y + monitor.height):
                        target_monitor = monitor
                        break
                
                if target_monitor:
                    # 🔥 ESTIMAR WORK AREA (recortar barra)
                    barra_altura = self.estimativa_segura_barra_tarefas(target_monitor.height)
                    work_area = (
                        target_monitor.x,
                        target_monitor.y,
                        target_monitor.x + target_monitor.width,
                        target_monitor.y + target_monitor.height - barra_altura
                    )
                    
                    screenshot = ImageGrab.grab(bbox=work_area)
                    rel_x = x - work_area[0]
                    rel_y = y - work_area[1]
                    
                    metodo_utilizado = f"ScreenInfo Work Area {work_area}"
                    print(f"✅ CAPTURA SEM BARRA - ScreenInfo Work Area | Coord: ({rel_x},{rel_y})")
                    
                    return screenshot, (rel_x, rel_y), metodo_utilizado
                    
            except Exception as e:
                print(f"⚠️  ScreenInfo falhou: {e}")

            # 🔥 ESTRATÉGIA 4: Fallback - recorte manual do monitor primário
            try:
                # Capturar tela completa primeiro
                screenshot_full = pyautogui.screenshot()
                screen_width, screen_height = screenshot_full.size
                
                # Verificar se as coordenadas estão no monitor primário
                if 0 <= x < screen_width and 0 <= y < screen_height:
                    # Recortar barra do primário
                    barra_altura = self.estimativa_segura_barra_tarefas(screen_height)
                    work_area = (0, 0, screen_width, screen_height - barra_altura)
                    screenshot = screenshot_full.crop(work_area)
                    
                    rel_x = x
                    rel_y = y
                    if y > screen_height - barra_altura:
                        rel_y = screen_height - barra_altura - 5
                        
                    metodo_utilizado = f"Fallback Primário Recortado"
                    print(f"⚠️  CAPTURA SEM BARRA (Primário) | Coord: ({rel_x},{rel_y})")
                    
                    return screenshot, (rel_x, rel_y), metodo_utilizado
                else:
                    # Coordenadas fora do primário - retornar tela completa como fallback
                    metodo_utilizado = "Fallback - Tela Completa (fora do primário)"
                    print(f"❌ Coordenadas ({x},{y}) fora do monitor primário, usando tela completa")
                    return screenshot_full, (x, y), metodo_utilizado
                    
            except Exception as e:
                print(f"❌ Fallback falhou: {e}")

            # 🔥 ESTRATÉGIA 5: Último recurso
            screenshot = pyautogui.screenshot()
            metodo_utilizado = "Fallback Extremo - Tela Completa"
            print(f"❌ TODOS OS MÉTODOS FALHARAM - Retornando tela completa")
            
            return screenshot, (x, y), metodo_utilizado

        except Exception as e:
            print(f"❌ Falha crítica na captura sem barra: {e}")
            # Fallback extremo
            screenshot = pyautogui.screenshot()
            return screenshot, (x, y), f"Erro Crítico: {str(e)}"
        
    def estimativa_segura_barra_tarefas(self, altura_tela):
        """
        Estimativa conservadora da altura da barra de tarefas
        """
        if altura_tela >= 2160:  # 4K
            return 100
        elif altura_tela >= 1440:  # QHD
            return 80
        elif altura_tela >= 1080:  # Full HD
            return 70
        else:  # Resoluções menores
            return 60
        
    # 🔥 NOVA FUNÇÃO: APLICAR TIMESTAMP MODERNO COM FUNDO
    def aplicar_timestamp_moderno(self, caminho_imagem, evidencia_meta):
        """Aplica o timestamp com fundo semi-transparente e texto centralizado"""
        img = Image.open(caminho_imagem).convert("RGBA")
        draw = ImageDraw.Draw(img)
        
        # Calcular posição em pixels
        img_width, img_height = img.size
        pos_x_percent = evidencia_meta["timestamp_posicao"]["x"]
        pos_y_percent = evidencia_meta["timestamp_posicao"]["y"]
        
        # Configurações do texto
        texto = evidencia_meta["timestamp_texto"]
        texto_cor = evidencia_meta["timestamp_cor"]  # Branco
        fundo_cor = evidencia_meta.get("timestamp_fundo", "#000000B2")  # Preto 70%
        tamanho = evidencia_meta.get("timestamp_tamanho", self.TIMESTAMP_TAMANHO_PADRAO)
        
        # Converter cor de fundo para RGBA
        if fundo_cor.startswith("#") and len(fundo_cor) == 9:  # Formato #RRGGBBAA
            r = int(fundo_cor[1:3], 16)
            g = int(fundo_cor[3:5], 16)
            b = int(fundo_cor[5:7], 16)
            a = int(fundo_cor[7:9], 16)
            fundo_rgba = (r, g, b, a)
        else:
            fundo_rgba = (0, 0, 0, 178)  # Fallback: preto 70%
        
        # Usar fonte
        try:
            font = ImageFont.truetype("arial.ttf", tamanho)
        except:
            font = ImageFont.load_default()
        
        # 🔥 CALCULAR TAMANHO DO TEXTO PARA CENTRALIZAR
        bbox = draw.textbbox((0, 0), texto, font=font)
        texto_largura = bbox[2] - bbox[0]
        texto_altura = bbox[3] - bbox[1]
        
        # 🔥 DEFINIR PADDING E CANTOS ARREDONDADOS
        padding_horizontal = 20
        padding_vertical = 12
        borda_radius = 8
        
        # 🔥 CALCULAR POSIÇÃO FINAL DO FUNDO (centralizado na posição especificada)
        fundo_largura = texto_largura + (padding_horizontal * 2)
        fundo_altura = texto_altura + (padding_vertical * 2)
        
        # Calcular coordenadas do fundo baseado na posição percentual
        fundo_x1 = int((img_width * pos_x_percent) - (fundo_largura / 2))  # Centralizado horizontalmente
        fundo_y1 = int((img_height * pos_y_percent) - (fundo_altura / 2))   # Centralizado verticalmente
        fundo_x2 = fundo_x1 + fundo_largura
        fundo_y2 = fundo_y1 + fundo_altura
        
        # 🔥 GARANTIR que o fundo não saia dos limites da imagem
        margem = 10
        if fundo_x1 < margem:
            fundo_x1 = margem
            fundo_x2 = fundo_x1 + fundo_largura
        elif fundo_x2 > img_width - margem:
            fundo_x2 = img_width - margem
            fundo_x1 = fundo_x2 - fundo_largura
            
        if fundo_y1 < margem:
            fundo_y1 = margem
            fundo_y2 = fundo_y1 + fundo_altura
        elif fundo_y2 > img_height - margem:
            fundo_y2 = img_height - margem
            fundo_y1 = fundo_y2 - fundo_altura
        
        # 🔥 CORREÇÃO CRÍTICA: Calcular a posição vertical do texto considerando a métrica da fonte
        bbox = draw.textbbox((0, 0), texto, font=font)
        texto_ascendente = -bbox[1]  # A parte "acima" da linha de base
        texto_largura = bbox[2] - bbox[0]
        texto_altura_total = bbox[3] - bbox[1]

        # 🔥 CALCULAR POSIÇÃO DO TEXTO (verdadeiramente centralizado no fundo)
        texto_x = fundo_x1 + padding_horizontal
        # Ajuste fino para centralização vertical perfeita
        texto_y = fundo_y1 + (fundo_altura - texto_altura_total) // 2 + texto_ascendente
        
        # 🔥 DESENHAR FUNDO COM CANTOS ARREDONDADOS
        # Criar máscara para cantos arredondados
        mask = Image.new("L", (fundo_largura, fundo_altura), 0)
        mask_draw = ImageDraw.Draw(mask)
        mask_draw.rounded_rectangle(
            [0, 0, fundo_largura, fundo_altura],
            radius=borda_radius,
            fill=255
        )
        
        # Aplicar fundo semi-transparente
        fundo_img = Image.new("RGBA", (fundo_largura, fundo_altura), fundo_rgba)
        img.paste(fundo_img, (fundo_x1, fundo_y1), mask)
        
        # 🔥 DESENHAR TEXTO BRANCO VERDADEIRAMENTE CENTRALIZADO
        draw.text((texto_x, texto_y), texto, fill=texto_cor, font=font)
        
        # Salvar a imagem
        img.save(caminho_imagem)

    # ---------- Nova janela de configuração ----------
    def mostrar_janela_configuracao(self):
        config_window = tk.Toplevel(root)
        config_window.title("Configuração de Gravação")
        config_window.geometry("600x600")
        config_window.resizable(False, False)
        
        config_window.transient(root)
        config_window.grab_set()
        
        main_frame = ttk.Frame(config_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="PrintF - Configuração de Gravação", 
                 font=("Arial", 16, "bold")).pack(pady=10)
        
        # Seleção de template
        ttk.Label(main_frame, text="Selecione o template DOCX:").pack(anchor="w", pady=(10, 5))
        
        template_frame = ttk.Frame(main_frame)
        template_frame.pack(fill=tk.X, pady=5)
        
        self.template_var = tk.StringVar()
        template_entry = ttk.Entry(template_frame, textvariable=self.template_var, width=40)
        template_entry.pack(side=tk.LEFT, padx=(0, 5), fill=tk.X, expand=True)
        
        def selecionar_template():
            template_path = filedialog.askopenfilename(
                title="Selecione o template DOCX",
                filetypes=[("Documentos Word", "*.docx")]
            )
            if template_path:
                self.template_var.set(template_path)
        
        ttk.Button(template_frame, text="Procurar", command=selecionar_template).pack(side=tk.RIGHT)
        
        # Seleção de diretório de destino
        ttk.Label(main_frame, text="Selecione o diretório de destino:").pack(anchor="w", pady=(20, 5))
        
        dir_frame = ttk.Frame(main_frame)
        dir_frame.pack(fill=tk.X, pady=5)
        
        self.dir_var = tk.StringVar()
        dir_entry = ttk.Entry(dir_frame, textvariable=self.dir_var, width=40)
        dir_entry.pack(side=tk.LEFT, padx=(0, 5), fill=tk.X, expand=True)
        
        def selecionar_diretorio():
            dir_path = filedialog.askdirectory(title="Selecione o diretório para salvar")
            if dir_path:
                self.dir_var.set(dir_path)
        
        ttk.Button(dir_frame, text="Procurar", command=selecionar_diretorio).pack(side=tk.RIGHT)
        
        # 🔥 NOVO: Seleção do modo de captura (APENAS 2 OPÇÕES)
        ttk.Label(main_frame, text="Modo de Captura da Barra de Tarefas:", 
                 font=("Arial", 11, "bold")).pack(anchor="w", pady=(20, 10))
        
        # Variável para os RadioButtons
        self.modo_captura_var = tk.StringVar(value="manter")  # Valor padrão
        
        # Frame para os RadioButtons
        modo_frame = ttk.Frame(main_frame)
        modo_frame.pack(fill=tk.X, pady=5)
        
        # RadioButton 1: Manter barra completa
        rb1 = ttk.Radiobutton(
            modo_frame, 
            text="Manter barra de tarefas (data/hora visível na barra do Windows)",
            variable=self.modo_captura_var, 
            value="manter"
        )
        rb1.pack(anchor="w", pady=2)
        
        # RadioButton 2: Ocultar barra
        rb2 = ttk.Radiobutton(
            modo_frame, 
            text="Ocultar barra de tarefas (data/hora será adicionada na imagem)",
            variable=self.modo_captura_var, 
            value="ocultar"
        )
        rb2.pack(anchor="w", pady=2)
        
        # Checkbox para manter evidências
        ttk.Label(main_frame, text="Opções de saída:", font=("Arial", 11, "bold")).pack(anchor="w", pady=(20, 10))
        
        # Variável para o checkbox - valor padrão True (marcado)
        self.manter_evidencias_var = tk.BooleanVar(value=True)
        
        # Checkbox
        checkbox_frame = ttk.Frame(main_frame)
        checkbox_frame.pack(fill=tk.X, pady=5)
        
        manter_checkbox = ttk.Checkbutton(
            checkbox_frame, 
            text="Manter arquivos de evidência (prints) na pasta após gerar o DOCX",
            variable=self.manter_evidencias_var
        )
        manter_checkbox.pack(anchor="w")
        
        # Label informativa
        info_label = ttk.Label(
            main_frame, 
            text="Se desmarcado, os arquivos de print serão excluídos após a geração do DOCX.", 
            font=("Arial", 9), 
            foreground="gray",
            justify=tk.LEFT
        )
        info_label.pack(anchor="w", pady=(5, 15))
        
        # Frame para os botões na parte inferior
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(20, 0))
        
        def iniciar_com_config():
            if not self.template_var.get() or not self.dir_var.get():
                messagebox.showerror("Erro", "Por favor, selecione o template e o diretório de destino.")
                return
            
            if not os.path.exists(self.template_var.get()):
                messagebox.showerror("Erro", "O arquivo de template selecionado não existe.")
                return
            
            # 🔥 Armazena a escolha do modo de captura
            self.modo_captura = self.modo_captura_var.get()
            
            # 🔥 VERIFICAÇÃO ADICIONAL: Limpar qualquer estado residual
            self.gravando = False
            self.pausado = False
            self.prints = []
            
            # VALIDAÇÃO SIMPLES: BLOQUEAR APENAS SE TIVER ARQUIVOS
            dir_path = self.dir_var.get()
            
            if os.path.exists(dir_path):
                try:
                    # Verificar se existe algum arquivo na pasta raiz
                    for item in os.listdir(dir_path):
                        item_path = os.path.join(dir_path, item)
                        # Ignorar ocultos e verificar apenas arquivos (não pastas)
                        if not item.startswith('.') and os.path.isfile(item_path):
                            messagebox.showerror(
                                "Arquivos na Pasta", 
                                f"A pasta selecionada contém arquivos.\n\n"
                                f"Para evitar misturar evidências, a pasta deve estar vazia "
                                f"ou conter apenas outras pastas.\n\n"                                
                            )
                            return
                            
                except PermissionError:
                    messagebox.showerror("Erro de Permissão", "Sem permissão para acessar a pasta selecionada.")
                    return
                except Exception as e:
                    messagebox.showerror("Erro", f"Erro ao verificar a pasta: {str(e)}")
                    return
            
            # Armazenar a escolha do usuário
            self.manter_evidencias = self.manter_evidencias_var.get()
            
            self.template_path = self.template_var.get()
            self.output_dir = self.dir_var.get()
            self.evidence_dir = self.dir_var.get()
            config_window.destroy()            
              
            self.iniciar_gravacao()
        
        # Centralizar os botões horizontalmente
        button_container = ttk.Frame(btn_frame)
        button_container.pack(expand=True)
        
        ttk.Button(button_container, text="Iniciar Gravação", command=iniciar_com_config).pack(side=tk.LEFT, padx=10)
        ttk.Button(button_container, text="Cancelar", command=config_window.destroy).pack(side=tk.LEFT, padx=10)
        
        # Forçar atualização da interface e ajustar tamanho se necessário
        config_window.update_idletasks()
        
        # Se a janela for muito grande para a tela, ajustar
        screen_width = config_window.winfo_screenwidth()
        screen_height = config_window.winfo_screenheight()
        
        if config_window.winfo_height() > screen_height:
            config_window.geometry(f"600x{screen_height-100}")
        
        root.wait_window(config_window)
        return self.template_path is not None and self.output_dir is not None
        
    # ---------- Captura de telas ----------
    def iniciar_gravacao(self):
        # ------ RESET CRÍTICO: Limpar estado da gravação anterior ------
        self.gravando = False
        self.pausado = False
        self.prints = []
        self.evidencia_count = 0
        self.current_index = 0
        self.metadata = {"evidencias": [], "proximo_id": 1}
        
        os.makedirs(self.output_dir, exist_ok=True)

        # Inicializar metadados
        self.metadata_path = os.path.join(self.output_dir, "evidencias_metadata.json")
        self.metadata = {"evidencias": [], "proximo_id": 1}

        # 🔥 IMPORTANTE: Criar NOVO documento Word
        try:
            if os.path.exists(self.template_path):
                self.doc = Document(self.template_path)
                self.using_template = True
                print("Template carregado com sucesso!")
            else:
                self.doc = Document()
                self.using_template = False
                print("Template não encontrado. Criando documento vazio.")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar template: {str(e)}")
            self.doc = Document()
            self.using_template = False
            
        messagebox.showinfo("Gravação", "▶ Clique em OK para começar a gravar!")
        
        print(f"Iniciando gravação com modo: {self.modo_captura}")

        self.gravando = True
        self.pausado = False

        minimizar_janela()
        
        # 🔥 Parar listener anterior se existir
        if self.listener_mouse:
            self.listener_mouse.stop()
        
        # 🔥 CORREÇÃO: Iniciar o listener do mouse
        self.listener_mouse = mouse.Listener(on_click=self.on_click)
        self.listener_mouse.start()
        
        print("✅ Gravador iniciado - capturando cliques")

    def pausar(self):
        if self.gravando and not self.pausado:
            self.pausado = True
            
            # 🔥 CORREÇÃO CRÍTICA: Parar o listener do mouse quando pausado
            if self.listener_mouse:
                self.listener_mouse.stop()
                print("⏸️ Listener do mouse PAUSADO")
            
            messagebox.showinfo("Gravação", "⏸ Gravação pausada! Cliques NÃO serão capturados.")

    def retomar(self):
        if self.gravando and self.pausado:            
            # 🔥 CORREÇÃO: Mostrar mensagem primeiro e só retomar após o OK
            messagebox.showinfo("Gravação", "▶ Gravação retomada! Cliques serão capturados novamente.")
            
            # Só depois do OK mudar o estado e reiniciar o listener
            self.pausado = False
            
            # 🔥 CORREÇÃO CRÍTICA: Reiniciar o listener do mouse quando retomado
            if self.listener_mouse:
                # Parar listener anterior se ainda estiver ativo
                try:
                    self.listener_mouse.stop()
                except:
                    pass
                
                # Criar novo listener
                self.listener_mouse = mouse.Listener(on_click=self.on_click)
                self.listener_mouse.start()
                print("▶️ Listener do mouse RETOMADO")

    def finalizar(self):
        if self.gravando:
            self.gravando = False
            self.pausado = False  # 🔥 Garantir que não fique em estado pausado
            
            if self.listener_mouse:
                self.listener_mouse.stop()
                self.listener_mouse = None
                print("⏹️ Listener do mouse FINALIZADO")
            
            print("Gravação finalizada - usando captura híbrida sem alterações na barra")
            
            messagebox.showinfo("Gravação", "⏹ Gravação finalizada!")
            if self.prints:
                self.gerar_docx()
            else:
                messagebox.showinfo("Info", "Nenhuma evidência capturada.")

    # 🔥 MÉTODO on_click ATUALIZADO PARA USAR CAPTURA INTELIGENTE
    def on_click(self, x, y, button, pressed):
        # 🔥 CORREÇÃO: Verificação mais robusta do estado
        if not self.gravando:
            return
            
        if self.pausado:
            # 🔥 CORREÇÃO: Se estiver pausado, não fazer nada
            return
            
        if pressed and self.gravando and not self.pausado:
            print(f"🎯 Clique capturado em ({x}, {y})")
            
            # 🔥 USAR CAPTURA INTELIGENTE (multi-monitor)
            screenshot, coordenadas_relativas, metodo = self.capture_inteligente(x, y)
            
            if not screenshot:
                print("Erro: Não foi possível capturar a tela")
                return

            # Gerar nome único com ID sequencial e timestamp
            evidencia_id = self.metadata["proximo_id"]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            nome_arquivo = f"evidencia_{evidencia_id:04d}_{timestamp}.png"
            caminho_print = os.path.join(self.output_dir, nome_arquivo)

            try:
                click_x, click_y = coordenadas_relativas
                
                img = screenshot.convert("RGBA")
                overlay = Image.new("RGBA", img.size, (255, 255, 255, 0))
                draw = ImageDraw.Draw(overlay)
                r = 40
                
                draw.ellipse((click_x-r, click_y-r, click_x+r, click_y+r), fill=(255, 255, 0, 100))
                final_img = Image.alpha_composite(img, overlay)
                final_img.convert("RGB").save(caminho_print, "PNG")
                
                # Adicionar aos metadados
                timestamp_captura = datetime.now().isoformat()
                timestamp_texto = datetime.now().strftime('%d/%m/%Y %H:%M:%S')

                # 🔥 DIFERENÇA CRÍTICA: Só adiciona dados de timestamp se for modo "ocultar"
                metadados_timestamp = {}
                if self.modo_captura == "ocultar":
                    metadados_timestamp = {
                        "timestamp_texto": timestamp_texto,
                        # 🔥 ALTERADO: Usar posição padrão centralizada no canto inferior direito
                        "timestamp_posicao": {"x": self.TIMESTAMP_POSICAO_PADRAO_X, "y": self.TIMESTAMP_POSICAO_PADRAO_Y},
                        "timestamp_cor": "#FFFFFF",
                        "timestamp_fundo": "#000000B2",
                        "timestamp_tamanho": self.TIMESTAMP_TAMANHO_PADRAO,
                    }
                else:
                    metadados_timestamp = {
                        "timestamp_texto": "",
                        # 🔥 ALTERADO: Usar posição padrão centralizada no canto inferior direito
                        "timestamp_posicao": {"x": self.TIMESTAMP_POSICAO_PADRAO_X, "y": self.TIMESTAMP_POSICAO_PADRAO_Y},
                        "timestamp_cor": "#FFFFFF", 
                        "timestamp_fundo": "#000000B2",
                        "timestamp_tamanho": self.TIMESTAMP_TAMANHO_PADRAO,
                    }

                self.metadata["evidencias"].append({
                    "id": evidencia_id,
                    "arquivo": nome_arquivo,
                    "timestamp": timestamp_captura,
                    "excluida": False,
                    "comentario": "",
                    "metodo_captura": metodo,
                    **metadados_timestamp
                })
                self.metadata["proximo_id"] += 1
                self._salvar_metadata()
                
                self.prints.append(caminho_print)
                print(f"✅ Print salvo: {caminho_print} | Método: {metodo} | Modo: {self.modo_captura}")
                
            except Exception as e:
                print(f"❌ Erro ao processar captura: {e}")
                try:
                    # Fallback: salvar screenshot diretamente
                    screenshot.save(caminho_print)
                    
                    self.metadata["evidencias"].append({
                        "id": evidencia_id,
                        "arquivo": nome_arquivo,
                        "timestamp": datetime.now().isoformat(),
                        "excluida": False,
                        "comentario": "",
                        "metodo_captura": f"Fallback - {metodo}"
                    })
                    self.metadata["proximo_id"] += 1
                    self._salvar_metadata()
                    
                    self.prints.append(caminho_print)
                    print(f"✅ Print salvo (fallback): {caminho_print}")
                except Exception as fallback_error:
                    print(f"❌ Erro no fallback: {fallback_error}")

    # ---------- Navegação e Geração do DOCX ----------
    def gerar_docx(self):
        if not self.prints:
            messagebox.showinfo("Info", "Nenhuma evidência para processar.")
            return
            
        self.current_index = 0
        self.mostrar_janela_navegacao()

    def mostrar_janela_navegacao(self):
        """Janela principal de navegação pelas evidências"""
        self.popup = tk.Toplevel(root)
        self.popup.title("Navegação de Evidências")
        self.popup.geometry("1200x800")
        self.popup.resizable(True, True)
        
        # Configurar grid para melhor organização
        self.popup.grid_columnconfigure(0, weight=1)
        self.popup.grid_rowconfigure(0, weight=1)  # A área da imagem expande
        
        # Frame da imagem (maior para melhor visualização)
        img_frame = tk.Frame(self.popup, bg="white")
        img_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        img_frame.grid_rowconfigure(0, weight=1)
        img_frame.grid_columnconfigure(0, weight=1)
        
        self.current_img_label = tk.Label(img_frame, bg="white")
        self.current_img_label.grid(row=0, column=0, sticky="nsew")
        
        # Frame do comentário (abaixo da imagem)
        comment_frame = tk.Frame(self.popup)
        comment_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=(0, 5))
                
        tk.Label(comment_frame, text="Comentário:", font=("Arial", 11)).pack(anchor="w")
        
        # Criar um frame para o campo de entrada
        comment_entry_frame = tk.Frame(comment_frame)
        comment_entry_frame.pack(fill=tk.X, pady=2)
        
        # Campo de comentário
        self.comment_entry = tk.Entry(comment_entry_frame, font=("Arial", 10))
        self.comment_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.comment_entry.bind("<FocusOut>", lambda e: self.salvar_comentario())
        
        # Frame principal para os botões de navegação e ação
        buttons_main_frame = tk.Frame(self.popup)
        buttons_main_frame.grid(row=2, column=0, sticky="nsew", padx=10, pady=5)
        
        # Frame para centralizar os botões de navegação
        nav_frame = tk.Frame(buttons_main_frame)
        nav_frame.pack(expand=True, pady=2)
        
        # Botões de navegação (centralizados)
        tk.Button(nav_frame, text="⏮️ Primeira", command=self.primeira_evidencia, 
                 width=12).pack(side=tk.LEFT, padx=2)
        tk.Button(nav_frame, text="◀️ Anterior", command=self.anterior_evidencia,
                 width=12).pack(side=tk.LEFT, padx=2)
        
        # Indicador de posição
        self.pos_label = tk.Label(nav_frame, text="", font=("Arial", 12, "bold"))
        self.pos_label.pack(side=tk.LEFT, padx=15)
        
        tk.Button(nav_frame, text="▶️ Próxima", command=self.proxima_evidencia,
                 width=12).pack(side=tk.LEFT, padx=2)
        tk.Button(nav_frame, text="⏭️ Última", command=self.ultima_evidencia,
                 width=12).pack(side=tk.LEFT, padx=2)
        
        # Pular para específica
        tk.Button(nav_frame, text="🔢 Ir para...", command=self.ir_para_especifica,
                 width=12).pack(side=tk.LEFT, padx=2)
        
        # Botões de ação no mesmo nível (Editar e Excluir Print)
        action_frame = tk.Frame(buttons_main_frame)
        action_frame.pack(expand=True, pady=2)
        
        tk.Button(action_frame, text="✏️ Editar Print", command=self.editar_evidencia_atual,
                 width=15).pack(side=tk.LEFT, padx=5)
        tk.Button(action_frame, text="🗑️ Excluir Print", command=self.excluir_evidencia_atual,
                 width=15).pack(side=tk.LEFT, padx=5)
        
        # Frame de controle (parte inferior)
        control_frame = tk.Frame(self.popup)
        control_frame.grid(row=3, column=0, sticky="ew", padx=10, pady=5)
        
        # Frame para centralizar os botões de controle
        control_buttons_frame = tk.Frame(control_frame)
        control_buttons_frame.pack(expand=True)
        
        # Botões na ordem solicitada: Cancelar primeiro, depois Gerar Evidência
        tk.Button(control_buttons_frame, text="❌ Cancelar", command=self.cancelar_processamento,
                 bg="#f44336", fg="white", font=("Arial", 12), 
                 width=15, height=1).pack(side=tk.LEFT, padx=5)
        
        tk.Button(control_buttons_frame, text="✅ Gerar Evidência", command=self.finalizar_processamento,
                 bg="#4CAF50", fg="white", font=("Arial", 12, "bold"), 
                 width=20, height=1).pack(side=tk.LEFT, padx=5)
        
        # Carregar primeira evidência
        self.atualizar_exibicao()
        
        self.popup.protocol("WM_DELETE_WINDOW", self.cancelar_processamento)
        self.popup.grab_set()
    
    def atualizar_exibicao(self):
        """Atualiza a exibição da evidência atual"""
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        
        try:
            # Carrega a imagem original
            img = Image.open(caminho_print).convert("RGBA")
            
            # Obter os metadados do timestamp para esta evidência
            nome_arquivo = os.path.basename(caminho_print)
            timestamp_data = self.obter_timestamp_metadata(nome_arquivo)
            
            # 🔥 DIFERENÇA CRÍTICA: Só aplica timestamp se for modo "ocultar" E tiver texto
            if timestamp_data and timestamp_data["texto"] and self.modo_captura == "ocultar":
                draw = ImageDraw.Draw(img)
                
                # Calcular posição em pixels
                img_width, img_height = img.size
                pos_x_percent = timestamp_data["x"]
                pos_y_percent = timestamp_data["y"]
                
                # Configurações do texto
                texto = timestamp_data["texto"]
                texto_cor = timestamp_data["cor"]
                fundo_cor = timestamp_data.get("fundo", "#000000B2")
                tamanho = timestamp_data.get("tamanho", self.TIMESTAMP_TAMANHO_PADRAO)
                
                # Usar fonte
                try:
                    font = ImageFont.truetype("arial.ttf", tamanho)
                except:
                    font = ImageFont.load_default()
                
                # 🔥 CALCULAR TAMANHO DO TEXTO PARA CENTRALIZAR
                bbox = draw.textbbox((0, 0), texto, font=font)
                texto_largura = bbox[2] - bbox[0]
                texto_altura = bbox[3] - bbox[1]
                
                # 🔥 DEFINIR PADDING E CANTOS ARREDONDADOS
                padding_horizontal = 20
                padding_vertical = 12
                borda_radius = 8
                
                # 🔥 CALCULAR POSIÇÃO FINAL DO FUNDO (centralizado na posição especificada)
                fundo_largura = texto_largura + (padding_horizontal * 2)
                fundo_altura = texto_altura + (padding_vertical * 2)
                
                # Calcular coordenadas do fundo baseado na posição percentual
                fundo_x1 = int((img_width * pos_x_percent) - (fundo_largura / 2))  # Centralizado horizontalmente
                fundo_y1 = int((img_height * pos_y_percent) - (fundo_altura / 2))   # Centralizado verticalmente
                fundo_x2 = fundo_x1 + fundo_largura
                fundo_y2 = fundo_y1 + fundo_altura
                
                # 🔥 GARANTIR que o fundo não saia dos limites da imagem
                margem = 10
                if fundo_x1 < margem:
                    fundo_x1 = margem
                    fundo_x2 = fundo_x1 + fundo_largura
                elif fundo_x2 > img_width - margem:
                    fundo_x2 = img_width - margem
                    fundo_x1 = fundo_x2 - fundo_largura
                    
                if fundo_y1 < margem:
                    fundo_y1 = margem
                    fundo_y2 = fundo_y1 + fundo_altura
                elif fundo_y2 > img_height - margem:
                    fundo_y2 = img_height - margem
                    fundo_y1 = fundo_y2 - fundo_altura
                
                # 🔥 CORREÇÃO CRÍTICA: Calcular a posição vertical do texto considerando a métrica da fonte
                bbox = draw.textbbox((0, 0), texto, font=font)
                texto_ascendente = -bbox[1]  # A parte "acima" da linha de base
                texto_largura = bbox[2] - bbox[0]
                texto_altura_total = bbox[3] - bbox[1]

                # 🔥 CALCULAR POSIÇÃO DO TEXTO (verdadeiramente centralizado no fundo)
                texto_x = fundo_x1 + padding_horizontal
                # Ajuste fino para centralização vertical perfeita
                texto_y = fundo_y1 + (fundo_altura - texto_altura_total) // 2 + texto_ascendente

                # 🔥 DESENHAR FUNDO COM CANTOS ARREDONDADOS
                # Criar máscara para cantos arredondados
                mask = Image.new("L", (fundo_largura, fundo_altura), 0)
                mask_draw = ImageDraw.Draw(mask)
                mask_draw.rounded_rectangle(
                    [0, 0, fundo_largura, fundo_altura],
                    radius=borda_radius,
                    fill=255
                )

                # Aplicar fundo semi-transparente
                if fundo_cor.startswith("#") and len(fundo_cor) == 9:
                    r = int(fundo_cor[1:3], 16)
                    g = int(fundo_cor[3:5], 16)
                    b = int(fundo_cor[5:7], 16)
                    a = int(fundo_cor[7:9], 16)
                    fundo_rgba = (r, g, b, a)
                else:
                    fundo_rgba = (0, 0, 0, 178)

                fundo_img = Image.new("RGBA", (fundo_largura, fundo_altura), fundo_rgba)
                img.paste(fundo_img, (fundo_x1, fundo_y1), mask)

                # 🔥 DESENHAR TEXTO BRANCO VERDADEIRAMENTE CENTRALIZADO
                draw.text((texto_x, texto_y), texto, fill=texto_cor, font=font)
            
            # Obter o tamanho da área disponível para a imagem
            self.popup.update()
            available_width = self.popup.winfo_width() - 40
            available_height = self.popup.winfo_height() - 250
            
            # Ajustar a imagem para caber na área disponível
            img.thumbnail((available_width, available_height))
            self.current_img_tk = ImageTk.PhotoImage(img)
            self.current_img_label.config(image=self.current_img_tk)
            
            # Atualiza indicador de posição
            self.pos_label.config(text=f"Evidência {self.current_index + 1} de {len(self.prints)}")
            
            # Carrega comentário salvo
            comentario = self.obter_comentario(nome_arquivo)
            self.comment_entry.delete(0, tk.END)
            self.comment_entry.insert(0, comentario)
            
        except Exception as e:
            print(f"Erro ao carregar imagem: {e}")

    def obter_comentario(self, nome_arquivo):
        """Obtém o comentário salvo nos metadados"""
        for evidencia in self.metadata["evidencias"]:
            if evidencia["arquivo"] == nome_arquivo:
                return evidencia.get("comentario", "")
        return ""

    def obter_timestamp_metadata(self, nome_arquivo):
        """Obtém os metadados do timestamp para um arquivo"""
        for evidencia in self.metadata["evidencias"]:
            if evidencia["arquivo"] == nome_arquivo:
                return {
                    "x": evidencia["timestamp_posicao"]["x"],
                    "y": evidencia["timestamp_posicao"]["y"],
                    "cor": evidencia["timestamp_cor"],
                    "fundo": evidencia.get("timestamp_fundo", "#000000B2"),
                    "tamanho": evidencia.get("timestamp_tamanho", self.TIMESTAMP_TAMANHO_PADRAO),
                    "texto": evidencia["timestamp_texto"]
                }
        return None

    def atualizar_posicao_timestamp(self, nome_arquivo, nova_posicao):
        """Atualiza a posição do timestamp nos metadados"""
        for evidencia in self.metadata["evidencias"]:
            if evidencia["arquivo"] == nome_arquivo:
                evidencia["timestamp_posicao"]["x"] = nova_posicao[0]
                evidencia["timestamp_posicao"]["y"] = nova_posicao[1]
                self._salvar_metadata()
                break

    def aplicar_timestamp_na_imagem(self, caminho_imagem, evidencia_meta):
        """Aplica o timestamp na imagem conforme posição salva"""
        # 🔥 ALTERADO: Usar a nova função moderna
        if self.modo_captura != "manter":
            self.aplicar_timestamp_moderno(caminho_imagem, evidencia_meta)

    def salvar_comentario(self):
        """Salva o comentário da evidência atual"""
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        nome_arquivo = os.path.basename(caminho_print)
        comentario = self.comment_entry.get()
        
        # Atualiza metadados
        for evidencia in self.metadata["evidencias"]:
            if evidencia["arquivo"] == nome_arquivo:
                evidencia["comentario"] = comentario
                break
                
        self._salvar_metadata()        

    # Métodos de navegação
    def primeira_evidencia(self):
        self.salvar_comentario()
        self.current_index = 0
        self.atualizar_exibicao()

    def anterior_evidencia(self):
        self.salvar_comentario()
        if self.current_index > 0:
            self.current_index -= 1
            self.atualizar_exibicao()

    def proxima_evidencia(self):
        self.salvar_comentario()
        if self.current_index < len(self.prints) - 1:
            self.current_index += 1
            self.atualizar_exibicao()

    def ultima_evidencia(self):
        self.salvar_comentario()
        self.current_index = len(self.prints) - 1
        self.atualizar_exibicao()

    def ir_para_especifica(self):
        self.salvar_comentario()
        if not self.prints:
            return
            
        numero = simpledialog.askinteger("Navegar", 
                                       f"Digite o número da evidência (1-{len(self.prints)}):",
                                       minvalue=1, maxvalue=len(self.prints))
        if numero:
            self.current_index = numero - 1
            self.atualizar_exibicao()

    def editar_evidencia_atual(self):
        self.salvar_comentario()
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        self.abrir_editor(caminho_print, self.popup)
        self.atualizar_exibicao()

    def excluir_evidencia_atual(self):
        self.salvar_comentario()
        if not self.prints or self.current_index >= len(self.prints):
            return
            
        caminho_print = self.prints[self.current_index]
        nome_arquivo = os.path.basename(caminho_print)
        
        if messagebox.askyesno("Confirmar Exclusão", 
                             "Tem certeza que deseja excluir este print?"):
            try:
                # Remove arquivo físico
                os.remove(caminho_print)
                
                # Marca como excluída nos metadados
                for evidencia in self.metadata["evidencias"]:
                    if evidencia["arquivo"] == nome_arquivo:
                        evidencia["excluida"] = True
                        break
                
                self._salvar_metadata()
                
                # Recarrega a lista de evidências
                self.recarregar_evidencias()
                
                if not self.prints:
                    messagebox.showinfo("Info", "Todas as evidências foram processadas.")
                    self.finalizar_processamento()
                    return
                
                # Ajusta o índice se necessário
                if self.current_index >= len(self.prints):
                    self.current_index = len(self.prints) - 1
                
                self.atualizar_exibicao()
                messagebox.showinfo("Sucesso", "Evidência excluída!")
                
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao excluir: {str(e)}")

    def finalizar_processamento(self):
        """Processa todas as evidências e gera o DOCX"""
        self.salvar_comentario()

        # 🔥 DIFERENÇA CRÍTICA: Só aplica timestamp se for modo "ocultar"
        if self.modo_captura == "ocultar":
            # Aplicar timestamp em todas as evidências
            for caminho_print in self.prints:
                nome_arquivo = os.path.basename(caminho_print)
                # Encontrar os metadados da evidência
                for evidencia in self.metadata["evidencias"]:
                    if evidencia["arquivo"] == nome_arquivo and evidencia["timestamp_texto"]:
                        # 🔥 ALTERADO: Usar a nova função moderna
                        self.aplicar_timestamp_moderno(caminho_print, evidencia)
                        break

        # Agora adicionar as imagens ao DOCX
        for caminho_print in self.prints:
            nome_arquivo = os.path.basename(caminho_print)
            comentario = self.obter_comentario(nome_arquivo)
            
            self.doc.add_picture(caminho_print, width=Inches(6))
            if comentario.strip():
                self.doc.add_paragraph(comentario)
        
        self.salvar_docx()
        if self.popup:
            self.popup.destroy()

    def cancelar_processamento(self):
        self.salvar_comentario()
        if messagebox.askyesno("Confirmar Cancelamento", 
                              "Tem certeza que deseja cancelar o processamento?\n\n"
                              "⚠️ TODOS os arquivos de print serão EXCLUÍDOS permanentemente!"):
            
            try:
                # Excluir todos os arquivos de print
                prints_excluidos = 0
                for caminho_print in self.prints:
                    try:
                        if os.path.exists(caminho_print):
                            os.remove(caminho_print)
                            prints_excluidos += 1
                            print(f"Print excluído: {caminho_print}")
                    except Exception as e:
                        print(f"Erro ao excluir print {caminho_print}: {e}")

                # Excluir arquivo de metadados
                if self.metadata_path and os.path.exists(self.metadata_path):
                    try:
                        os.remove(self.metadata_path)
                        print(f"Metadata excluído: {self.metadata_path}")
                    except Exception as e:
                        print(f"Erro ao excluir metadata: {e}")

                # Mostrar mensagem de confirmação
                messagebox.showinfo(
                    "Cancelado", 
                    f"Processamento cancelado!\n\n"
                    f"Foram excluídos:\n"
                    f"• {prints_excluidos} arquivos de print\n"
                    f"• Arquivo de metadados"
                )

                # Fechar a janela de navegação
                if self.popup:
                    self.popup.destroy()

            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao excluir arquivos: {str(e)}")

    def salvar_docx(self):
        if self.template_path:
            nome_base = os.path.basename(self.template_path)
            if nome_base.lower().endswith('.docx'):
                nome_base = nome_base[:-5]
            nome_arquivo = f"{nome_base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        else:
            nome_arquivo = f"Evidencias_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        caminho_save = os.path.join(self.output_dir, nome_arquivo)
        
        try:
            # Salva o documento
            self.doc.save(caminho_save)

            # 🔥 CORREÇÃO: SEMPRE excluir o arquivo de metadados (controle interno)
            if self.metadata_path and os.path.exists(self.metadata_path):
                try:
                    os.remove(self.metadata_path)
                    print(f"Metadata excluído: {self.metadata_path}")
                except Exception as e:
                    print(f"Erro ao excluir metadata: {e}")

            # Verificar se deve excluir os prints (apenas os arquivos de imagem)
            manter = self.manter_evidencias if self.manter_evidencias is not None else True

            if not manter:
                # Excluir todos os arquivos de print (apenas as imagens)
                prints_excluidos = 0
                for caminho_print in self.prints:
                    try:
                        if os.path.exists(caminho_print):
                            os.remove(caminho_print)
                            prints_excluidos += 1
                            print(f"Print excluído: {caminho_print}")
                    except Exception as e:
                        print(f"Erro ao excluir print {caminho_print}: {e}")
                
                mensagem_exclusao = f"\n\nExclusão realizada:\n- {prints_excluidos} arquivos de evidência excluídos"
            else:
                mensagem_exclusao = "\n\nArquivos de evidência mantidos na pasta."

            # Mensagem de sucesso
            messagebox.showinfo(
                "Concluído", 
                f"Documento gerado com sucesso!\nSalvo em:\n{caminho_save}{mensagem_exclusao}"
            )
            
            # Abrir a pasta onde foi salvo
            if os.name == 'nt':
                os.startfile(self.output_dir)
            elif os.name == 'posix':
                import subprocess
                if sys.platform == 'darwin':
                    subprocess.Popen(['open', self.output_dir])
                else:
                    subprocess.Popen(['xdg-open', self.output_dir])
                    
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao salvar documento: {str(e)}")

    # ---------- Editor de prints (mantido completo) ----------
    def abrir_editor(self, caminho_print, parent):
        editor = tk.Toplevel(parent)
        editor.title("Editor de Evidência")
        editor.geometry("1200x800")
        
        # Frame principal
        main_frame = tk.Frame(editor)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Frame para ferramentas e opções
        tools_frame = tk.Frame(main_frame)
        tools_frame.pack(side=tk.TOP, fill=tk.X, pady=5)
        
        # Frame para a área de desenho
        canvas_frame = tk.Frame(main_frame)
        canvas_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        # 🔥 CORREÇÃO: SEMPRE carregar a imagem ORIGINAL (sem timestamp)
        img_original = Image.open(caminho_print).convert("RGBA")
        
        # Obter metadados do timestamp
        nome_arquivo = os.path.basename(caminho_print)
        timestamp_data = self.obter_timestamp_metadata(nome_arquivo)
        
        # 🔥 CORREÇÃO CRÍTICA: Sempre usar a posição salva nos metadados
        if timestamp_data:
            self.timestamp_pos = (timestamp_data["x"], timestamp_data["y"])
        else:
            # 🔥 Usar a posição padrão definida nas constantes
            self.timestamp_pos = (self.TIMESTAMP_POSICAO_PADRAO_X, self.TIMESTAMP_POSICAO_PADRAO_Y)
            # Criar dados básicos se não existirem
            timestamp_data = {
                "texto": datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
                "cor": "#FFFFFF",
                "fundo": "#000000B2"
            }
                
        # Usar a imagem original como base para edição
        self.original_img = img_original
        img_w, img_h = self.original_img.size        
     
        # Calcula o fator de escala para exibição
        max_w, max_h = 1000, 700
        scale = min(max_w / img_w, max_h / img_h)
        self.scale_factor = scale
        disp_w, disp_h = int(img_w * scale), int(img_h * scale)
        
        # Cria cópia da imagem para edição (SEM timestamp)
        self.editing_img = self.original_img.copy()
        self.display_img = self.editing_img.resize((disp_w, disp_h), Image.LANCZOS)

        # Variáveis para controle
        self.current_tk_img = ImageTk.PhotoImage(self.display_img)
        self.elements = []  # Lista de elementos desenhados
        self.undo_stack = []  # PILHA PARA DESFAZER AÇÕES
        self.temp_element = None
        
        # 🔥 CORREÇÃO: Controle do timestamp - sempre usar a posição dos metadados
        self.moving_timestamp = False
        self.timestamp_drag_data = {"x": 0, "y": 0, "item": None}
        self.last_mouse_pos = None
        
        # Canvas para a imagem
        self.canvas = tk.Canvas(canvas_frame, width=disp_w, height=disp_h, cursor="cross", bg="gray")
        self.canvas.pack(padx=5, pady=5)
        self.canvas_img = self.canvas.create_image(0, 0, anchor="nw", image=self.current_tk_img)
            
        # Variáveis de controle
        tool_var = tk.StringVar(value="rectangle")
        color_var = tk.StringVar(value="#FF0000")
        width_var = tk.IntVar(value=3)
        
        # Variável para controle do modo mover timestamp
        move_timestamp_var = tk.BooleanVar(value=False)
        
        # 🔥 CORREÇÃO: Adicionar as funções que estavam faltando
        def start_move_timestamp(event):
            """Inicia o movimento do timestamp"""
            if move_timestamp_var.get():
                self.moving_timestamp = True
                self.timestamp_drag_data["x"] = event.x
                self.timestamp_drag_data["y"] = event.y
                self.last_mouse_pos = (event.x, event.y)
                self.canvas.config(cursor="fleur")  # Cursor de movimento

        def stop_move_timestamp(event):
            """Finaliza o movimento do timestamp"""
            if self.moving_timestamp:
                self.moving_timestamp = False
                if move_timestamp_var.get():
                    self.canvas.config(cursor="hand2")
                else:
                    self.canvas.config(cursor="cross")

        def toggle_move_timestamp():
            """Ativa/desativa o modo de mover timestamp"""
            current_state = move_timestamp_var.get()
            move_timestamp_var.set(not current_state)

            if move_timestamp_var.get():
                # Ativa modo mover timestamp
                move_btn.config(relief=tk.SUNKEN, bg="#4CAF50", fg="white", 
                              text="📅 MODO MOVER ATIVO")
                self.canvas.config(cursor="hand2")
            else:
                # Desativa modo mover timestamp
                move_btn.config(relief=tk.RAISED, bg="SystemButtonFace", fg="black",
                              text="📅 Mover Data/Hora")
                self.canvas.config(cursor="cross")
            
            refresh_display()

        # FERRAMENTAS DE DESENHO
        tk.Label(tools_frame, text="Ferramentas:", font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=5)
        
        icon_frame = tk.Frame(tools_frame)
        icon_frame.pack(side=tk.LEFT, padx=5)
        
        def criar_botao_ferramenta(parent, texto, valor, variavel):
            btn = tk.Radiobutton(parent, text=texto, font=("Arial", 12), 
                               variable=variavel, value=valor, indicatoron=0, 
                               width=3, height=2, relief=tk.RAISED,
                               cursor="hand2")
            return btn

        # Criar botões para cada ferramenta
        tool_icons = {
            "rectangle": "⬜",   # Retângulo
            "circle": "🔴",      # Círculo  
            "arrow": "👉",       # Seta
            "text": "🆎"         # Texto
        }

        for tool_value, icon in tool_icons.items():
            btn = criar_botao_ferramenta(icon_frame, icon, tool_value, tool_var)
            btn.pack(side=tk.LEFT, padx=2)

        # Destacar o botão selecionado inicialmente
        for widget in icon_frame.winfo_children():
            if isinstance(widget, tk.Radiobutton) and widget.cget("value") == "rectangle":
                widget.config(relief=tk.SUNKEN, bg="#e3f2fd")
                break

        def update_button_appearance(*args):
            selected_tool = tool_var.get()
            for widget in icon_frame.winfo_children():
                if isinstance(widget, tk.Radiobutton):
                    if widget.cget("value") == selected_tool:
                        widget.config(relief=tk.SUNKEN, bg="#e3f2fd")
                    else:
                        widget.config(relief=tk.RAISED, bg="SystemButtonFace")

        tool_var.trace("w", update_button_appearance)
        
        # Controles de cor e espessura
        color_frame = tk.Frame(tools_frame)
        color_frame.pack(side=tk.LEFT, padx=20)
        
        tk.Label(color_frame, text="Cor:").pack(side=tk.LEFT)
        
        colors = ["#FF0000", "#00FF00", "#FFFF00", "#0000FF", "#000000", "#FFFFFF"]
        color_buttons_frame = tk.Frame(color_frame)
        color_buttons_frame.pack(side=tk.LEFT, padx=5)
        
        for color in colors:
            btn = tk.Button(color_buttons_frame, bg=color, width=2, height=1, 
                           command=lambda c=color: self.set_color(color_var, c, color_preview))
            btn.pack(side=tk.LEFT, padx=1)
        
        custom_btn = tk.Button(color_frame, text="Personalizada", 
                              command=lambda: self.choose_custom_color(editor, color_var, color_preview))
        custom_btn.pack(side=tk.LEFT, padx=5)
        
        color_preview = tk.Frame(color_frame, width=30, height=20, bg=color_var.get())
        color_preview.pack(side=tk.LEFT, padx=5)
        
        width_frame = tk.Frame(tools_frame)
        width_frame.pack(side=tk.LEFT, padx=20)
        
        tk.Label(width_frame, text="Espessura:").pack(side=tk.LEFT)
        tk.Scale(width_frame, from_=1, to=10, variable=width_var, orient=tk.HORIZONTAL, 
                length=100, showvalue=1).pack(side=tk.LEFT, padx=5)
        
        def undo_action():
            if self.elements:
                removed_element = self.elements.pop()
                self.undo_stack.append(removed_element)
                refresh_display()
        
        undo_btn = tk.Button(tools_frame, text="↩️ Desfazer (Ctrl+Z)", command=undo_action)
        undo_btn.pack(side=tk.LEFT, padx=20)
        
        # 🔥 CORREÇÃO MELHORADA: FUNÇÃO PRINCIPAL DE ATUALIZAÇÃO DA TELA COM MESMA LÓGICA DO TIMESTAMP
        def refresh_display():
            """Redesenha toda a cena: imagem base + timestamp visual + elementos"""
            # Limpa o canvas
            self.canvas.delete("all")
            
            # Redesenha a imagem ORIGINAL (sem timestamp)
            self.display_img = self.editing_img.resize((disp_w, disp_h), Image.LANCZOS)
            self.current_tk_img = ImageTk.PhotoImage(self.display_img)
            self.canvas.create_image(0, 0, anchor="nw", image=self.current_tk_img)
            
            # 🔥 CORREÇÃO CRÍTICA: Usar a MESMA lógica do aplicar_timestamp_moderno
            if self.timestamp_pos and self.modo_captura != "manter":
                img_width, img_height = self.original_img.size
                pos_x_percent = self.timestamp_pos[0]
                pos_y_percent = self.timestamp_pos[1]
                
                texto = timestamp_data["texto"] if timestamp_data else datetime.now().strftime('%d/%m/%Y %H:%M:%S')
                cor = timestamp_data["cor"] if timestamp_data else "#FFFFFF"
                fundo_cor = "#000000B2"  # Usar o mesmo fundo semi-transparente
                tamanho = 12  # Tamanho reduzido para visualização no canvas
                
                # Usar tkfont para medir o texto (mesma lógica do método principal)
                font = tkfont.Font(family="Arial", size=tamanho, weight="bold")
                text_width = font.measure(texto)
                text_height = font.metrics("linespace")
                
                # 🔥 USAR MESMO CÁLCULO DO MÉTODO PRINCIPAL
                padding_horizontal = 10  # Reduzido para exibição no canvas
                padding_vertical = 6     # Reduzido para exibição no canvas
                
                # Calcular dimensões do fundo
                fundo_largura = text_width + (padding_horizontal * 2)
                fundo_altura = text_height + (padding_vertical * 2)
                
                # Calcular posição (centralizada como no método principal)
                pos_x = int(pos_x_percent * img_width * self.scale_factor)
                pos_y = int(pos_y_percent * img_height * self.scale_factor)
                
                fundo_x1 = pos_x - (fundo_largura // 2)  # Centralizado horizontalmente
                fundo_y1 = pos_y - (fundo_altura // 2)   # Centralizado verticalmente
                fundo_x2 = fundo_x1 + fundo_largura
                fundo_y2 = fundo_y1 + fundo_altura
                
                # Calcular posição do texto (centralizado no fundo)
                texto_x = fundo_x1 + padding_horizontal
                texto_y = fundo_y1 + padding_vertical
                
                # 🔥 DESENHAR FUNDO PRETO SEMI-TRANSPARENTE
                self.canvas.create_rectangle(
                    fundo_x1, fundo_y1, fundo_x2, fundo_y2,
                    fill="#000000", outline="", stipple="gray50",  # Stipple para efeito de transparência
                    tags="timestamp_bg"
                )
                
                # 🔥 DESENHAR TEXTO BRANCO CENTRALIZADO
                self.canvas.create_text(
                    texto_x, texto_y, 
                    text=texto, 
                    fill=cor, 
                    font=("Arial", tamanho, "bold"), 
                    anchor="nw",
                    tags="timestamp"
                )
            
            # Redesenha elementos de desenho (setas, círculos, etc.)
            for element in self.elements:
                elem_type, coords, color, width, text = element
                scaled_coords = [int(c * self.scale_factor) for c in coords]
                
                if elem_type == "circle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_oval(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "rectangle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "arrow":
                    x1, y1, x2, y2 = scaled_coords
                    self.draw_arrow_on_canvas(x1, y1, x2, y2, color, width)
                elif elem_type == "text":
                    x, y = scaled_coords
                    self.canvas.create_text(x, y, text=text, fill=color, font=("Arial", 12), anchor="nw")
            
            # Desenha elemento temporário durante a criação
            if self.temp_element:
                elem_type, coords, color, width, text = self.temp_element
                scaled_coords = [int(c * self.scale_factor) for c in coords]
                
                if elem_type == "circle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_oval(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "rectangle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "arrow":
                    x1, y1, x2, y2 = scaled_coords
                    self.draw_arrow_on_canvas(x1, y1, x2, y2, color, width)

        def on_motion_timestamp(event):
            if self.moving_timestamp and move_timestamp_var.get():
                current_mouse_pos = (event.x, event.y)
                
                if (abs(current_mouse_pos[0] - self.last_mouse_pos[0]) > 2 or 
                    abs(current_mouse_pos[1] - self.last_mouse_pos[1]) > 2):
                    
                    dx = event.x - self.timestamp_drag_data["x"]
                    dy = event.y - self.timestamp_drag_data["y"]
                    
                    img_width, img_height = self.original_img.size
                    
                    # 🔥 CORREÇÃO: Calcular a nova posição considerando o deslocamento
                    # Converter coordenadas de tela para percentuais
                    new_x = self.timestamp_pos[0] + (dx / (img_width * self.scale_factor))
                    new_y = self.timestamp_pos[1] + (dy / (img_height * self.scale_factor))
                    
                    # Limitar aos limites da imagem
                    new_x = max(0.05, min(0.95, new_x))  # Deixar margem para o fundo
                    new_y = max(0.05, min(0.95, new_y))  # Deixar margem para o fundo
                    
                    self.timestamp_pos = (new_x, new_y)
                    refresh_display()
                    
                    self.timestamp_drag_data["x"] = event.x
                    self.timestamp_drag_data["y"] = event.y
                    self.last_mouse_pos = current_mouse_pos

        def stop_move_timestamp(event):
            if self.moving_timestamp:
                self.moving_timestamp = False
                if move_timestamp_var.get():
                    self.canvas.config(cursor="hand2")
                else:
                    self.canvas.config(cursor="cross")

        # FUNÇÕES PARA DESENHO
        start_xy = {"x": None, "y": None}
        
        def on_button_press(event):
            if move_timestamp_var.get():
                start_move_timestamp(event)
            else:
                start_xy["x"], start_xy["y"] = event.x, event.y

        def on_motion(event):
            if self.moving_timestamp and move_timestamp_var.get():
                on_motion_timestamp(event)
            elif start_xy["x"] is not None and not move_timestamp_var.get():
                sx, sy = start_xy["x"], start_xy["y"]
                ex, ey = event.x, event.y
                
                ix1, iy1 = int(sx / self.scale_factor), int(sy / self.scale_factor)
                ix2, iy2 = int(ex / self.scale_factor), int(ey / self.scale_factor)
                
                tool = tool_var.get()
                color = color_var.get()
                width = width_var.get()
                
                if tool == "circle":
                    radius = int(((ix2 - ix1)**2 + (iy2 - iy1)**2)**0.5)
                    self.temp_element = ("circle", [ix1-radius, iy1-radius, ix1+radius, iy1+radius], color, width, "")
                elif tool == "rectangle":
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.temp_element = ("rectangle", [x1_norm, y1_norm, x2_norm, y2_norm], color, width, "")
                elif tool == "arrow":
                    self.temp_element = ("arrow", [ix1, iy1, ix2, iy2], color, width, "")
                
                refresh_display()

        def on_button_release(event):
            if self.moving_timestamp:
                stop_move_timestamp(event)
            elif start_xy["x"] is not None and not move_timestamp_var.get():
                sx, sy = start_xy["x"], start_xy["y"]
                ex, ey = event.x, event.y
                
                ix1, iy1 = int(sx / self.scale_factor), int(sy / self.scale_factor)
                ix2, iy2 = int(ex / self.scale_factor), int(ey / self.scale_factor)
                
                tool = tool_var.get()
                color = color_var.get()
                width = width_var.get()
                
                self.undo_stack.clear()
                
                if tool == "circle":
                    radius = int(((ix2 - ix1)**2 + (iy2 - iy1)**2)**0.5)
                    self.elements.append(("circle", [ix1-radius, iy1-radius, ix1+radius, iy1+radius], color, width, ""))
                
                elif tool == "rectangle":
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.elements.append(("rectangle", [x1_norm, y1_norm, x2_norm, y2_norm], color, width, ""))
                
                elif tool == "arrow":
                    self.elements.append(("arrow", [ix1, iy1, ix2, iy2], color, width, ""))
                
                elif tool == "text":
                    text = simpledialog.askstring("Texto", "Digite o texto:", parent=editor)
                    if text:
                        self.elements.append(("text", [ix1, iy1], color, width, text))
                        refresh_display()
                
                self.temp_element = None
                refresh_display()
            
            start_xy["x"], start_xy["y"] = None, None

        def on_key_press(event):
            if event.keysym.lower() == 'z' and (event.state & 0x4):
                undo_action()

        editor.bind('<Control-z>', on_key_press)
        editor.bind('<Control-Z>', on_key_press)
        
        self.canvas.bind("<Button-1>", on_button_press)
        self.canvas.bind("<B1-Motion>", on_motion)
        self.canvas.bind("<ButtonRelease-1>", on_button_release)
        
        def update_cursor(*args):
            if move_timestamp_var.get():
                self.canvas.config(cursor="hand2")
            elif tool_var.get() == "text":
                self.canvas.config(cursor="xterm")
            else:
                self.canvas.config(cursor="cross")
        
        tool_var.trace("w", update_cursor)
        move_timestamp_var.trace("w", lambda *args: update_cursor())
        
        refresh_display()
        
        # BOTÕES DE AÇÃO
        action_frame = tk.Frame(editor)
        action_frame.pack(side=tk.BOTTOM, pady=10)
        
        def salvar_edicao():
            """Salva a imagem com elementos de desenho e atualiza a posição do timestamp"""
            try:
                # 🔥 CORREÇÃO: Carrega a imagem ORIGINAL (sem timestamp)
                final_img = Image.open(caminho_print).convert("RGBA")
                draw = ImageDraw.Draw(final_img)
                
                # Aplica elementos de desenho
                for element in self.elements:
                    elem_type, coords, color, width, text = element
                    
                    if elem_type == "circle":
                        x1, y1, x2, y2 = coords
                        draw.ellipse([x1, y1, x2, y2], outline=color, width=width)
                    
                    elif elem_type == "rectangle":
                        x1, y1, x2, y2 = coords
                        draw.rectangle([x1, y1, x2, y2], outline=color, width=width)
                    
                    elif elem_type == "arrow":
                        x1, y1, x2, y2 = coords
                        draw.line([x1, y1, x2, y2], fill=color, width=width)
                        
                        angle = math.atan2(y2 - y1, x2 - x1)
                        arrow_size = 15
                        x3 = x2 - arrow_size * math.cos(angle - math.pi/6)
                        y3 = y2 - arrow_size * math.sin(angle - math.pi/6)
                        x4 = x2 - arrow_size * math.cos(angle + math.pi/6)
                        y4 = y2 - arrow_size * math.sin(angle + math.pi/6)
                        
                        draw.polygon([x2, y2, x3, y3, x4, y4], fill=color, outline=color)
                    
                    elif elem_type == "text":
                        x, y = coords
                        try:
                            font_text = ImageFont.truetype("arial.ttf", 20)
                        except:
                            font_text = ImageFont.load_default()
                        draw.text((x, y), text, fill=color, font=font_text)
                
                # 🔥 CORREÇÃO: NÃO aplica o timestamp aqui - será aplicado apenas na geração do DOCX
                # Apenas atualiza a posição nos metadados
                
                # Salva a imagem com elementos de desenho (sem timestamp)
                final_img.save(caminho_print)
                
                # 🔥 CORREÇÃO: Atualiza metadados com nova posição do timestamp
                for evidencia in self.metadata["evidencias"]:
                    if evidencia["arquivo"] == nome_arquivo:
                        evidencia["timestamp_posicao"]["x"] = self.timestamp_pos[0]
                        evidencia["timestamp_posicao"]["y"] = self.timestamp_pos[1]
                        # 🔥 GARANTIR que o tamanho sempre exista e seja consistente
                        evidencia["timestamp_tamanho"] = evidencia.get("timestamp_tamanho", self.TIMESTAMP_TAMANHO_PADRAO)
                        break
                                
                self._salvar_metadata()
                messagebox.showinfo("Sucesso", "Edição salva! A data/hora será aplicada na geração do documento.")
                editor.destroy()
                
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao salvar: {str(e)}")

        def cancelar_edicao():
            if messagebox.askyesno("Confirmar", "Descartar todas as alterações?"):
                editor.destroy()
        
        tk.Button(action_frame, text="💾 Salvar", command=salvar_edicao, width=15).pack(side=tk.LEFT, padx=5)
        
        if self.modo_captura != "manter":
            move_btn = tk.Button(action_frame, text="📅 Mover Data/Hora", 
                            command=toggle_move_timestamp, relief=tk.RAISED,
                            cursor="hand2", width=18)    
            move_btn.pack(side=tk.LEFT, padx=5)
 
        
        tk.Button(action_frame, text="❌ Cancelar", command=cancelar_edicao, width=15).pack(side=tk.LEFT, padx=5)

        editor.transient(parent)
        editor.grab_set()
        editor.focus_set()
        parent.wait_window(editor)

    # Métodos auxiliares para o editor
    def set_color(self, color_var, color, preview_widget):
        color_var.set(color)
        preview_widget.config(bg=color)

    def choose_custom_color(self, parent, color_var, preview_widget):
        # Fecha qualquer janela de cor anterior que possa estar aberta
        if hasattr(self, 'color_chooser_window') and self.color_chooser_window:
            try:
                self.color_chooser_window.destroy()
            except:
                pass
        
        # Abre a nova janela de seleção de cor
        color = colorchooser.askcolor(title="Escolha uma cor", initialcolor=color_var.get(), parent=parent)
        if color[1]:
            color_var.set(color[1])
            preview_widget.config(bg=color[1])

    def draw_arrow_on_canvas(self, x1, y1, x2, y2, color, width):
        # Desenha a linha da seta
        self.canvas.create_line(x1, y1, x2, y2, fill=color, width=width)
        
        # Calcula o ângulo da seta
        angle = math.atan2(y2 - y1, x2 - x1)
        
        # Desenha the ponta da seta (triângulo)
        arrow_size = 15
        x3 = x2 - arrow_size * math.cos(angle - math.pi/6)
        y3 = y2 - arrow_size * math.sin(angle - math.pi/6)
        x4 = x2 - arrow_size * math.cos(angle + math.pi/6)
        y4 = y2 - arrow_size * math.sin(angle + math.pi/6)
        
        self.canvas.create_polygon(x2, y2, x3, y3, x4, y4, fill=color, outline=color)

# ------------------ Funções utilitárias ------------------
def minimizar_janela():
    try:
        root.iconify()
    except:
        pass

# ------------------ Interface ------------------
gravador = GravadorDocx()

def iniciar():
    if gravador.mostrar_janela_configuracao():
        pass

def pausar():
    gravador.pausar()

def retomar():
    gravador.retomar()

def finalizar():
    gravador.finalizar()

# ------------------ Janela principal ------------------
if __name__ == "__main__":
    root = tk.Tk()
    root.title("PrintF - Capturar Evidências")
    root.geometry("500x400")

    def on_closing():
        root.quit()

    root.protocol("WM_DELETE_WINDOW", on_closing)

    tk.Label(root, text="📷 PrintF - Capturar Evidências", font=("Arial", 16, "bold")).pack(pady=10)
    tk.Button(root, text="▶ Iniciar Gravação (F8)", command=lambda: root.after(0, iniciar), width=30).pack(pady=5)
    tk.Button(root, text="⏸ Pausar Gravação (F6)", command=lambda: root.after(0, pausar), width=30).pack(pady=5)
    tk.Button(root, text="▶ Retomar Gravação (F7)", command=lambda: root.after(0, retomar), width=30).pack(pady=5)
    tk.Button(root, text="⏹ Finalizar Gravação (F9)", command=lambda: root.after(0, finalizar), width=30).pack(pady=5)
    tk.Button(root, text="❌ Fechar Aplicativo (F12)", command=on_closing, width=30).pack(pady=8)

# ------------------ Atalhos globais ------------------
def create_global_hotkeys():
    """Cria atalhos globais que funcionam mesmo com browser em foco"""
    try:
        from pynput import keyboard as pynput_keyboard
        
        def on_activate_f6():
            print("F6 pressionado - Pausando")
            root.after(0, pausar)
            
        def on_activate_f7():
            print("F7 pressionado - Retomando") 
            root.after(0, retomar)
            
        def on_activate_f8():
            print("F8 pressionado - Iniciando")
            root.after(0, iniciar)
            
        def on_activate_f9():
            print("F9 pressionado - Finalizando")
            root.after(0, finalizar)
            
        def on_activate_f12():
            print("F12 pressionado - Fechando")
            root.after(0, root.quit)

        # 🔥 CORREÇÃO: Criar hotkeys globais com pynput
        with pynput_keyboard.GlobalHotKeys({
                '<f6>': on_activate_f6,
                '<f7>': on_activate_f7, 
                '<f8>': on_activate_f8,
                '<f9>': on_activate_f9,
                '<f12>': on_activate_f12}) as h:
            
            print("✅ Atalhos globais configurados - F6, F7, F8, F9, F12")
            h.join()
            
    except Exception as e:
        print(f"❌ Erro nos atalhos globais: {e}")
        # Fallback para o método antigo
        def on_press(key):
            try:
                if key == pynput_keyboard.Key.f6:
                    root.after(0, pausar)
                elif key == pynput_keyboard.Key.f7:
                    root.after(0, retomar)
                elif key == pynput_keyboard.Key.f8:
                    root.after(0, iniciar)
                elif key == pynput_keyboard.Key.f9:
                    root.after(0, finalizar)
                elif key == pynput_keyboard.Key.f12:
                    root.after(0, root.quit)
            except Exception as e:
                print(f"Erro no atalho: {e}")

        listener = pynput_keyboard.Listener(on_press=on_press)
        listener.start()

# Iniciar os hotkeys globais em uma thread separada
import threading
hotkey_thread = threading.Thread(target=create_global_hotkeys, daemon=True)
hotkey_thread.start()

# ------------------ Inicia interface ------------------
root.mainloop()