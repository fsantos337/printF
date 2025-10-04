"""
PrintF - Sistema de Captura de Evidências
Aplicação para capturar screenshots e gerar documentos Word com evidências.
"""

import os
import sys
import json
import math
import time
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass, field

import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, colorchooser, ttk
import tkinter.font as tkfont

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import pyautogui
from pynput import mouse, keyboard
from PIL import Image, ImageTk, ImageDraw, ImageFont, ImageGrab

# Imports opcionais
try:
    import win32gui
    import win32con
    import win32api
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False

try:
    import mss
    MSS_AVAILABLE = True
except ImportError:
    MSS_AVAILABLE = False


# ==================== CONSTANTES ====================
class CaptureMode:
    """Modos de captura disponíveis"""
    MANTER = "manter"  # Mantém barra de tarefas
    OCULTAR = "ocultar"  # Oculta barra de tarefas


class Colors:
    """Cores padrão da aplicação"""
    RED = "#FF0000"
    GREEN = "#00FF00"
    YELLOW = "#FFFF00"
    BLUE = "#0000FF"
    BLACK = "#000000"
    WHITE = "#FFFFFF"
    
    TIMESTAMP_TEXT = "#FFFFFF"
    TIMESTAMP_BG = "#000000B2"


class Dimensions:
    """Dimensões padrão"""
    CLICK_RADIUS = 40
    ARROW_SIZE = 15
    TIMESTAMP_SIZE = 24
    TIMESTAMP_PADDING = 10
    TIMESTAMP_BORDER_RADIUS = 8


# ==================== DATACLASSES ====================
@dataclass
class TimestampConfig:
    """Configuração do timestamp"""
    texto: str
    posicao: Tuple[float, float]
    cor: str = Colors.TIMESTAMP_TEXT
    fundo: str = Colors.TIMESTAMP_BG
    tamanho: int = Dimensions.TIMESTAMP_SIZE


@dataclass
class EvidenciaMetadata:
    """Metadados de uma evidência"""
    id: int
    arquivo: str
    timestamp: str
    excluida: bool = False
    comentario: str = ""
    metodo_captura: str = ""
    timestamp_config: Optional[TimestampConfig] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Converte para dicionário"""
        data = {
            "id": self.id,
            "arquivo": self.arquivo,
            "timestamp": self.timestamp,
            "excluida": self.excluida,
            "comentario": self.comentario,
            "metodo_captura": self.metodo_captura
        }
        
        if self.timestamp_config:
            data.update({
                "timestamp_texto": self.timestamp_config.texto,
                "timestamp_posicao": {
                    "x": self.timestamp_config.posicao[0],
                    "y": self.timestamp_config.posicao[1]
                },
                "timestamp_cor": self.timestamp_config.cor,
                "timestamp_fundo": self.timestamp_config.fundo,
                "timestamp_tamanho": self.timestamp_config.tamanho
            })
        
        return data
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'EvidenciaMetadata':
        """Cria instância a partir de dicionário"""
        timestamp_config = None
        if "timestamp_texto" in data:
            timestamp_config = TimestampConfig(
                texto=data["timestamp_texto"],
                posicao=(
                    data["timestamp_posicao"]["x"],
                    data["timestamp_posicao"]["y"]
                ),
                cor=data.get("timestamp_cor", Colors.TIMESTAMP_TEXT),
                fundo=data.get("timestamp_fundo", Colors.TIMESTAMP_BG),
                tamanho=data.get("timestamp_tamanho", Dimensions.TIMESTAMP_SIZE)
            )
        
        return cls(
            id=data["id"],
            arquivo=data["arquivo"],
            timestamp=data["timestamp"],
            excluida=data.get("excluida", False),
            comentario=data.get("comentario", ""),
            metodo_captura=data.get("metodo_captura", ""),
            timestamp_config=timestamp_config
        )


@dataclass
class DrawElement:
    """Elemento de desenho no editor"""
    tipo: str  # "circle", "rectangle", "arrow", "text"
    coords: List[int]
    cor: str
    largura: int
    texto: str = ""


# ==================== GERENCIADOR DE CAPTURA ====================
class CapturaManager:
    """Gerencia a captura de tela em diferentes modos"""
    
    def __init__(self, modo: str = CaptureMode.OCULTAR):
        self.modo = modo
    
    def capturar(self, x: int, y: int) -> Tuple[Optional[Image.Image], Tuple[int, int], str]:
        """
        Captura a tela baseado no modo configurado
        
        Args:
            x: Coordenada X do clique
            y: Coordenada Y do clique
            
        Returns:
            Tupla (imagem, coordenadas_relativas, método_usado)
        """
        if self.modo == CaptureMode.MANTER:
            return self._capturar_tela_completa(x, y)
        else:
            return self._capturar_area_trabalho(x, y)
    
    def _capturar_tela_completa(self, x: int, y: int) -> Tuple[Optional[Image.Image], Tuple[int, int], str]:
        """Captura tela completa incluindo barra de tarefas"""
        try:
            # Estratégia 1: Win32 API + MSS
            if WIN32_AVAILABLE:
                try:
                    monitor_handle = win32api.MonitorFromPoint(
                        (x, y), 
                        win32con.MONITOR_DEFAULTTONEAREST
                    )
                    monitor_info = win32gui.GetMonitorInfo(monitor_handle)
                    monitor_area = monitor_info["Monitor"]
                    
                    if MSS_AVAILABLE:
                        with mss.mss() as sct:
                            monitor_mss = {
                                "left": monitor_area[0],
                                "top": monitor_area[1],
                                "width": monitor_area[2] - monitor_area[0],
                                "height": monitor_area[3] - monitor_area[1]
                            }
                            screenshot = sct.grab(monitor_mss)
                            img = Image.frombytes(
                                "RGB", 
                                screenshot.size, 
                                screenshot.bgra, 
                                "raw", 
                                "BGRX"
                            )
                            
                            rel_x = x - monitor_area[0]
                            rel_y = y - monitor_area[1]
                            
                            metodo = f"Win32 + MSS Monitor Completo {monitor_area}"
                            print(f"✅ CAPTURA WIN32+MSS - Monitor {monitor_area} | Coord: ({rel_x},{rel_y})")
                            
                            return img, (rel_x, rel_y), metodo
                    else:
                        screenshot = ImageGrab.grab(bbox=monitor_area)
                        rel_x = x - monitor_area[0]
                        rel_y = y - monitor_area[1]
                        
                        metodo = f"Win32 Monitor Completo {monitor_area}"
                        print(f"✅ CAPTURA WIN32 - Monitor {monitor_area} | Coord: ({rel_x},{rel_y})")
                        
                        return screenshot, (rel_x, rel_y), metodo
                        
                except Exception as e:
                    print(f"❌ Win32 falhou, tentando MSS: {e}")
            
            # Estratégia 2: MSS como alternativa
            if MSS_AVAILABLE:
                with mss.mss() as sct:
                    monitor_encontrado = None
                    
                    for monitor in sct.monitors[1:]:
                        if (monitor["left"] <= x < monitor["left"] + monitor["width"] and
                            monitor["top"] <= y < monitor["top"] + monitor["height"]):
                            monitor_encontrado = monitor
                            break
                    
                    if not monitor_encontrado:
                        monitor_encontrado = sct.monitors[1]
                        print("⚠️ Monitor não encontrado, usando primário")
                    
                    screenshot = sct.grab(monitor_encontrado)
                    img = Image.frombytes(
                        "RGB", 
                        screenshot.size, 
                        screenshot.bgra, 
                        "raw", 
                        "BGRX"
                    )
                    
                    rel_x = x - monitor_encontrado["left"]
                    rel_y = y - monitor_encontrado["top"]
                    
                    metodo = f"MSS Monitor Completo {monitor_encontrado['width']}x{monitor_encontrado['height']}"
                    print(f"✅ CAPTURA MSS - Monitor {monitor_encontrado} | Coord: ({rel_x},{rel_y})")
                    
                    return img, (rel_x, rel_y), metodo
            
            # Fallback: pyautogui
            screenshot = pyautogui.screenshot()
            return screenshot, (x, y), "Fallback - pyautogui (apenas primário)"
            
        except Exception as e:
            print(f"❌ Falha na captura completa: {e}")
            screenshot = pyautogui.screenshot()
            return screenshot, (x, y), f"Fallback - Erro: {str(e)}"
    
    def _capturar_area_trabalho(self, x: int, y: int) -> Tuple[Optional[Image.Image], Tuple[int, int], str]:
        """Captura apenas área de trabalho (sem barra de tarefas)"""
        try:
            screenshot = pyautogui.screenshot()
            metodo = "PyAutoGUI - Área de Trabalho (sem barra)"
            print(f"✅ CAPTURA PYAUTOGUI - Área de Trabalho | Coord: ({x},{y})")
            
            return screenshot, (x, y), metodo
            
        except Exception as e:
            print(f"❌ Falha na captura com pyautogui: {e}")
            screenshot = pyautogui.screenshot()
            return screenshot, (x, y), f"Fallback - Erro: {str(e)}"


# ==================== GERENCIADOR DE METADADOS ====================
class MetadataManager:
    """Gerencia os metadados das evidências"""
    
    def __init__(self, caminho: str):
        self.caminho = caminho
        self.evidencias: List[EvidenciaMetadata] = []
        self.proximo_id = 1
        self._carregar()
    
    def _carregar(self):
        """Carrega metadados do arquivo JSON"""
        if os.path.exists(self.caminho):
            try:
                with open(self.caminho, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.proximo_id = data.get("proximo_id", 1)
                    self.evidencias = [
                        EvidenciaMetadata.from_dict(ev) 
                        for ev in data.get("evidencias", [])
                    ]
            except Exception as e:
                print(f"Erro ao carregar metadados: {e}")
                self.evidencias = []
                self.proximo_id = 1
    
    def salvar(self):
        """Salva metadados no arquivo JSON"""
        data = {
            "evidencias": [ev.to_dict() for ev in self.evidencias],
            "proximo_id": self.proximo_id
        }
        
        try:
            with open(self.caminho, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Erro ao salvar metadados: {e}")
    
    def adicionar_evidencia(self, evidencia: EvidenciaMetadata):
        """Adiciona uma nova evidência"""
        self.evidencias.append(evidencia)
        self.proximo_id += 1
        self.salvar()
    
    def obter_evidencia(self, arquivo: str) -> Optional[EvidenciaMetadata]:
        """Obtém evidência pelo nome do arquivo"""
        for ev in self.evidencias:
            if ev.arquivo == arquivo:
                return ev
        return None
    
    def atualizar_evidencia(self, arquivo: str, **kwargs):
        """Atualiza campos de uma evidência"""
        for ev in self.evidencias:
            if ev.arquivo == arquivo:
                for key, value in kwargs.items():
                    if hasattr(ev, key):
                        setattr(ev, key, value)
                self.salvar()
                break
    
    def listar_evidencias_ativas(self) -> List[EvidenciaMetadata]:
        """Lista evidências não excluídas"""
        return [ev for ev in self.evidencias if not ev.excluida]


# ==================== PROCESSADOR DE IMAGENS ====================
class ImageProcessor:
    """Processa imagens para adicionar marcações e timestamps"""
    
    @staticmethod
    def adicionar_marcacao_clique(
        img: Image.Image, 
        x: int, 
        y: int, 
        raio: int = Dimensions.CLICK_RADIUS
    ) -> Image.Image:
        """Adiciona círculo amarelo no ponto de clique"""
        img_rgba = img.convert("RGBA")
        overlay = Image.new("RGBA", img_rgba.size, (255, 255, 255, 0))
        draw = ImageDraw.Draw(overlay)
        
        draw.ellipse(
            (x - raio, y - raio, x + raio, y + raio),
            fill=(255, 255, 0, 100)
        )
        
        return Image.alpha_composite(img_rgba, overlay)
    
    @staticmethod
    def aplicar_timestamp_moderno(
        img_path: str,
        timestamp_config: TimestampConfig
    ):
        """Aplica timestamp com fundo moderno na imagem"""
        img = Image.open(img_path).convert("RGBA")
        draw = ImageDraw.Draw(img)
        
        img_width, img_height = img.size
        pos_x = int(timestamp_config.posicao[0] * img_width)
        pos_y = int(timestamp_config.posicao[1] * img_height)
        
        # Converter cor de fundo
        fundo_rgba = ImageProcessor._hex_to_rgba(timestamp_config.fundo)
        
        # Configurar fonte
        try:
            font = ImageFont.truetype("arial.ttf", timestamp_config.tamanho)
        except:
            font = ImageFont.load_default()
        
        # Calcular dimensões do texto
        bbox = draw.textbbox((0, 0), timestamp_config.texto, font=font)
        texto_largura = bbox[2] - bbox[0]
        texto_altura = bbox[3] - bbox[1]
        
        # Criar fundo com cantos arredondados
        padding = Dimensions.TIMESTAMP_PADDING
        radius = Dimensions.TIMESTAMP_BORDER_RADIUS
        
        fundo_coords = (
            pos_x - padding,
            pos_y - padding,
            pos_x + texto_largura + padding,
            pos_y + texto_altura + padding
        )
        
        # Criar máscara para cantos arredondados
        mask_size = (
            fundo_coords[2] - fundo_coords[0],
            fundo_coords[3] - fundo_coords[1]
        )
        mask = Image.new("L", mask_size, 0)
        mask_draw = ImageDraw.Draw(mask)
        mask_draw.rounded_rectangle(
            [0, 0, mask_size[0], mask_size[1]],
            radius=radius,
            fill=255
        )
        
        # Aplicar fundo
        fundo_img = Image.new("RGBA", mask_size, fundo_rgba)
        img.paste(fundo_img, (fundo_coords[0], fundo_coords[1]), mask)
        
        # Desenhar texto
        draw.text(
            (pos_x, pos_y),
            timestamp_config.texto,
            fill=timestamp_config.cor,
            font=font
        )
        
        img.save(img_path)
    
    @staticmethod
    def _hex_to_rgba(hex_color: str) -> Tuple[int, int, int, int]:
        """Converte cor hexadecimal para RGBA"""
        if hex_color.startswith("#") and len(hex_color) == 9:
            r = int(hex_color[1:3], 16)
            g = int(hex_color[3:5], 16)
            b = int(hex_color[5:7], 16)
            a = int(hex_color[7:9], 16)
            return (r, g, b, a)
        return (0, 0, 0, 178)


# ==================== CLASSE PRINCIPAL ====================
class GravadorDocx:
    """Classe principal para gravação de evidências"""
    
    def __init__(self):
        self.gravando = False
        self.pausado = False
        self.output_dir = os.getcwd()
        self.listener_mouse: Optional[mouse.Listener] = None
        self.prints: List[str] = []
        self.doc: Optional[Document] = None
        self.using_template = False
        self.template_path: Optional[str] = None
        self.current_index = 0
        self.evidence_dir: Optional[str] = None
        self.metadata_manager: Optional[MetadataManager] = None
        self.manter_evidencias: Optional[bool] = None
        self.modo_captura = CaptureMode.OCULTAR
        self.captura_manager: Optional[CapturaManager] = None
        
        # Componentes da UI
        self.popup: Optional[tk.Toplevel] = None
        self.current_img_label: Optional[tk.Label] = None
        self.current_img_tk: Optional[ImageTk.PhotoImage] = None
        self.comment_entry: Optional[tk.Entry] = None
        self.pos_label: Optional[tk.Label] = None
    
    def mostrar_janela_configuracao(self) -> bool:
        """Mostra janela de configuração inicial"""
        config_window = tk.Toplevel(root)
        config_window.title("Configuração de Gravação")
        config_window.geometry("600x600")
        config_window.resizable(False, False)
        config_window.transient(root)
        config_window.grab_set()
        
        main_frame = ttk.Frame(config_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Título
        ttk.Label(
            main_frame,
            text="PrintF - Configuração de Gravação",
            font=("Arial", 16, "bold")
        ).pack(pady=10)
        
        # Seleção de template
        self._criar_selecao_template(main_frame)
        
        # Seleção de diretório
        self._criar_selecao_diretorio(main_frame)
        
        # Modo de captura
        modo_captura_var = self._criar_selecao_modo_captura(main_frame)
        
        # Opção de manter evidências
        manter_evidencias_var = self._criar_opcao_manter_evidencias(main_frame)
        
        # Botões de ação
        self._criar_botoes_configuracao(
            main_frame,
            config_window,
            modo_captura_var,
            manter_evidencias_var
        )
        
        root.wait_window(config_window)
        return self.template_path is not None and self.output_dir is not None
    
    def _criar_selecao_template(self, parent: ttk.Frame):
        """Cria seleção de template"""
        ttk.Label(parent, text="Selecione o template DOCX:").pack(
            anchor="w", pady=(10, 5)
        )
        
        template_frame = ttk.Frame(parent)
        template_frame.pack(fill=tk.X, pady=5)
        
        self.template_var = tk.StringVar()
        template_entry = ttk.Entry(
            template_frame,
            textvariable=self.template_var,
            width=40
        )
        template_entry.pack(side=tk.LEFT, padx=(0, 5), fill=tk.X, expand=True)
        
        def selecionar_template():
            template_path = filedialog.askopenfilename(
                title="Selecione o template DOCX",
                filetypes=[("Documentos Word", "*.docx")]
            )
            if template_path:
                self.template_var.set(template_path)
        
        ttk.Button(
            template_frame,
            text="Procurar",
            command=selecionar_template
        ).pack(side=tk.RIGHT)
    
    def _criar_selecao_diretorio(self, parent: ttk.Frame):
        """Cria seleção de diretório"""
        ttk.Label(parent, text="Selecione o diretório de destino:").pack(
            anchor="w", pady=(20, 5)
        )
        
        dir_frame = ttk.Frame(parent)
        dir_frame.pack(fill=tk.X, pady=5)
        
        self.dir_var = tk.StringVar()
        dir_entry = ttk.Entry(dir_frame, textvariable=self.dir_var, width=40)
        dir_entry.pack(side=tk.LEFT, padx=(0, 5), fill=tk.X, expand=True)
        
        def selecionar_diretorio():
            dir_path = filedialog.askdirectory(
                title="Selecione o diretório para salvar"
            )
            if dir_path:
                self.dir_var.set(dir_path)
        
        ttk.Button(dir_frame, text="Procurar", command=selecionar_diretorio).pack(
            side=tk.RIGHT
        )
    
    def _criar_selecao_modo_captura(self, parent: ttk.Frame) -> tk.StringVar:
        """Cria seleção do modo de captura"""
        ttk.Label(
            parent,
            text="Modo de Captura da Barra de Tarefas:",
            font=("Arial", 11, "bold")
        ).pack(anchor="w", pady=(20, 10))
        
        modo_captura_var = tk.StringVar(value=CaptureMode.OCULTAR)
        
        modo_frame = ttk.Frame(parent)
        modo_frame.pack(fill=tk.X, pady=5)
        
        ttk.Radiobutton(
            modo_frame,
            text="Manter barra de tarefas (data/hora visível na barra do Windows)",
            variable=modo_captura_var,
            value=CaptureMode.MANTER
        ).pack(anchor="w", pady=2)
        
        ttk.Radiobutton(
            modo_frame,
            text="Ocultar barra de tarefas (data/hora será adicionada na imagem)",
            variable=modo_captura_var,
            value=CaptureMode.OCULTAR
        ).pack(anchor="w", pady=2)
        
        return modo_captura_var
    
    def _criar_opcao_manter_evidencias(self, parent: ttk.Frame) -> tk.BooleanVar:
        """Cria opção de manter evidências"""
        ttk.Label(
            parent,
            text="Opções de saída:",
            font=("Arial", 11, "bold")
        ).pack(anchor="w", pady=(20, 10))
        
        manter_evidencias_var = tk.BooleanVar(value=True)
        
        checkbox_frame = ttk.Frame(parent)
        checkbox_frame.pack(fill=tk.X, pady=5)
        
        ttk.Checkbutton(
            checkbox_frame,
            text="Manter arquivos de evidência (prints) na pasta após gerar o DOCX",
            variable=manter_evidencias_var
        ).pack(anchor="w")
        
        ttk.Label(
            parent,
            text="Se desmarcado, os arquivos de print serão excluídos após a geração do DOCX.",
            font=("Arial", 9),
            foreground="gray",
            justify=tk.LEFT
        ).pack(anchor="w", pady=(5, 15))
        
        return manter_evidencias_var
    
    def _criar_botoes_configuracao(
        self,
        parent: ttk.Frame,
        config_window: tk.Toplevel,
        modo_captura_var: tk.StringVar,
        manter_evidencias_var: tk.BooleanVar
    ):
        """Cria botões de configuração"""
        btn_frame = ttk.Frame(parent)
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=(20, 0))
        
        def iniciar_com_config():
            if not self._validar_configuracao():
                return
            
            self._aplicar_configuracao(modo_captura_var, manter_evidencias_var)
            config_window.destroy()
            self.iniciar_gravacao()
        
        button_container = ttk.Frame(btn_frame)
        button_container.pack(expand=True)
        
        ttk.Button(
            button_container,
            text="Iniciar Gravação",
            command=iniciar_com_config
        ).pack(side=tk.LEFT, padx=10)
        
        ttk.Button(
            button_container,
            text="Cancelar",
            command=config_window.destroy
        ).pack(side=tk.LEFT, padx=10)
    
    def _validar_configuracao(self) -> bool:
        """Valida configuração antes de iniciar"""
        if not self.template_var.get() or not self.dir_var.get():
            messagebox.showerror(
                "Erro",
                "Por favor, selecione o template e o diretório de destino."
            )
            return False
        
        if not os.path.exists(self.template_var.get()):
            messagebox.showerror(
                "Erro",
                "O arquivo de template selecionado não existe."
            )
            return False
        
        # Verificar se diretório contém arquivos
        dir_path = self.dir_var.get()
        if os.path.exists(dir_path):
            try:
                for item in os.listdir(dir_path):
                    item_path = os.path.join(dir_path, item)
                    if not item.startswith('.') and os.path.isfile(item_path):
                        messagebox.showerror(
                            "Arquivos na Pasta",
                            "A pasta selecionada contém arquivos.\n\n"
                            "Para evitar misturar evidências, a pasta deve estar vazia "
                            "ou conter apenas outras pastas."
                        )
                        return False
            except PermissionError:
                messagebox.showerror(
                    "Erro de Permissão",
                    "Sem permissão para acessar a pasta selecionada."
                )
                return False
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao verificar a pasta: {str(e)}")
                return False
        
        return True
    
    def _aplicar_configuracao(
        self,
        modo_captura_var: tk.StringVar,
        manter_evidencias_var: tk.BooleanVar
    ):
        """Aplica configuração escolhida"""
        self.modo_captura = modo_captura_var.get()
        self.manter_evidencias = manter_evidencias_var.get()
        self.template_path = self.template_var.get()
        self.output_dir = self.dir_var.get()
        self.evidence_dir = self.dir_var.get()
        
        # Limpar estado
        self.gravando = False
        self.pausado = False
        self.prints = []
    
    def iniciar_gravacao(self):
        """Inicia a gravação de evidências"""
        self._resetar_estado()
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Inicializar gerenciadores
        metadata_path = os.path.join(self.output_dir, "evidencias_metadata.json")
        self.metadata_manager = MetadataManager(metadata_path)
        self.captura_manager = CapturaManager(self.modo_captura)
        
        # Carregar documento
        self._carregar_documento()
        
        messagebox.showinfo("Gravação", "▶ Clique em OK para começar a gravar!")
        print(f"Iniciando gravação com modo: {self.modo_captura}")
        
        self.gravando = True
        self.pausado = False
        
        minimizar_janela()
        
        # Iniciar listener de mouse
        if self.listener_mouse:
            self.listener_mouse.stop()
        
        self.listener_mouse = mouse.Listener(on_click=self.on_click)
        self.listener_mouse.start()
    
    def _resetar_estado(self):
        """Reseta o estado da gravação"""
        self.gravando = False
        self.pausado = False
        self.prints = []
        self.current_index = 0
    
    def _carregar_documento(self):
        """Carrega o documento Word"""
        try:
            if os.path.exists(self.template_path):
                self.doc = Document(self.template_path)
                self.using_template = True
                print("Template carregado com sucesso!")
            else:
                self.doc = Document()
                self.using_template = False
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar template: {str(e)}")
            self.doc = Document()
            self.using_template = False
    
    def pausar(self):
        """Pausa a gravação"""
        if self.gravando and not self.pausado:
            self.pausado = True
            messagebox.showinfo("Gravação", "⏸ Gravação pausada!")
    
    def retomar(self):
        """Retoma a gravação"""
        if self.gravando and self.pausado:
            self.pausado = False
            messagebox.showinfo("Gravação", "▶ Gravação retomada!")
    
    def finalizar(self):
        """Finaliza a gravação"""
        if self.gravando:
            self.gravando = False
            
            if self.listener_mouse:
                self.listener_mouse.stop()
                self.listener_mouse = None
            
            print("Gravação finalizada")
            messagebox.showinfo("Gravação", "⏹ Gravação finalizada!")
            
            if self.prints:
                self.gerar_docx()
            else:
                messagebox.showinfo("Info", "Nenhuma evidência capturada.")
    
    def on_click(self, x: int, y: int, button, pressed: bool):
        """Handler de clique do mouse"""
        if not (pressed and self.gravando and not self.pausado):
            return
        
        # Capturar screenshot
        screenshot, coords_rel, metodo = self.captura_manager.capturar(x, y)
        
        if not screenshot:
            print("Erro: Não foi possível capturar a tela")
            return
        
        # Gerar nome único
        evidencia_id = self.metadata_manager.proximo_id
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nome_arquivo = f"evidencia_{evidencia_id:04d}_{timestamp}.png"
        caminho_print = os.path.join(self.output_dir, nome_arquivo)
        
        try:
            # Adicionar marcação de clique
            click_x, click_y = coords_rel
            img_final = ImageProcessor.adicionar_marcacao_clique(
                screenshot, click_x, click_y
            )
            
            # Salvar imagem
            img_final.convert("RGB").save(caminho_print, "PNG")
            
            # Criar metadados
            timestamp_config = self._criar_timestamp_config()
            evidencia = EvidenciaMetadata(
                id=evidencia_id,
                arquivo=nome_arquivo,
                timestamp=datetime.now().isoformat(),
                metodo_captura=metodo,
                timestamp_config=timestamp_config
            )
            
            self.metadata_manager.adicionar_evidencia(evidencia)
            self.prints.append(caminho_print)
            
            print(f"Print salvo: {caminho_print} | Método: {metodo} | Modo: {self.modo_captura}")
            
        except Exception as e:
            print(f"Erro ao processar captura: {e}")
            self._salvar_screenshot_fallback(screenshot, caminho_print, evidencia_id, metodo)
    
    def _criar_timestamp_config(self) -> Optional[TimestampConfig]:
        """Cria configuração de timestamp baseada no modo"""
        if self.modo_captura == CaptureMode.OCULTAR:
            return TimestampConfig(
                texto=datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
                posicao=(0.75, 0.90)
            )
        return TimestampConfig(texto="", posicao=(0.75, 0.90))
    
    def _salvar_screenshot_fallback(
        self,
        screenshot: Image.Image,
        caminho: str,
        evidencia_id: int,
        metodo: str
    ):
        """Salva screenshot em caso de falha no processamento"""
        try:
            screenshot.save(caminho)
            
            evidencia = EvidenciaMetadata(
                id=evidencia_id,
                arquivo=os.path.basename(caminho),
                timestamp=datetime.now().isoformat(),
                metodo_captura=f"Fallback - {metodo}"
            )
            
            self.metadata_manager.adicionar_evidencia(evidencia)
            self.prints.append(caminho)
        except Exception as e:
            print(f"Erro no fallback: {e}")
    
    def gerar_docx(self):
        """Inicia processo de geração do DOCX"""
        if not self.prints:
            messagebox.showinfo("Info", "Nenhuma evidência para processar.")
            return
        
        self.current_index = 0
        self.mostrar_janela_navegacao()
    
    def mostrar_janela_navegacao(self):
        """Mostra janela de navegação das evidências"""
        self.popup = tk.Toplevel(root)
        self.popup.title("Navegação de Evidências")
        self.popup.geometry("1200x800")
        self.popup.resizable(True, True)
        
        self.popup.grid_columnconfigure(0, weight=1)
        self.popup.grid_rowconfigure(0, weight=1)
        
        # Frame da imagem
        img_frame = tk.Frame(self.popup, bg="white")
        img_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        img_frame.grid_rowconfigure(0, weight=1)
        img_frame.grid_columnconfigure(0, weight=1)
        
        self.current_img_label = tk.Label(img_frame, bg="white")
        self.current_img_label.grid(row=0, column=0, sticky="nsew")
        
        # Frame do comentário
        self._criar_frame_comentario()
        
        # Frame de navegação
        self._criar_frame_navegacao()
        
        # Frame de controle
        self._criar_frame_controle()
        
        self.atualizar_exibicao()
        self.popup.protocol("WM_DELETE_WINDOW", self.cancelar_processamento)
        self.popup.grab_set()
    
    def _criar_frame_comentario(self):
        """Cria frame para comentários"""
        comment_frame = tk.Frame(self.popup)
        comment_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=(0, 5))
        
        tk.Label(comment_frame, text="Comentário:", font=("Arial", 11)).pack(anchor="w")
        
        comment_entry_frame = tk.Frame(comment_frame)
        comment_entry_frame.pack(fill=tk.X, pady=2)
        
        self.comment_entry = tk.Entry(comment_entry_frame, font=("Arial", 10))
        self.comment_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.comment_entry.bind("<FocusOut>", lambda e: self.salvar_comentario())
    
    def _criar_frame_navegacao(self):
        """Cria frame de navegação"""
        buttons_main_frame = tk.Frame(self.popup)
        buttons_main_frame.grid(row=2, column=0, sticky="nsew", padx=10, pady=5)
        
        # Botões de navegação
        nav_frame = tk.Frame(buttons_main_frame)
        nav_frame.pack(expand=True, pady=2)
        
        tk.Button(
            nav_frame, text="⏮️ Primeira",
            command=self.primeira_evidencia, width=12
        ).pack(side=tk.LEFT, padx=2)
        
        tk.Button(
            nav_frame, text="◀️ Anterior",
            command=self.anterior_evidencia, width=12
        ).pack(side=tk.LEFT, padx=2)
        
        self.pos_label = tk.Label(nav_frame, text="", font=("Arial", 12, "bold"))
        self.pos_label.pack(side=tk.LEFT, padx=15)
        
        tk.Button(
            nav_frame, text="▶️ Próxima",
            command=self.proxima_evidencia, width=12
        ).pack(side=tk.LEFT, padx=2)
        
        tk.Button(
            nav_frame, text="⏭️ Última",
            command=self.ultima_evidencia, width=12
        ).pack(side=tk.LEFT, padx=2)
        
        tk.Button(
            nav_frame, text="🔢 Ir para...",
            command=self.ir_para_especifica, width=12
        ).pack(side=tk.LEFT, padx=2)
        
        # Botões de ação
        action_frame = tk.Frame(buttons_main_frame)
        action_frame.pack(expand=True, pady=2)
        
        tk.Button(
            action_frame, text="✏️ Editar Print",
            command=self.editar_evidencia_atual, width=15
        ).pack(side=tk.LEFT, padx=5)
        
        tk.Button(
            action_frame, text="🗑️ Excluir Print",
            command=self.excluir_evidencia_atual, width=15
        ).pack(side=tk.LEFT, padx=5)
    
    def _criar_frame_controle(self):
        """Cria frame de controle"""
        control_frame = tk.Frame(self.popup)
        control_frame.grid(row=3, column=0, sticky="ew", padx=10, pady=5)
        
        control_buttons_frame = tk.Frame(control_frame)
        control_buttons_frame.pack(expand=True)
        
        tk.Button(
            control_buttons_frame, text="❌ Cancelar",
            command=self.cancelar_processamento,
            bg="#f44336", fg="white",
            font=("Arial", 12), width=15, height=1
        ).pack(side=tk.LEFT, padx=5)
        
        tk.Button(
            control_buttons_frame, text="✅ Gerar Evidência",
            command=self.finalizar_processamento,
            bg="#4CAF50", fg="white",
            font=("Arial", 12, "bold"), width=20, height=1
        ).pack(side=tk.LEFT, padx=5)
    
    def atualizar_exibicao(self):
        """Atualiza exibição da evidência atual"""
        if not self.prints or self.current_index >= len(self.prints):
            return
        
        caminho_print = self.prints[self.current_index]
        nome_arquivo = os.path.basename(caminho_print)
        
        try:
            img = Image.open(caminho_print).convert("RGBA")
            
            # Obter metadados
            evidencia = self.metadata_manager.obter_evidencia(nome_arquivo)
            
            # Aplicar timestamp visual se necessário
            if evidencia and evidencia.timestamp_config and \
               evidencia.timestamp_config.texto and \
               self.modo_captura == CaptureMode.OCULTAR:
                img = self._adicionar_timestamp_visual(img, evidencia.timestamp_config)
            
            # Redimensionar para exibição
            self.popup.update()
            available_width = self.popup.winfo_width() - 40
            available_height = self.popup.winfo_height() - 250
            
            img.thumbnail((available_width, available_height))
            self.current_img_tk = ImageTk.PhotoImage(img)
            self.current_img_label.config(image=self.current_img_tk)
            
            # Atualizar indicador
            self.pos_label.config(
                text=f"Evidência {self.current_index + 1} de {len(self.prints)}"
            )
            
            # Carregar comentário
            comentario = evidencia.comentario if evidencia else ""
            self.comment_entry.delete(0, tk.END)
            self.comment_entry.insert(0, comentario)
            
        except Exception as e:
            print(f"Erro ao carregar imagem: {e}")
    
    def _adicionar_timestamp_visual(
        self,
        img: Image.Image,
        config: TimestampConfig
    ) -> Image.Image:
        """Adiciona timestamp visual na imagem (apenas para preview)"""
        draw = ImageDraw.Draw(img)
        
        img_width, img_height = img.size
        scale_factor = min(1000 / img_width, 700 / img_height)
        
        pos_x = int(config.posicao[0] * img_width)
        pos_y = int(config.posicao[1] * img_height)
        
        # Configurar fonte para preview
        tamanho = 12
        try:
            font = tkfont.Font(family="Arial", size=tamanho, weight="bold")
        except:
            font = None
        
        if font:
            text_width = font.measure(config.texto)
            text_height = font.metrics("linespace")
        else:
            text_width = len(config.texto) * 8
            text_height = 15
        
        padding = 8
        
        # Desenhar fundo
        fundo_coords = (
            pos_x - padding,
            pos_y - padding,
            pos_x + text_width + padding,
            pos_y + text_height + padding
        )
        
        draw.rectangle(fundo_coords, fill="#000000", outline="")
        
        # Desenhar texto
        draw.text((pos_x, pos_y), config.texto, fill=config.cor)
        
        return img
    
    def salvar_comentario(self):
        """Salva comentário da evidência atual"""
        if not self.prints or self.current_index >= len(self.prints):
            return
        
        caminho_print = self.prints[self.current_index]
        nome_arquivo = os.path.basename(caminho_print)
        comentario = self.comment_entry.get()
        
        self.metadata_manager.atualizar_evidencia(
            nome_arquivo,
            comentario=comentario
        )
    
    # Métodos de navegação
    def primeira_evidencia(self):
        """Vai para primeira evidência"""
        self.salvar_comentario()
        self.current_index = 0
        self.atualizar_exibicao()
    
    def anterior_evidencia(self):
        """Vai para evidência anterior"""
        self.salvar_comentario()
        if self.current_index > 0:
            self.current_index -= 1
            self.atualizar_exibicao()
    
    def proxima_evidencia(self):
        """Vai para próxima evidência"""
        self.salvar_comentario()
        if self.current_index < len(self.prints) - 1:
            self.current_index += 1
            self.atualizar_exibicao()
    
    def ultima_evidencia(self):
        """Vai para última evidência"""
        self.salvar_comentario()
        self.current_index = len(self.prints) - 1
        self.atualizar_exibicao()
    
    def ir_para_especifica(self):
        """Vai para evidência específica"""
        self.salvar_comentario()
        if not self.prints:
            return
        
        numero = simpledialog.askinteger(
            "Navegar",
            f"Digite o número da evidência (1-{len(self.prints)}):",
            minvalue=1,
            maxvalue=len(self.prints)
        )
        
        if numero:
            self.current_index = numero - 1
            self.atualizar_exibicao()
    
    def editar_evidencia_atual(self):
        """Abre editor para evidência atual"""
        self.salvar_comentario()
        if not self.prints or self.current_index >= len(self.prints):
            return
        
        caminho_print = self.prints[self.current_index]
        editor = EditorEvidencia(self, caminho_print, self.popup)
        editor.abrir()
        self.atualizar_exibicao()
    
    def excluir_evidencia_atual(self):
        """Exclui evidência atual"""
        self.salvar_comentario()
        if not self.prints or self.current_index >= len(self.prints):
            return
        
        caminho_print = self.prints[self.current_index]
        nome_arquivo = os.path.basename(caminho_print)
        
        if not messagebox.askyesno(
            "Confirmar Exclusão",
            "Tem certeza que deseja excluir este print?"
        ):
            return
        
        try:
            os.remove(caminho_print)
            self.metadata_manager.atualizar_evidencia(nome_arquivo, excluida=True)
            
            # Recarregar lista
            self.prints = [
                os.path.join(self.output_dir, ev.arquivo)
                for ev in self.metadata_manager.listar_evidencias_ativas()
            ]
            
            if not self.prints:
                messagebox.showinfo("Info", "Todas as evidências foram processadas.")
                self.finalizar_processamento()
                return
            
            if self.current_index >= len(self.prints):
                self.current_index = len(self.prints) - 1
            
            self.atualizar_exibicao()
            messagebox.showinfo("Sucesso", "Evidência excluída!")
            
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao excluir: {str(e)}")
    
    def finalizar_processamento(self):
        """Finaliza processamento e gera DOCX"""
        self.salvar_comentario()
        
        # Aplicar timestamps se modo ocultar
        if self.modo_captura == CaptureMode.OCULTAR:
            for caminho_print in self.prints:
                nome_arquivo = os.path.basename(caminho_print)
                evidencia = self.metadata_manager.obter_evidencia(nome_arquivo)
                
                if evidencia and evidencia.timestamp_config and \
                   evidencia.timestamp_config.texto:
                    ImageProcessor.aplicar_timestamp_moderno(
                        caminho_print,
                        evidencia.timestamp_config
                    )
        
        # Adicionar imagens ao documento
        for caminho_print in self.prints:
            nome_arquivo = os.path.basename(caminho_print)
            evidencia = self.metadata_manager.obter_evidencia(nome_arquivo)
            
            self.doc.add_picture(caminho_print, width=Inches(6))
            
            if evidencia and evidencia.comentario.strip():
                self.doc.add_paragraph(evidencia.comentario)
        
        self.salvar_docx()
        
        if self.popup:
            self.popup.destroy()
    
    def cancelar_processamento(self):
        """Cancela processamento e exclui arquivos"""
        self.salvar_comentario()
        
        if not messagebox.askyesno(
            "Confirmar Cancelamento",
            "Tem certeza que deseja cancelar o processamento?\n\n"
            "⚠️ TODOS os arquivos de print serão EXCLUÍDOS permanentemente!"
        ):
            return
        
        try:
            prints_excluidos = 0
            for caminho_print in self.prints:
                try:
                    if os.path.exists(caminho_print):
                        os.remove(caminho_print)
                        prints_excluidos += 1
                except Exception as e:
                    print(f"Erro ao excluir print {caminho_print}: {e}")
            
            # Excluir metadados
            if self.metadata_manager and os.path.exists(self.metadata_manager.caminho):
                os.remove(self.metadata_manager.caminho)
            
            messagebox.showinfo(
                "Cancelado",
                f"Processamento cancelado!\n\n"
                f"Foram excluídos:\n"
                f"• {prints_excluidos} arquivos de print\n"
                f"• Arquivo de metadados"
            )
            
            if self.popup:
                self.popup.destroy()
                
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao excluir arquivos: {str(e)}")
    
    def salvar_docx(self):
        """Salva o documento DOCX"""
        # Gerar nome do arquivo
        if self.template_path:
            nome_base = os.path.basename(self.template_path)
            if nome_base.lower().endswith('.docx'):
                nome_base = nome_base[:-5]
            nome_arquivo = f"{nome_base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        else:
            nome_arquivo = f"Evidencias_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        caminho_save = os.path.join(self.output_dir, nome_arquivo)
        
        try:
            self.doc.save(caminho_save)
            
            # Sempre excluir metadados
            if self.metadata_manager and os.path.exists(self.metadata_manager.caminho):
                os.remove(self.metadata_manager.caminho)
            
            # Excluir prints se configurado
            mensagem_exclusao = ""
            if not self.manter_evidencias:
                prints_excluidos = 0
                for caminho_print in self.prints:
                    try:
                        if os.path.exists(caminho_print):
                            os.remove(caminho_print)
                            prints_excluidos += 1
                    except Exception as e:
                        print(f"Erro ao excluir print {caminho_print}: {e}")
                
                mensagem_exclusao = f"\n\nExclusão realizada:\n- {prints_excluidos} arquivos de evidência excluídos"
            else:
                mensagem_exclusao = "\n\nArquivos de evidência mantidos na pasta."
            
            messagebox.showinfo(
                "Concluído",
                f"Documento gerado com sucesso!\nSalvo em:\n{caminho_save}{mensagem_exclusao}"
            )
            
            # Abrir pasta
            self._abrir_pasta_destino()
            
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao salvar documento: {str(e)}")
    
    def _abrir_pasta_destino(self):
        """Abre a pasta de destino no explorador"""
        if os.name == 'nt':
            os.startfile(self.output_dir)
        elif os.name == 'posix':
            import subprocess
            if sys.platform == 'darwin':
                subprocess.Popen(['open', self.output_dir])
            else:
                subprocess.Popen(['xdg-open', self.output_dir])


# ==================== EDITOR DE EVIDÊNCIAS ====================
class EditorEvidencia:
    """Editor gráfico de evidências"""
    
    def __init__(self, gravador: GravadorDocx, caminho_print: str, parent: tk.Toplevel):
        self.gravador = gravador
        self.caminho_print = caminho_print
        self.parent = parent
        self.nome_arquivo = os.path.basename(caminho_print)
        
        # Estado do editor
        self.elements: List[DrawElement] = []
        self.undo_stack: List[DrawElement] = []
        self.temp_element: Optional[DrawElement] = None
        
        # Controle de timestamp
        self.moving_timestamp = False
        self.timestamp_drag_data = {"x": 0, "y": 0, "item": None}
        self.last_mouse_pos: Optional[Tuple[int, int]] = None
        
        # Carregar metadados
        self.evidencia = self.gravador.metadata_manager.obter_evidencia(self.nome_arquivo)
        if self.evidencia and self.evidencia.timestamp_config:
            self.timestamp_pos = self.evidencia.timestamp_config.posicao
        else:
            self.timestamp_pos = (0.75, 0.90)
        
        # Carregar imagem original
        self.original_img = Image.open(caminho_print).convert("RGBA")
        self.editing_img = self.original_img.copy()
        
        # Calcular escala
        img_w, img_h = self.original_img.size
        max_w, max_h = 1000, 700
        self.scale_factor = min(max_w / img_w, max_h / img_h)
        self.disp_w = int(img_w * self.scale_factor)
        self.disp_h = int(img_h * self.scale_factor)
        
        self.display_img = self.editing_img.resize(
            (self.disp_w, self.disp_h),
            Image.LANCZOS
        )
        
        # Componentes UI
        self.editor: Optional[tk.Toplevel] = None
        self.canvas: Optional[tk.Canvas] = None
        self.current_tk_img: Optional[ImageTk.PhotoImage] = None
    
    def abrir(self):
        """Abre a janela do editor"""
        self.editor = tk.Toplevel(self.parent)
        self.editor.title("Editor de Evidência")
        self.editor.geometry("1200x800")
        
        # Frames principais
        main_frame = tk.Frame(self.editor)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        tools_frame = tk.Frame(main_frame)
        tools_frame.pack(side=tk.TOP, fill=tk.X, pady=5)
        
        canvas_frame = tk.Frame(main_frame)
        canvas_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
        
        # Criar canvas
        self.canvas = tk.Canvas(
            canvas_frame,
            width=self.disp_w,
            height=self.disp_h,
            cursor="cross",
            bg="gray"
        )
        self.canvas.pack(padx=5, pady=5)
        
        # Variáveis de controle
        tool_var = tk.StringVar(value="rectangle")
        color_var = tk.StringVar(value=Colors.RED)
        width_var = tk.IntVar(value=3)
        move_timestamp_var = tk.BooleanVar(value=False)
        
        # Criar ferramentas
        self._criar_ferramentas(
            tools_frame,
            tool_var,
            color_var,
            width_var,
            move_timestamp_var
        )
        
        # Criar botões de ação
        self._criar_botoes_acao(move_timestamp_var)
        
        # Configurar eventos
        self._configurar_eventos(tool_var, move_timestamp_var)
        
        # Desenhar inicial
        self.refresh_display()
        
        self.editor.transient(self.parent)
        self.editor.grab_set()
        self.editor.focus_set()
        self.parent.wait_window(self.editor)
    
    def _criar_ferramentas(
        self,
        parent: tk.Frame,
        tool_var: tk.StringVar,
        color_var: tk.StringVar,
        width_var: tk.IntVar,
        move_timestamp_var: tk.BooleanVar
    ):
        """Cria as ferramentas de desenho"""
        # Ferramentas de desenho
        tk.Label(parent, text="Ferramentas:", font=("Arial", 10, "bold")).pack(
            side=tk.LEFT, padx=5
        )
        
        icon_frame = tk.Frame(parent)
        icon_frame.pack(side=tk.LEFT, padx=5)
        
        tools = {
            "rectangle": "⬜",
            "circle": "🔴",
            "arrow": "👉",
            "text": "🆎"
        }
        
        for value, icon in tools.items():
            btn = tk.Radiobutton(
                icon_frame,
                text=icon,
                font=("Arial", 12),
                variable=tool_var,
                value=value,
                indicatoron=0,
                width=3,
                height=2,
                relief=tk.RAISED,
                cursor="hand2"
            )
            btn.pack(side=tk.LEFT, padx=2)
        
        # Seleção de cor
        self._criar_selecao_cor(parent, color_var)
        
        # Controle de espessura
        width_frame = tk.Frame(parent)
        width_frame.pack(side=tk.LEFT, padx=20)
        
        tk.Label(width_frame, text="Espessura:").pack(side=tk.LEFT)
        tk.Scale(
            width_frame,
            from_=1,
            to=10,
            variable=width_var,
            orient=tk.HORIZONTAL,
            length=100,
            showvalue=1
        ).pack(side=tk.LEFT, padx=5)
        
        # Botão desfazer
        tk.Button(
            parent,
            text="↩️ Desfazer (Ctrl+Z)",
            command=self.undo_action
        ).pack(side=tk.LEFT, padx=20)
    
    def _criar_selecao_cor(self, parent: tk.Frame, color_var: tk.StringVar):
        """Cria seleção de cor"""
        color_frame = tk.Frame(parent)
        color_frame.pack(side=tk.LEFT, padx=20)
        
        tk.Label(color_frame, text="Cor:").pack(side=tk.LEFT)
        
        colors = [
            Colors.RED, Colors.GREEN, Colors.YELLOW,
            Colors.BLUE, Colors.BLACK, Colors.WHITE
        ]
        
        color_buttons_frame = tk.Frame(color_frame)
        color_buttons_frame.pack(side=tk.LEFT, padx=5)
        
        color_preview = tk.Frame(color_frame, width=30, height=20, bg=color_var.get())
        
        for color in colors:
            btn = tk.Button(
                color_buttons_frame,
                bg=color,
                width=2,
                height=1,
                command=lambda c=color: self._set_color(color_var, c, color_preview)
            )
            btn.pack(side=tk.LEFT, padx=1)
        
        tk.Button(
            color_frame,
            text="Personalizada",
            command=lambda: self._choose_custom_color(color_var, color_preview)
        ).pack(side=tk.LEFT, padx=5)
        
        color_preview.pack(side=tk.LEFT, padx=5)
    
    def _set_color(self, color_var: tk.StringVar, color: str, preview: tk.Frame):
        """Define cor selecionada"""
        color_var.set(color)
        preview.config(bg=color)
    
    def _choose_custom_color(self, color_var: tk.StringVar, preview: tk.Frame):
        """Abre seletor de cor personalizada"""
        color = colorchooser.askcolor(
            title="Escolha uma cor",
            initialcolor=color_var.get(),
            parent=self.editor
        )
        if color[1]:
            color_var.set(color[1])
            preview.config(bg=color[1])
    
    def _criar_botoes_acao(self, move_timestamp_var: tk.BooleanVar):
        """Cria botões de ação"""
        action_frame = tk.Frame(self.editor)
        action_frame.pack(side=tk.BOTTOM, pady=10)
        
        tk.Button(
            action_frame,
            text="💾 Salvar",
            command=self.salvar_edicao,
            width=15
        ).pack(side=tk.LEFT, padx=5)
        
        self.move_btn = tk.Button(
            action_frame,
            text="📅 Mover Data/Hora",
            command=lambda: self._toggle_move_timestamp(move_timestamp_var),
            relief=tk.RAISED,
            cursor="hand2",
            width=18
        )
        self.move_btn.pack(side=tk.LEFT, padx=5)
        
        tk.Button(
            action_frame,
            text="❌ Cancelar",
            command=self.cancelar_edicao,
            width=15
        ).pack(side=tk.LEFT, padx=5)
    
    def _toggle_move_timestamp(self, move_var: tk.BooleanVar):
        """Alterna modo de mover timestamp"""
        current = move_var.get()
        move_var.set(not current)
        
        if move_var.get():
            self.move_btn.config(
                relief=tk.SUNKEN,
                bg="#4CAF50",
                fg="white",
                text="📅 MODO MOVER ATIVO"
            )
            self.canvas.config(cursor="hand2")
        else:
            self.move_btn.config(
                relief=tk.RAISED,
                bg="SystemButtonFace",
                fg="black",
                text="📅 Mover Data/Hora"
            )
            self.canvas.config(cursor="cross")
        
        self.refresh_display()
    
    def _configurar_eventos(self, tool_var: tk.StringVar, move_var: tk.BooleanVar):
        """Configura eventos do canvas"""
        start_xy = {"x": None, "y": None}
        
        def on_button_press(event):
            if move_var.get():
                self._start_move_timestamp(event)
            else:
                start_xy["x"], start_xy["y"] = event.x, event.y
        
        def on_motion(event):
            if self.moving_timestamp and move_var.get():
                self._on_motion_timestamp(event)
            elif start_xy["x"] is not None and not move_var.get():
                self._handle_draw_motion(event, start_xy, tool_var)
        
        def on_button_release(event):
            if self.moving_timestamp:
                self._stop_move_timestamp()
            elif start_xy["x"] is not None and not move_var.get():
                self._handle_draw_release(event, start_xy, tool_var)
            
            start_xy["x"], start_xy["y"] = None, None
        
        def on_key_press(event):
            if event.keysym.lower() == 'z' and (event.state & 0x4):
                self.undo_action()
        
        self.canvas.bind("<Button-1>", on_button_press)
        self.canvas.bind("<B1-Motion>", on_motion)
        self.canvas.bind("<ButtonRelease-1>", on_button_release)
        self.editor.bind('<Control-z>', on_key_press)
        self.editor.bind('<Control-Z>', on_key_press)
    
    def _start_move_timestamp(self, event):
        """Inicia movimento do timestamp"""
        items = self.canvas.find_overlapping(
            event.x - 10, event.y - 10,
            event.x + 10, event.y + 10
        )
        
        for item in items:
            tags = self.canvas.gettags(item)
            if "timestamp" in tags or "timestamp_bg" in tags:
                self.moving_timestamp = True
                self.timestamp_drag_data["x"] = event.x
                self.timestamp_drag_data["y"] = event.y
                self.last_mouse_pos = (event.x, event.y)
                self.canvas.config(cursor="fleur")
                break
    
    def _on_motion_timestamp(self, event):
        """Processa movimento do timestamp"""
        current_pos = (event.x, event.y)
        
        if abs(current_pos[0] - self.last_mouse_pos[0]) > 2 or \
           abs(current_pos[1] - self.last_mouse_pos[1]) > 2:
            
            dx = event.x - self.timestamp_drag_data["x"]
            dy = event.y - self.timestamp_drag_data["y"]
            
            img_width, img_height = self.original_img.size
            new_x = self.timestamp_pos[0] + (dx / self.scale_factor / img_width)
            new_y = self.timestamp_pos[1] + (dy / self.scale_factor / img_height)
            
            new_x = max(0.02, min(0.98, new_x))
            new_y = max(0.02, min(0.98, new_y))
            
            self.timestamp_pos = (new_x, new_y)
            self.refresh_display()
            
            self.timestamp_drag_data["x"] = event.x
            self.timestamp_drag_data["y"] = event.y
            self.last_mouse_pos = current_pos
    
    def _stop_move_timestamp(self):
        """Para movimento do timestamp"""
        self.moving_timestamp = False
        self.canvas.config(cursor="cross")
    
    def _handle_draw_motion(self, event, start_xy: dict, tool_var: tk.StringVar):
        """Processa movimento durante desenho"""
        sx, sy = start_xy["x"], start_xy["y"]
        ex, ey = event.x, event.y
        
        ix1 = int(sx / self.scale_factor)
        iy1 = int(sy / self.scale_factor)
        ix2 = int(ex / self.scale_factor)
        iy2 = int(ey / self.scale_factor)
        
        tool = tool_var.get()
        
        if tool == "circle":
            radius = int(((ix2 - ix1)**2 + (iy2 - iy1)**2)**0.5)
            coords = [ix1 - radius, iy1 - radius, ix1 + radius, iy1 + radius]
            self.temp_element = DrawElement("circle", coords, "", 0)
        
        elif tool == "rectangle":
            coords = [min(ix1, ix2), min(iy1, iy2), max(ix1, ix2), max(iy1, iy2)]
            self.temp_element = DrawElement("rectangle", coords, "", 0)
        
        elif tool == "arrow":
            self.temp_element = DrawElement("arrow", [ix1, iy1, ix2, iy2], "", 0)
        
        self.refresh_display()
    
    def _handle_draw_release(self, event, start_xy: dict, tool_var: tk.StringVar):
        """Processa liberação do mouse após desenho"""
        # Obter variáveis de cor e largura do contexto
        color_var = tk.StringVar(value=Colors.RED)
        width_var = tk.IntVar(value=3)
        
        sx, sy = start_xy["x"], start_xy["y"]
        ex, ey = event.x, event.y
        
        ix1 = int(sx / self.scale_factor)
        iy1 = int(sy / self.scale_factor)
        ix2 = int(ex / self.scale_factor)
        iy2 = int(ey / self.scale_factor)
        
        tool = tool_var.get()
        color = color_var.get()
        width = width_var.get()
        
        self.undo_stack.clear()
        
        if tool == "circle":
            radius = int(((ix2 - ix1)**2 + (iy2 - iy1)**2)**0.5)
            coords = [ix1 - radius, iy1 - radius, ix1 + radius, iy1 + radius]
            self.elements.append(DrawElement("circle", coords, color, width))
        
        elif tool == "rectangle":
            coords = [min(ix1, ix2), min(iy1, iy2), max(ix1, ix2), max(iy1, iy2)]
            self.elements.append(DrawElement("rectangle", coords, color, width))
        
        elif tool == "arrow":
            self.elements.append(DrawElement("arrow", [ix1, iy1, ix2, iy2], color, width))
        
        elif tool == "text":
            text = simpledialog.askstring("Texto", "Digite o texto:", parent=self.editor)
            if text:
                self.elements.append(DrawElement("text", [ix1, iy1], color, width, text))
        
        self.temp_element = None
        self.refresh_display()
    
    def refresh_display(self):
        """Atualiza a exibição do canvas"""
        self.canvas.delete("all")
        
        # Desenhar imagem base
        self.display_img = self.editing_img.resize(
            (self.disp_w, self.disp_h),
            Image.LANCZOS
        )
        self.current_tk_img = ImageTk.PhotoImage(self.display_img)
        self.canvas.create_image(0, 0, anchor="nw", image=self.current_tk_img)
        
        # Desenhar timestamp se modo ocultar
        if self.gravador.modo_captura == CaptureMode.OCULTAR:
            self._desenhar_timestamp_preview()
        
        # Desenhar elementos
        for element in self.elements:
            self._desenhar_elemento(element)
        
        # Desenhar elemento temporário
        if self.temp_element:
            self._desenhar_elemento(self.temp_element)
    
    def _desenhar_timestamp_preview(self):
        """Desenha preview do timestamp no canvas"""
        if not self.timestamp_pos:
            return
        
        img_width, img_height = self.original_img.size
        pos_x = int(self.timestamp_pos[0] * img_width * self.scale_factor)
        pos_y = int(self.timestamp_pos[1] * img_height * self.scale_factor)
        
        texto = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        
        # Calcular dimensões do texto
        font = tkfont.Font(family="Arial", size=12, weight="bold")
        text_width = font.measure(texto)
        text_height = font.metrics("linespace")
        
        padding = 8
        
        # Desenhar fundo
        self.canvas.create_rectangle(
            pos_x - padding,
            pos_y - padding,
            pos_x + text_width + padding,
            pos_y + text_height + padding,
            fill="#000000",
            outline="",
            tags="timestamp_bg"
        )
        
        # Desenhar texto
        self.canvas.create_text(
            pos_x, pos_y,
            text=texto,
            fill=Colors.WHITE,
            font=("Arial", 12, "bold"),
            anchor="nw",
            tags="timestamp"
        )
    
    def _desenhar_elemento(self, element: DrawElement):
        """Desenha um elemento no canvas"""
        scaled_coords = [int(c * self.scale_factor) for c in element.coords]
        
        if element.tipo == "circle":
            x1, y1, x2, y2 = scaled_coords
            self.canvas.create_oval(
                x1, y1, x2, y2,
                outline=element.cor,
                width=element.largura
            )
        
        elif element.tipo == "rectangle":
            x1, y1, x2, y2 = scaled_coords
            self.canvas.create_rectangle(
                x1, y1, x2, y2,
                outline=element.cor,
                width=element.largura
            )
        
        elif element.tipo == "arrow":
            x1, y1, x2, y2 = scaled_coords
            self._desenhar_seta(x1, y1, x2, y2, element.cor, element.largura)
        
        elif element.tipo == "text":
            x, y = scaled_coords
            self.canvas.create_text(
                x, y,
                text=element.texto,
                fill=element.cor,
                font=("Arial", 12),
                anchor="nw"
            )
    
    def _desenhar_seta(self, x1: int, y1: int, x2: int, y2: int, cor: str, largura: int):
        """Desenha uma seta no canvas"""
        self.canvas.create_line(x1, y1, x2, y2, fill=cor, width=largura)
        
        angle = math.atan2(y2 - y1, x2 - x1)
        arrow_size = Dimensions.ARROW_SIZE
        
        x3 = x2 - arrow_size * math.cos(angle - math.pi / 6)
        y3 = y2 - arrow_size * math.sin(angle - math.pi / 6)
        x4 = x2 - arrow_size * math.cos(angle + math.pi / 6)
        y4 = y2 - arrow_size * math.sin(angle + math.pi / 6)
        
        self.canvas.create_polygon(x2, y2, x3, y3, x4, y4, fill=cor, outline=cor)
    
    def undo_action(self):
        """Desfaz última ação"""
        if self.elements:
            removed = self.elements.pop()
            self.undo_stack.append(removed)
            self.refresh_display()
    
    def salvar_edicao(self):
        """Salva a edição"""
        try:
            final_img = Image.open(self.caminho_print).convert("RGBA")
            draw = ImageDraw.Draw(final_img)
            
            # Aplicar elementos de desenho
            for element in self.elements:
                self._aplicar_elemento_em_imagem(draw, element)
            
            final_img.save(self.caminho_print)
            
            # Atualizar posição do timestamp nos metadados
            if self.evidencia and self.evidencia.timestamp_config:
                self.evidencia.timestamp_config.posicao = self.timestamp_pos
                self.gravador.metadata_manager.atualizar_evidencia(
                    self.nome_arquivo,
                    timestamp_config=self.evidencia.timestamp_config
                )
            
            messagebox.showinfo(
                "Sucesso",
                "Edição salva! A data/hora será aplicada na geração do documento."
            )
            self.editor.destroy()
            
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao salvar: {str(e)}")
    
    def _aplicar_elemento_em_imagem(self, draw: ImageDraw.ImageDraw, element: DrawElement):
        """Aplica elemento de desenho na imagem final"""
        if element.tipo == "circle":
            x1, y1, x2, y2 = element.coords
            draw.ellipse([x1, y1, x2, y2], outline=element.cor, width=element.largura)
        
        elif element.tipo == "rectangle":
            x1, y1, x2, y2 = element.coords
            draw.rectangle([x1, y1, x2, y2], outline=element.cor, width=element.largura)
        
        elif element.tipo == "arrow":
            x1, y1, x2, y2 = element.coords
            draw.line([x1, y1, x2, y2], fill=element.cor, width=element.largura)
            
            angle = math.atan2(y2 - y1, x2 - x1)
            arrow_size = Dimensions.ARROW_SIZE
            
            x3 = x2 - arrow_size * math.cos(angle - math.pi / 6)
            y3 = y2 - arrow_size * math.sin(angle - math.pi / 6)
            x4 = x2 - arrow_size * math.cos(angle + math.pi / 6)
            y4 = y2 - arrow_size * math.sin(angle + math.pi / 6)
            
            draw.polygon([x2, y2, x3, y3, x4, y4], fill=element.cor, outline=element.cor)
        
        elif element.tipo == "text":
            x, y = element.coords
            try:
                font = ImageFont.truetype("arial.ttf", 20)
            except:
                font = ImageFont.load_default()
            draw.text((x, y), element.texto, fill=element.cor, font=font)
    
    def cancelar_edicao(self):
        """Cancela a edição"""
        if messagebox.askyesno("Confirmar", "Descartar todas as alterações?"):
            self.editor.destroy()


# ==================== INTERFACE PRINCIPAL ====================
def minimizar_janela():
    """Minimiza a janela principal"""
    try:
        root.iconify()
    except:
        pass


def criar_interface_principal(gravador: GravadorDocx):
    """Cria a interface principal"""
    root.title("PrintF - Capturar Evidências")
    root.geometry("500x400")
    
    tk.Label(
        root,
        text="📷 PrintF - Capturar Evidências",
        font=("Arial", 16, "bold")
    ).pack(pady=10)
    
    tk.Button(
        root,
        text="▶ Iniciar Gravação (F8)",
        command=lambda: root.after(0, gravador.mostrar_janela_configuracao),
        width=30
    ).pack(pady=5)
    
    tk.Button(
        root,
        text="⏸ Pausar Gravação (F6)",
        command=lambda: root.after(0, gravador.pausar),
        width=30
    ).pack(pady=5)
    
    tk.Button(
        root,
        text="▶ Retomar Gravação (F7)",
        command=lambda: root.after(0, gravador.retomar),
        width=30
    ).pack(pady=5)
    
    tk.Button(
        root,
        text="⏹ Finalizar Gravação (F9)",
        command=lambda: root.after(0, gravador.finalizar),
        width=30
    ).pack(pady=5)
    
    tk.Button(
        root,
        text="❌ Fechar Aplicativo (F12)",
        command=root.quit,
        width=30
    ).pack(pady=8)


def configurar_atalhos_globais(gravador: GravadorDocx):
    """Configura atalhos de teclado globais"""
    def on_press(key):
        try:
            if key == keyboard.Key.f6:
                root.after(0, gravador.pausar)
            elif key == keyboard.Key.f7:
                root.after(0, gravador.retomar)
            elif key == keyboard.Key.f8:
                root.after(0, gravador.mostrar_janela_configuracao)
            elif key == keyboard.Key.f9:
                root.after(0, gravador.finalizar)
            elif key == keyboard.Key.f12:
                root.after(0, root.quit)
        except Exception:
            pass
    
    listener = keyboard.Listener(on_press=on_press)
    listener.start()


# ==================== MAIN ====================
if __name__ == "__main__":
    root = tk.Tk()
    gravador = GravadorDocx()
    
    criar_interface_principal(gravador)
    configurar_atalhos_globais(gravador)
    
    root.protocol("WM_DELETE_WINDOW", root.quit)
    root.mainloop()
