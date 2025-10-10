import csv
import json
import os
import re
import threading
import time
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox, scrolledtext, ttk
from typing import Dict, List, Optional, Tuple
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


class ConfigManager:
    """Gerencia o carregamento e salvamento da configuração de campos"""
    
    DEFAULT_CONFIG = [
        {"label": "Campo1:", "key": "campo1"},
        {"label": "Campo2:", "key": "campo2"},
        {"label": "Campo3:", "key": "campo3"},
        {"label": "Campo4:", "key": "campo4"},
        {"label": "Campo5:", "key": "campo5"},
        {"label": "Campo6:", "key": "campo6"}
    ]

    def __init__(self, config_file: str = 'config_campos.json'):
        self.config_file = Path(config_file)

    def load_config(self) -> List[Dict]:
        """Carrega a configuração do arquivo JSON ou cria uma padrão"""
        try:
            if self.config_file.exists():
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    print(f"✅ Configuração carregada de '{self.config_file}'")
                    return config
            else:
                return self._create_default_config()
        except Exception as e:
            print(f"⚠️ Erro ao carregar configuração: {e}")
            return self.DEFAULT_CONFIG

    def _create_default_config(self) -> List[Dict]:
        """Cria arquivo de configuração padrão"""
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.DEFAULT_CONFIG, f, indent=4, ensure_ascii=False)
            print(f"ℹ️ Arquivo '{self.config_file}' criado com configuração padrão")
            return self.DEFAULT_CONFIG
        except Exception as e:
            print(f"❌ Erro ao criar configuração padrão: {e}")
            return self.DEFAULT_CONFIG


class CSVReader:
    """Responsável pela leitura de arquivos CSV"""
    
    ENCODINGS = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252', 'windows-1252']

    @staticmethod
    def read_csv(file_path: str) -> Optional[List[str]]:
        """Lê um arquivo CSV e retorna a lista de nomes"""
        try:
            return CSVReader._read_with_pandas(file_path) or CSVReader._read_manual(file_path)
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao ler o CSV: {e}")
            return None

    @staticmethod
    def _read_with_pandas(file_path: str) -> Optional[List[str]]:
        """Tenta ler o CSV usando pandas"""
        for encoding in CSVReader.ENCODINGS:
            try:
                df = pd.read_csv(file_path, encoding=encoding, engine='python', 
                               on_bad_lines='skip')
                if 'Nome' in df.columns:
                    nomes = df['Nome'].dropna().str.strip()
                    return nomes[nomes != ''].tolist()
            except Exception:
                continue
        return None

    @staticmethod
    def _read_manual(file_path: str) -> Optional[List[str]]:
        """Leitura manual do CSV como fallback"""
        for encoding in CSVReader.ENCODINGS:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return CSVReader._parse_csv_lines(file.readlines())
            except Exception:
                continue
        return None

    @staticmethod
    def _parse_csv_lines(lines: List[str]) -> Optional[List[str]]:
        """Parseia as linhas do CSV manualmente"""
        if not lines:
            return None

        header_line = next((line for line in lines if 'Nome' in line), None)
        if not header_line:
            return None

        headers = header_line.strip().split(',')
        if 'Nome' not in headers:
            return None

        nome_index = headers.index('Nome')
        nomes = []

        for line in lines[1:]:
            try:
                if '"' in line:
                    reader = csv.reader([line])
                    parts = next(reader)
                else:
                    parts = line.strip().split(',')

                if len(parts) > nome_index:
                    nome = parts[nome_index].strip().strip('"')
                    if nome:
                        nomes.append(nome)
            except Exception:
                continue

        return nomes if nomes else None


class DocumentProcessor:
    """Processa e gera documentos Word baseados em templates"""
    
    @staticmethod
    def clean_filename(filename: str, max_length: int = 100) -> str:
        """Limpa o nome do arquivo removendo caracteres inválidos"""
        cleaned = re.sub(r'[<>:"/\\|?*]', '_', filename)
        cleaned = cleaned.strip()[:max_length]
        return cleaned or "caso_teste"

    @staticmethod
    def adjust_template_fields(doc: Document, field_mapping: Dict[str, str]) -> None:
        """Ajusta os campos do template conforme o mapeamento fornecido"""
        for paragraph in doc.paragraphs:
            DocumentProcessor._adjust_paragraph_fields(paragraph, field_mapping)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        DocumentProcessor._adjust_paragraph_fields(paragraph, field_mapping)

    @staticmethod
    def _adjust_paragraph_fields(paragraph, field_mapping: Dict[str, str]) -> None:
        """Ajusta os campos em um parágrafo específico"""
        texto_original = paragraph.text.strip()
        if ':' in texto_original:
            field_key = texto_original.split(':', 1)[0].strip()
            if field_key in field_mapping:
                paragraph.text = f"{field_mapping[field_key]}: "

    @staticmethod
    def fill_template(doc: Document, data: Dict[str, str], field_mapping: Dict[str, str]) -> None:
        """Preenche o template com os dados fornecidos"""
        # Primeiro ajusta os campos do template de acordo com as labels do JSON
        DocumentProcessor.adjust_template_fields(doc, field_mapping)
        
        # Cria mapeamento label -> valor
        label_to_value = {}
        for original_key, label in field_mapping.items():
            label_to_value[label] = data.get(original_key, '')
        
        label_to_value['Caso de Teste'] = data.get('Caso de Teste', '')
        
        # Preenche parágrafos
        for paragraph in doc.paragraphs:
            DocumentProcessor._fill_paragraph(paragraph, label_to_value)
        
        # Preenche tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        DocumentProcessor._fill_paragraph(paragraph, label_to_value)

    @staticmethod
    def _fill_paragraph(paragraph, label_to_value: Dict[str, str]) -> None:
        """Preenche um parágrafo específico com os dados"""
        texto = paragraph.text.strip()
        if ':' in texto:
            field_name = texto.split(':', 1)[0].strip()
            if field_name in label_to_value:
                paragraph.text = f"{field_name}: {label_to_value[field_name]}"


class DefaultDocumentGenerator:
    """Gera documentos padrão quando nenhum template é fornecido"""
    
    @staticmethod
    def create_default_document(data: Dict[str, str], field_config: List[Dict]) -> Document:
        """Cria um documento padrão com estrutura organizada"""
        doc = Document()
        
        # Configurar estilos de fonte padrão
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        # Título do documento - Arial 20, negrito
        title = doc.add_heading('Evidências de Teste - Documentação', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        # Adicionar data e hora
        current_time = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        date_para = doc.add_paragraph(f"Gerado em: {current_time}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.style = doc.styles['Normal']
        doc.add_paragraph()
        
        # Seção de informações do teste
        info_heading = doc.add_heading('Informações do Teste', level=1)
        for run in info_heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        # Tabela para dados organizados
        table = doc.add_table(rows=len(field_config) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Cabeçalho da tabela
        header_cells = table.rows[0].cells
        header_cells[0].text = "Campo"
        header_cells[1].text = "Valor"
        
        # Formatar cabeçalho
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles['Normal']
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        # Preencher dados da configuração
        for i, campo_info in enumerate(field_config, 1):
            key = campo_info['key']
            label = campo_info['label'].rstrip(':')
            
            row_cells = table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = data.get(key, 'Não informado')
            
            # Aplicar estilo Arial 12 sem negrito
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles['Normal']
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        run.bold = False
                        run.font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        doc.add_paragraph()
        
        # Seção do caso de teste
        caso_heading = doc.add_heading('Caso de Teste', level=1)
        for run in caso_heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
        caso_teste_para = doc.add_paragraph()
        caso_run = caso_teste_para.add_run('Nome do Caso de Teste: ')
        caso_run.bold = True
        caso_run.font.name = 'Arial'
        caso_run.font.size = Pt(12)
        caso_run.font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        nome_run = caso_teste_para.add_run(data.get('Caso de Teste', 'Não informado'))
        nome_run.font.name = 'Arial'
        nome_run.font.size = Pt(12)
        nome_run.bold = False
        nome_run.font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        # Seção de descrição
        desc_heading = doc.add_heading('Descrição do Teste', level=2)
        for run in desc_heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
        desc_para = doc.add_paragraph(
            "Esta seção deve conter a descrição detalhada do caso de teste executado, "
            "incluindo pré-condições, passos de execução e resultados esperados."
        )
        desc_para.style = doc.styles['Normal']
        
        # Seção de evidências
        evid_heading = doc.add_heading('Evidências Coletadas', level=2)
        for run in evid_heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
        evid_para = doc.add_paragraph("Registro das evidências coletadas durante a execução do teste:")
        evid_para.style = doc.styles['Normal']
        
        # Tabela para evidências
        evidencias_table = doc.add_table(rows=5, cols=3)
        evidencias_table.style = 'Light Grid Accent 1'
        
        # Cabeçalho da tabela de evidências
        evidencias_header = evidencias_table.rows[0].cells
        headers = ['Etapa', 'Evidência', 'Resultado']
        for col, header in enumerate(headers):
            evidencias_header[col].text = header
            for paragraph in evidencias_header[col].paragraphs:
                paragraph.style = doc.styles['Normal']
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        # Linhas para preenchimento
        etapas = [
            'Pré-condições',
            'Configuração Inicial', 
            'Execução do Teste',
            'Pós-condições',
            'Resultado Final'
        ]
        
        for row, etapa in enumerate(etapas, 1):
            row_cells = evidencias_table.rows[row].cells
            row_cells[0].text = etapa
            row_cells[1].text = "[Descreva a evidência coletada]"
            row_cells[2].text = "[Resultado obtido - OK/Erro]"
            
            # Aplicar estilo Arial 12 sem negrito
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles['Normal']
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        run.bold = False
                        run.font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        doc.add_paragraph()
        
        # Seção de observações
        obs_heading = doc.add_heading('Observações e Comentários', level=2)
        for run in obs_heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = False
            run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
        obs_para = doc.add_paragraph("Adicione observações relevantes sobre a execução do teste:")
        obs_para.style = doc.styles['Normal']
        
        # Área para observações
        obs_list_para = doc.add_paragraph()
        obs_title_run = obs_list_para.add_run("Observações Gerais:\n")
        obs_title_run.bold = True
        obs_title_run.font.name = 'Arial'
        obs_title_run.font.size = Pt(12)
        obs_title_run.font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        obs_items = [
            "• [Insira observações sobre problemas encontrados]\n",
            "• [Comentários sobre o comportamento do sistema]\n",
            "• [Sugestões de melhorias]\n",
            "• [Outras informações relevantes]"
        ]
        
        for item in obs_items:
            item_run = obs_list_para.add_run(item)
            item_run.font.name = 'Arial'
            item_run.font.size = Pt(12)
            item_run.bold = False
            item_run.font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        # Rodapé informativo
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run("Documento gerado automaticamente pelo PrintF - Gerador de Templates")
        footer_run.italic = True
        footer_run.font.name = 'Arial'
        footer_run.font.size = Pt(12)
        footer_run.font.color.rgb = RGBColor(0, 0, 0) # Preto
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return doc


class TemplateGenerator:
    """Gerador de templates de exemplo"""
    
    @staticmethod
    def create_example_template(field_config: List[Dict]) -> bool:
        """Cria um template de exemplo com base na configuração"""
        try:
            doc = Document()
            
            # Configurar estilo normal para Arial 12
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)
            font.color.rgb = RGBColor(0, 0, 0) # Preto
            
            # Título principal - Arial 20, negrito
            main_title = doc.add_heading('Template Evidências de Teste', level=1)
            for run in main_title.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(20)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
            # Adicionar instruções
            info_para = doc.add_paragraph()
            info_title_run = info_para.add_run("Instruções: ")
            info_title_run.bold = True
            info_title_run.font.name = 'Arial'
            info_title_run.font.size = Pt(12)
            info_title_run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
            info_text_run = info_para.add_run("Este é um template de exemplo. Os campos abaixo serão preenchidos automaticamente.")
            info_text_run.font.name = 'Arial'
            info_text_run.font.size = Pt(12)
            info_text_run.bold = False
            info_text_run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
            doc.add_paragraph()
            
            # Adicionar campos da configuração
            for campo_info in field_config:
                campo_para = doc.add_paragraph()
                label_run = campo_para.add_run(f"{campo_info['label']} ")
                label_run.font.name = 'Arial'
                label_run.font.size = Pt(12)
                label_run.bold = False
                label_run.font.color.rgb = RGBColor(0, 0, 0) # Preto
                
                value_run = campo_para.add_run("[VALOR]")
                value_run.font.name = 'Arial'
                value_run.font.size = Pt(12)
                value_run.bold = False
                value_run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
            doc.add_paragraph()
            
            # Seção para caso de teste - Arial 20, negrito
            caso_title = doc.add_heading('Detalhes do Caso de Teste', level=2)
            for run in caso_title.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(20)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
            caso_para = doc.add_paragraph()
            caso_label_run = caso_para.add_run("Caso de Teste: ")
            caso_label_run.font.name = 'Arial'
            caso_label_run.font.size = Pt(12)
            caso_label_run.bold = False
            caso_label_run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
            caso_value_run = caso_para.add_run("[NOME_DO_CASO]")
            caso_value_run.font.name = 'Arial'
            caso_value_run.font.size = Pt(12)
            caso_value_run.bold = False
            caso_value_run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
            # Tabela para evidências
            table = doc.add_table(rows=3, cols=2)
            table.style = 'Table Grid'
            
            # Configurar células da tabela
            table.cell(0, 0).text = 'Resultado Esperado:'
            table.cell(0, 1).text = '[RESULTADO_ESPERADO]'
            table.cell(1, 0).text = 'Resultado Obtido:'
            table.cell(1, 1).text = '[RESULTADO_OBTIDO]'
            table.cell(2, 0).text = 'Observações:'
            table.cell(2, 1).text = '[OBSERVACOES]'
            
            # Aplicar estilo Arial 12 sem negrito para toda a tabela
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style = doc.styles['Normal']
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            run.bold = False
                            run.font.color.rgb = RGBColor(0, 0, 0) # Preto
            
            doc.save('template_evidencias.docx')
            return True
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao criar template: {e}")
            return False


# O restante do código permanece igual...
class GeradorTemplates:
    """Interface principal da aplicação"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("PrintF - Gerar Templates")
        self.root.geometry("800x700")
        self.root.configure(bg='#f0f0f0')
        
        self.config_manager = ConfigManager()
        self.csv_reader = CSVReader()
        self.doc_processor = DocumentProcessor()
        self.default_doc_generator = DefaultDocumentGenerator()
        
        self.campos_config = self.config_manager.load_config()
        self.campos_entries: Dict[str, tk.Entry] = {}
        
        self._setup_ui()
        
    def _setup_ui(self) -> None:
        """Configura a interface do usuário"""
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        self._configure_grid_weights()
        self._create_title_section(main_frame)
        self._create_dynamic_fields_section(main_frame)
        self._create_file_section(main_frame)
        self._create_control_buttons(main_frame)
        self._create_progress_section(main_frame)
        self._create_log_section(main_frame)
        
        self._set_default_template()

    def _configure_grid_weights(self) -> None:
        """Configura os pesos do grid para redimensionamento"""
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)

    def _create_title_section(self, parent) -> None:
        """Cria a seção do título"""
        titulo = ttk.Label(parent, text="📄 PrintF - Gerar Templates", 
                          font=("Arial", 16, "bold"))
        titulo.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        ttk.Separator(parent, orient='horizontal').grid(
            row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))

    def _create_dynamic_fields_section(self, parent) -> None:
        """Cria os campos dinâmicos baseados na configuração"""
        ttk.Label(parent, text="Dados dos Testes:", 
                 font=("Arial", 12, "bold")).grid(
                     row=2, column=0, columnspan=3, sticky=tk.W, pady=(0, 10))

        for i, campo_info in enumerate(self.campos_config):
            self._create_field_row(parent, campo_info, i)

    def _create_field_row(self, parent, campo_info: Dict, row_index: int) -> None:
        """Cria uma linha de campo na interface"""
        label_text = campo_info['label']
        campo_key = campo_info['key']
        
        ttk.Label(parent, text=label_text).grid(
            row=3 + row_index, column=0, sticky=tk.W, pady=2)
        
        entry = ttk.Entry(parent, width=40)
        entry.grid(row=3 + row_index, column=1, columnspan=2, 
                  sticky=(tk.W, tk.E), pady=2, padx=(5, 0))
        
        self.campos_entries[campo_key] = entry

    def _create_file_section(self, parent) -> None:
        """Cria a seção de seleção de arquivos"""
        next_row = 3 + len(self.campos_config)
        
        ttk.Separator(parent, orient='horizontal').grid(
            row=next_row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)

        ttk.Label(parent, text="Arquivos:", 
                 font=("Arial", 12, "bold")).grid(
                     row=next_row + 1, column=0, columnspan=3, sticky=tk.W, pady=(0, 10))

        # Campos de arquivo
        self.csv_entry = self._create_file_field(parent, "CSV:*", next_row + 2, self.selecionar_csv)
        self.template_entry = self._create_file_field(parent, "Template (Opcional):", next_row + 3, self.selecionar_template)
        self.pasta_entry = self._create_file_field(parent, "Pasta Saída:", next_row + 4, self.selecionar_pasta)

        # Info sobre campos obrigatórios
        info_label = ttk.Label(parent, text="* Campos obrigatórios", font=("Arial", 9), foreground="gray")
        info_label.grid(row=next_row + 5, column=0, columnspan=3, sticky=tk.W, pady=(5, 0))

    def _create_file_field(self, parent, label: str, row: int, command) -> ttk.Entry:
        """Cria um campo de seleção de arquivo"""
        ttk.Label(parent, text=label).grid(row=row, column=0, sticky=tk.W, pady=2)
        entry = ttk.Entry(parent, width=40)
        entry.grid(row=row, column=1, sticky=(tk.W, tk.E), pady=2, padx=(5, 0))
        ttk.Button(parent, text="Procurar", command=command).grid(
            row=row, column=2, padx=(5, 0))
        return entry

    def _create_control_buttons(self, parent) -> None:
        """Cria os botões de controle - REMOVIDO O BOTÃO DE TEMPLATE EXEMPLO"""
        next_row = 3 + len(self.campos_config) + 6
        
        ttk.Separator(parent, orient='horizontal').grid(
            row=next_row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)

        button_frame = ttk.Frame(parent)
        button_frame.grid(row=next_row + 1, column=0, columnspan=3, pady=10)
        
        self.gerar_btn = ttk.Button(button_frame, text="▶️ Gerar Documentos", 
                                   command=self.iniciar_processamento, style='Accent.TButton')
        self.gerar_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(button_frame, text="🔄 Limpar", command=self.limpar_campos).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="❌ Sair", command=self.root.quit).pack(side=tk.LEFT, padx=5)

    def _create_progress_section(self, parent) -> None:
        """Cria a seção de progresso"""
        next_row = 3 + len(self.campos_config) + 8
        
        ttk.Separator(parent, orient='horizontal').grid(
            row=next_row, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)

        ttk.Label(parent, text="Progresso:").grid(
            row=next_row + 1, column=0, columnspan=3, sticky=tk.W, pady=(0, 5))
        
        self.progress = ttk.Progressbar(parent, mode='determinate')
        self.progress.grid(row=next_row + 2, column=0, columnspan=3, 
                          sticky=(tk.W, tk.E), pady=(0, 10))

    def _create_log_section(self, parent) -> None:
        """Cria a seção de log"""
        next_row = 3 + len(self.campos_config) + 10
        
        ttk.Label(parent, text="Log de Execução:").grid(
            row=next_row, column=0, columnspan=3, sticky=tk.W, pady=(0, 5))
        
        self.log_text = scrolledtext.ScrolledText(parent, width=70, height=15, state='disabled')
        self.log_text.grid(row=next_row + 1, column=0, columnspan=3, 
                          sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        
        parent.rowconfigure(next_row + 1, weight=1)

    def _set_default_template(self) -> None:
        """Preenche template padrão se existir - AGORA CRIA AUTOMATICAMENTE SE NÃO EXISTIR"""
        template_path = 'template_evidencias.docx'
        if not Path(template_path).exists():
            # Cria template automaticamente se não existir
            self.log("📝 Criando template padrão automaticamente...")
            if self._criar_template_exemplo_automatico():
                self.template_entry.insert(0, template_path)
                self.log("✅ Template padrão criado com sucesso!")
        else:
            self.template_entry.insert(0, template_path)

    def _criar_template_exemplo_automatico(self) -> bool:
        """Cria template de exemplo automaticamente (sem interação do usuário)"""
        try:
            return TemplateGenerator.create_example_template(self.campos_config)
        except Exception as e:
            self.log(f"⚠️ Aviso: Não foi possível criar template automático: {e}")
            return False

    # Métodos de seleção de arquivos
    def selecionar_csv(self) -> None:
        self._select_file(self.csv_entry, "Selecionar arquivo CSV", 
                         [("CSV Files", "*.csv"), ("Todos os arquivos", "*.*")])

    def selecionar_template(self) -> None:
        self._select_file(self.template_entry, "Selecionar template DOCX", 
                         [("Word Documents", "*.docx"), ("Todos os arquivos", "*.*")])

    def selecionar_pasta(self) -> None:
        pasta = filedialog.askdirectory(title="Selecionar pasta de saída")
        if pasta:
            self.pasta_entry.delete(0, tk.END)
            self.pasta_entry.insert(0, pasta)

    def _select_file(self, entry_widget: ttk.Entry, title: str, filetypes: List[Tuple]) -> None:
        """Seleciona um arquivo e atualiza o campo de entrada"""
        arquivo = filedialog.askopenfilename(title=title, filetypes=filetypes)
        if arquivo:
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, arquivo)

    def limpar_campos(self) -> None:
        """Limpa todos os campos da interface"""
        for entry in [self.csv_entry, self.template_entry, self.pasta_entry]:
            entry.delete(0, tk.END)
        
        for entry in self.campos_entries.values():
            entry.delete(0, tk.END)
        
        self._clear_log()
        self.progress['value'] = 0

    def _clear_log(self) -> None:
        """Limpa o log de execução"""
        self.log_text.config(state='normal')
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state='disabled')

    def log(self, mensagem: str) -> None:
        """Adiciona mensagem ao log"""
        self.log_text.config(state='normal')
        self.log_text.insert(tk.END, mensagem + "\n")
        self.log_text.see(tk.END)
        self.log_text.config(state='disabled')
        self.root.update()

    def _get_fixed_data(self) -> Dict[str, str]:
        """Obtém os dados dos campos fixos"""
        dados = {}
        for campo_key, entry in self.campos_entries.items():
            valor = entry.get().strip()
            if not valor:
                raise ValueError(f"O campo '{campo_key}' é obrigatório!")
            dados[campo_key] = valor
        return dados

    def _validate_inputs(self, csv_path: str) -> bool:
        """Valida os inputs necessários (apenas CSV é obrigatório)"""
        if not csv_path:
            messagebox.showerror("Erro", "Selecione um arquivo CSV!")
            return False
        
        if not Path(csv_path).exists():
            messagebox.showerror("Erro", "Arquivo CSV não encontrado!")
            return False
        
        return True

    def processar_documentos(self) -> None:
        """Processa os documentos em lote - versão robusta que nunca falha"""
        try:
            # Obter dados fixos
            try:
                dados_fixos = self._get_fixed_data()
            except ValueError as e:
                messagebox.showerror("Erro", str(e))
                self.gerar_btn.config(state='normal')
                return
            
            csv_path = self.csv_entry.get()
            template_path = self.template_entry.get().strip()
            output_folder = self.pasta_entry.get().strip() or 'evidencias_geradas'
            
            if not self._validate_inputs(csv_path):
                self.gerar_btn.config(state='normal')
                return

            # Garantir que temos um template válido
            template_path = self._garantir_template_valido(template_path)
            
            # Criar pasta de saída
            try:
                Path(output_folder).mkdir(exist_ok=True)
            except Exception as e:
                self.log(f"⚠️ Aviso: Não foi possível criar a pasta '{output_folder}': {e}")
                self.log("📁 Usando pasta atual para salvar os documentos...")
                output_folder = '.'
            
            self.log("📖 Lendo arquivo CSV...")
            casos_teste = self.csv_reader.read_csv(csv_path)
            
            if not casos_teste:
                messagebox.showerror("Erro", "Não foi possível ler os casos de teste do CSV")
                self.gerar_btn.config(state='normal')
                return
            
            # Determinar modo de operação
            use_default_template = True
            if template_path and Path(template_path).exists():
                use_default_template = False
                self.log("📁 Usando template personalizado...")
            else:
                self.log("📝 Gerando documentos com template padrão...")
            
            self._process_test_cases(casos_teste, dados_fixos, template_path, output_folder, use_default_template)
            
        except Exception as e:
            self.log(f"❌ Erro inesperado: {e}")
            messagebox.showerror("Erro", f"Erro inesperado: {e}")
            self.gerar_btn.config(state='normal')

    def _garantir_template_valido(self, template_path: str) -> str:
        """Garante que temos um template válido, criando automaticamente se necessário"""
        if not template_path or not Path(template_path).exists():
            self.log("📝 Nenhum template válido encontrado, criando automaticamente...")
            if self._criar_template_exemplo_automatico():
                new_template_path = 'template_evidencias.docx'
                self.template_entry.delete(0, tk.END)
                self.template_entry.insert(0, new_template_path)
                self.log("✅ Template padrão criado e configurado automaticamente")
                return new_template_path
            else:
                self.log("⚠️ Não foi possível criar template, usando geração padrão...")
                return ""
        return template_path

    def _process_test_cases(self, casos_teste: List[str], dados_fixos: Dict[str, str], 
                           template_path: str, output_folder: str, use_default_template: bool) -> None:
        """Processa cada caso de teste individualmente"""
        self.log(f"📊 Encontrados {len(casos_teste)} casos de teste\n")
        self.progress['maximum'] = len(casos_teste)
        
        sucessos = 0
        erros = []
        arquivos_gerados = set()
        
        # Criar mapeamento de campos
        field_mapping = {campo['key']: campo['label'].rstrip(':').strip() 
                        for campo in self.campos_config}
        
        campo_nome = next(iter(dados_fixos.keys()))
        
        for i, caso_teste in enumerate(casos_teste, 1):
            try:
                self.progress['value'] = i
                self.log(f"🔄 Processando: {caso_teste}")
                
                if self._generate_single_document(caso_teste, dados_fixos, template_path, 
                                                output_folder, field_mapping, campo_nome, 
                                                arquivos_gerados, use_default_template):
                    sucessos += 1
                else:
                    erros.append((caso_teste, "Erro na geração"))
                    
            except Exception as e:
                self.log(f"❌ Erro no caso '{caso_teste}': {e}\n")
                erros.append((caso_teste, str(e)))
            
            time.sleep(0.05)  # Pequena pausa para não sobrecarregar
            
        self._show_final_results(sucessos, len(erros), len(casos_teste), 
                               output_folder, arquivos_gerados)
        self.gerar_btn.config(state='normal')

    def _generate_single_document(self, caso_teste: str, dados_fixos: Dict[str, str], 
                                template_path: str, output_folder: str, 
                                field_mapping: Dict[str, str], campo_nome: str,
                                arquivos_gerados: set, use_default_template: bool) -> bool:
        """Gera um único documento - versão robusta"""
        try:
            # Usar template se fornecido e existir, caso contrário criar documento padrão
            if not use_default_template:
                try:
                    doc = Document(template_path)
                    dados_completos = dados_fixos.copy()
                    dados_completos['Caso de Teste'] = caso_teste
                    self.doc_processor.fill_template(doc, dados_completos, field_mapping)
                except Exception as e:
                    self.log(f"⚠️ Erro ao usar template personalizado: {e}. Usando template padrão...")
                    doc = self.default_doc_generator.create_default_document(
                        dados_fixos.copy(), self.campos_config)
                    doc = self.default_doc_generator.create_default_document(
                        {**dados_fixos, 'Caso de Teste': caso_teste}, self.campos_config)
            else:
                # Criar documento padrão com todos os dados
                dados_completos = dados_fixos.copy()
                dados_completos['Caso de Teste'] = caso_teste
                doc = self.default_doc_generator.create_default_document(
                    dados_completos, self.campos_config)
            
            # Gerar nome do arquivo
            nome_base = self.doc_processor.clean_filename(caso_teste)
            nome_arquivo = self._generate_unique_filename(
                f"Evidencia_{dados_fixos[campo_nome]}_{nome_base}.docx", arquivos_gerados)
            
            caminho_completo = Path(output_folder) / nome_arquivo
            
            # Tentar salvar o documento
            try:
                doc.save(caminho_completo)
                self.log(f"✅ Salvo: {nome_arquivo}")
                return True
            except Exception as e:
                # Fallback: tentar salvar com nome diferente
                try:
                    nome_alternativo = f"Evidencia_{i}_{datetime.now().strftime('%H%M%S')}.docx"
                    caminho_alternativo = Path(output_folder) / nome_alternativo
                    doc.save(caminho_alternativo)
                    self.log(f"✅ Salvo (nome alternativo): {nome_alternativo}")
                    arquivos_gerados.add(nome_alternativo)
                    return True
                except Exception as e2:
                    self.log(f"❌ Erro ao salvar documento: {e2}")
                    return False
            
        except Exception as e:
            self.log(f"❌ Erro crítico ao gerar documento: {e}")
            return False

    def _generate_unique_filename(self, filename: str, existing_files: set) -> str:
        """Gera um nome de arquivo único"""
        contador = 1
        nome_original = filename
        
        while filename in existing_files:
            base, ext = os.path.splitext(nome_original)
            filename = f"{base}_{contador}{ext}"
            contador += 1
        
        existing_files.add(filename)
        return filename

    def _show_final_results(self, sucessos: int, erros: int, total: int, 
                           output_folder: str, arquivos_gerados: set) -> None:
        """Exibe os resultados finais do processamento"""
        self.log("\n" + "=" * 50)
        self.log("🎉 PROCESSO CONCLUÍDO!")
        self.log("=" * 50)
        self.log(f"📊 Total processado: {total}")
        self.log(f"✅ Sucessos: {sucessos}")
        self.log(f"❌ Erros: {erros}")
        self.log(f"📁 Pasta: {Path(output_folder).absolute()}")
        
        if sucessos > 0:
            self.log(f"\n📋 Arquivos gerados:")
            for arquivo in sorted(arquivos_gerados)[:10]:  # Mostra apenas os 10 primeiros
                self.log(f"• {arquivo}")
            if len(arquivos_gerados) > 10:
                self.log(f"• ... e mais {len(arquivos_gerados) - 10} arquivos")
        
        if erros == 0:
            messagebox.showinfo("Sucesso", 
                              f"✅ Todos os {sucessos} documentos foram gerados com sucesso!\n"
                              f"📁 Pasta: {Path(output_folder).absolute()}")
        else:
            messagebox.showwarning("Concluído com avisos", 
                                 f"Processo concluído com {erros} erro(s).\n"
                                 f"✅ {sucessos} documentos gerados com sucesso.\n"
                                 f"📁 Pasta: {Path(output_folder).absolute()}")

    def iniciar_processamento(self) -> None:
        """Inicia o processamento em thread separada"""
        self.gerar_btn.config(state='disabled')
        self._clear_log()
        self.progress['value'] = 0
        
        thread = threading.Thread(target=self.processar_documentos, daemon=True)
        thread.start()


def check_dependencies() -> bool:
    """Verifica se todas as dependências estão instaladas"""
    missing_deps = []
    
    try:
        import pandas as pd
    except ImportError:
        missing_deps.append("pandas")
    
    try:
        from docx import Document
    except ImportError:
        missing_deps.append("python-docx")
    
    if missing_deps:
        print("=" * 60)
        print("❌ DEPENDÊNCIAS NÃO INSTALADAS")
        print("=" * 60)
        print(f"Faltam as seguintes bibliotecas: {', '.join(missing_deps)}")
        print("\n💡 Para instalar, execute:")
        print(f"\n   pip install {' '.join(missing_deps)}")
        print("\nOu instale todas de uma vez:")
        print("\n   pip install pandas python-docx")
        print("\n" + "=" * 60)
        input("\nPressione Enter para sair...")
        return False
    
    return True


if __name__ == "__main__":
    if not check_dependencies():
        exit()
    
    root = tk.Tk()
    app = GeradorTemplates(root)
    root.mainloop()