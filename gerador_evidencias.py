import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, colorchooser, ttk
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pyautogui
from pynput import mouse, keyboard
from PIL import Image, ImageTk, ImageDraw, ImageFont
from datetime import datetime
import math
import re
import glob
import json  # Adicionado
import uuid   # Adicionado

# ------------------ Gerador Docx -----------------
class GeradorDocx:
    def __init__(self):
        self.output_dir = os.getcwd()
        self.prints = []            # lista de caminhos das imagens salvas
        self.doc = None
        self.using_template = False
        self.template_path = None
        self.current_index = 0  # Adicionado: controlar o índice atual
        self.evidence_dir = None  # Diretório das evidências
        self.metadata_path = None  # Adicionado
        self.metadata = {"evidencias": [], "proximo_id": 1}  # Adicionado

    def _salvar_metadata(self):
        """Salva os metadados no arquivo JSON"""
        if self.metadata_path:
            with open(self.metadata_path, 'w', encoding='utf-8') as f:
                json.dump(self.metadata, f, indent=2, ensure_ascii=False)

    def carregar_evidencias(self, dir_path):
        """Carrega as evidências baseadas nos metadados"""
        self.metadata_path = os.path.join(dir_path, "evidencias_metadata.json")
        
        if os.path.exists(self.metadata_path):
            try:
                with open(self.metadata_path, 'r', encoding='utf-8') as f:
                    self.metadata = json.load(f)
            except:
                self.metadata = {"evidencias": [], "proximo_id": 1}
        
        # Carrega evidências ativas (não excluídas)
        evidencias_ativas = []
        for evidencia in self.metadata["evidencias"]:
            if not evidencia.get("excluida", False):
                caminho = os.path.join(dir_path, evidencia["arquivo"])
                if os.path.exists(caminho):
                    evidencias_ativas.append(caminho)
        
        return evidencias_ativas

    def recarregar_evidencias(self):
        """Recarrega a lista de evidências"""
        if self.evidence_dir:
            self.prints = self.carregar_evidencias(self.evidence_dir)
            return True
        return False

    def obter_comentario(self, nome_arquivo):
        """Obtém o comentário salvo nos metadados"""
        for evidencia in self.metadata["evidencias"]:
            if evidencia["arquivo"] == nome_arquivo:
                return evidencia.get("comentario", "")
        return ""

    # ---------- Nova janela de configuração ----------
    def mostrar_janela_configuracao(self):
        config_window = tk.Toplevel(root)
        config_window.title("Configuração de Arquivo")
        config_window.geometry("600x500")
        config_window.resizable(False, False)
        
        config_window.transient(root)
        config_window.grab_set()
        
        main_frame = ttk.Frame(config_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(main_frame, text="PrintF - Configuração de Arquivo", 
                 font=("Arial", 16, "bold")).pack(pady=10)
        
        # Seleção de template
        ttk.Label(main_frame, text="Selecione o template DOCX:").pack(anchor="w", pady=(10, 5))
        
        template_frame = ttk.Frame(main_frame)
        template_frame.pack(fill=tk.X, pady=5)
        
        self.template_var = tk.StringVar()
        template_entry = ttk.Entry(template_frame, textvariable=self.template_var, width=40)
        template_entry.pack(side=tk.LEFT, padx=(0, 5), fill=tk.X, expand=True)
        
        def selecionar_template():
            template_path = filedialog.askopenfilename(
                title="Selecione o template DOCX",
                filetypes=[("Documentos Word", "*.docx")]
            )
            if template_path:
                self.template_var.set(template_path)
        
        ttk.Button(template_frame, text="Procurar", command=selecionar_template).pack(side=tk.RIGHT)
        
        # Seleção de diretório de evidências
        ttk.Label(main_frame, text="Selecione o diretório onde estão as evidências:").pack(anchor="w", pady=(10, 5))
        
        dir_frame = ttk.Frame(main_frame)
        dir_frame.pack(fill=tk.X, pady=5)
        
        self.dir_var = tk.StringVar()
        dir_entry = ttk.Entry(dir_frame, textvariable=self.dir_var, width=40)
        dir_entry.pack(side=tk.LEFT, padx=(0, 5), fill=tk.X, expand=True)
        
        def selecionar_diretorio():
            dir_path = filedialog.askdirectory(title="Selecione o diretório onde estão as evidências")
            if dir_path:
                self.dir_var.set(dir_path)
                atualizar_lista_arquivos(dir_path)
        
        ttk.Button(dir_frame, text="Procurar", command=selecionar_diretorio).pack(side=tk.RIGHT)
        
        # Frame para exibir a lista de arquivos
        file_list_frame = ttk.Frame(main_frame)
        file_list_frame.pack(fill=tk.BOTH, expand=True, pady=(5, 10))
        
        file_list_scrollbar = ttk.Scrollbar(file_list_frame)
        file_list_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.file_listbox = tk.Listbox(file_list_frame, yscrollcommand=file_list_scrollbar.set, height=8)
        self.file_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        file_list_scrollbar.config(command=self.file_listbox.yview)
        
        self.file_count_label = ttk.Label(main_frame, text="Nenhum arquivo PNG encontrado")
        self.file_count_label.pack(anchor="w", pady=(0, 10))
        
        def atualizar_lista_arquivos(dir_path):
            self.file_listbox.delete(0, tk.END)
            png_files = self.carregar_evidencias(dir_path)
            
            for file_path in png_files:
                filename = os.path.basename(file_path)
                # Mostra também o timestamp para referência
                timestamp = datetime.fromtimestamp(os.path.getmtime(file_path))
                self.file_listbox.insert(tk.END, f"{filename} ({timestamp.strftime('%H:%M:%S')})")
            
            if png_files:
                self.file_count_label.config(text=f"{len(png_files)} arquivo(s) PNG encontrado(s)")
            else:
                self.file_count_label.config(text="Nenhum arquivo PNG encontrado")
        
        # Botões
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=20)
        
        def iniciar_geracao():
            if not self.template_var.get() or not self.dir_var.get():
                messagebox.showerror("Erro", "Por favor, selecione o template e o diretório de evidências.")
                return
            
            if not os.path.exists(self.template_var.get()):
                messagebox.showerror("Erro", "O arquivo de template selecionado não existe.")
                return
            
            if not os.path.exists(self.dir_var.get()):
                messagebox.showerror("Erro", "O diretório de evidências selecionado não existe.")
                return
            
            png_files = self.carregar_evidencias(self.dir_var.get())
            if not png_files:
                messagebox.showerror("Erro", "Nenhuma evidência PNG encontrada no diretório selecionado.")
                return
            
            self.template_path = self.template_var.get()
            self.output_dir = self.dir_var.get()
            self.evidence_dir = self.dir_var.get()  # Salva o diretório de evidências
            self.prints = png_files
            self.current_index = 0  # Reiniciar índice
            
            config_window.destroy()            
            self.iniciar_processamento()
        
        ttk.Button(btn_frame, text="Gerar Documento", command=iniciar_geracao).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Cancelar", command=config_window.destroy).pack(side=tk.LEFT, padx=5)
        
        root.wait_window(config_window)
        return self.template_path is not None and self.output_dir is not None and self.prints

    def iniciar_processamento(self):
        os.makedirs(self.output_dir, exist_ok=True)

        try:
            if os.path.exists(self.template_path):
                self.doc = Document(self.template_path)
                self.using_template = True
            else:
                self.doc = Document()
                self.using_template = False
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar template: {str(e)}")
            self.doc = Document()
            self.using_template = False
        
        self.gerar_docx()

    def gerar_docx(self):
        # Processa as evidências usando índice em vez de loop for
        documento_salvo = False  # Flag para controlar se o documento já foi salvo
        
        while self.current_index < len(self.prints):
            caminho_print = self.prints[self.current_index]
            
            # Verifica se o arquivo ainda existe
            if not os.path.exists(caminho_print):
                # Recarrega a lista se o arquivo não existir mais
                if not self.recarregar_evidencias():
                    break
                if self.current_index >= len(self.prints):
                    break
                caminho_print = self.prints[self.current_index]
            
            resultado = self.mostrar_imagem(caminho_print)
            
            if resultado is False:  # Processamento cancelado
                break
            elif resultado is None:  # Exclusão ocorreu, não incrementa índice
                # Recarrega a lista após exclusão
                self.recarregar_evidencias()
                continue
            elif resultado == "ja_salvou":  # Já salvou via "Incluir Todos"
                documento_salvo = True
                break  # Sai do loop completamente
            else:  # Adicionou com sucesso, vai para próxima
                self.current_index += 1
        
        # Só salva se não salvou anteriormente (no incluir_todos)
        if not documento_salvo:
            self.salvar_docx()

    def mostrar_imagem(self, caminho_print):
        popup = tk.Toplevel(root)
        popup.title("Adicionar Comentário à Evidência")
        popup.geometry("950x750")
        popup.resizable(False, False)

        self.processamento_cancelado = False
        resultado = None

        # Verifica se o arquivo ainda existe
        if not os.path.exists(caminho_print):
            messagebox.showerror("Erro", f"Arquivo não encontrado: {os.path.basename(caminho_print)}")
            popup.destroy()
            return None

        img = Image.open(caminho_print)
        img.thumbnail((850, 550))
        img_tk = ImageTk.PhotoImage(img)
        label_img = tk.Label(popup, image=img_tk)
        label_img.image = img_tk
        label_img.pack(pady=10)

        tk.Label(popup, text="Comentário (opcional):").pack()
        entry = tk.Entry(popup, width=80)
        entry.pack(pady=5)

        # Mostra informações do arquivo
        file_info = f"Arquivo: {os.path.basename(caminho_print)}"
        timestamp = datetime.fromtimestamp(os.path.getmtime(caminho_print))
        file_info += f" - {timestamp.strftime('%H:%M:%S')}"
        tk.Label(popup, text=file_info, font=("Arial", 10)).pack()

        def editar_print():
            self.abrir_editor(caminho_print, popup)

        def adicionar():
            nonlocal resultado
            comentario = entry.get()
            
            # Atualiza comentário nos metadados
            nome_arquivo = os.path.basename(caminho_print)
            for evidencia in self.metadata["evidencias"]:
                if evidencia["arquivo"] == nome_arquivo:
                    evidencia["comentario"] = comentario
                    break
            self._salvar_metadata()
            
            self.doc.add_picture(caminho_print, width=Inches(5))
            if comentario.strip():
                self.doc.add_paragraph(comentario)
            resultado = True
            popup.destroy()

        def cancelar_processamento():
            if messagebox.askyesno("Confirmar Cancelamento", 
                                  "Tem certeza que deseja cancelar o processamento?"):
                self.processamento_cancelado = True
                popup.destroy()

        def incluir_todos():
            if messagebox.askyesno("Confirmar Inclusão", 
                                  "Deseja incluir todas as evidências restantes sem editar?\nAs evidências serão adicionadas sem comentários."):
                # Adicionar a evidência atual primeiro
                comentario = entry.get()
                
                # Atualiza comentário nos metadados
                nome_arquivo = os.path.basename(caminho_print)
                for evidencia in self.metadata["evidencias"]:
                    if evidencia["arquivo"] == nome_arquivo:
                        evidencia["comentario"] = comentario
                        break
                self._salvar_metadata()
                
                self.doc.add_picture(caminho_print, width=Inches(5))
                if comentario.strip():
                    self.doc.add_paragraph(comentario)
                
                # Adicionar todas as evidências restantes
                for i in range(self.current_index + 1, len(self.prints)):
                    print_path = self.prints[i]
                    if os.path.exists(print_path):  # Verifica se o arquivo existe
                        # Usa comentário salvo nos metadados
                        nome_arquivo_restante = os.path.basename(print_path)
                        comentario_restante = self.obter_comentario(nome_arquivo_restante)
                        
                        self.doc.add_picture(print_path, width=Inches(5))
                        if comentario_restante.strip():
                            self.doc.add_paragraph(comentario_restante)
                        else:
                            self.doc.add_paragraph("")
                
                # Atualiza o índice para o final
                self.current_index = len(self.prints)
             
                # Salva o documento imediatamente
                self.salvar_docx()
                
                # Fecha o popup após salvar
                popup.destroy()
                
                resultado = "ja_salvou"  # Retorna um valor especial para indicar que já salvou
            else:
                resultado = False

        def excluir_print():
            nonlocal resultado
            if messagebox.askyesno("Confirmar Exclusão", "Tem certeza que deseja excluir esta evidência?"):
                # Remove o arquivo
                try:
                    nome_arquivo = os.path.basename(caminho_print)
                    
                    # Marca como excluída nos metadados em vez de remover fisicamente
                    for evidencia in self.metadata["evidencias"]:
                        if evidencia["arquivo"] == nome_arquivo:
                            evidencia["excluida"] = True
                            break
                    
                    self._salvar_metadata()
                    
                    # Remove fisicamente o arquivo
                    os.remove(caminho_print)
                    print(f"Arquivo excluído: {caminho_print}")
                    
                    resultado = None  # Indica que houve exclusão
                    popup.destroy()
                    
                except Exception as e:
                    print(f"Erro ao excluir arquivo: {e}")
                    messagebox.showerror("Erro", f"Não foi possível excluir o arquivo: {e}")
                    resultado = False
                    popup.destroy()
                    return
            else:
                resultado = False  # Usuário cancelou a exclusão

        # Frame para botões de ação
        acoes_frame = tk.Frame(popup)
        acoes_frame.pack(pady=10)

        tk.Button(acoes_frame, text="✏ Editar Print", command=editar_print, width=15).pack(side=tk.LEFT, padx=5)
        tk.Button(acoes_frame, text="Adicionar e Próximo", command=adicionar, width=15).pack(side=tk.LEFT, padx=5)
        tk.Button(acoes_frame, text="🗑️ Excluir Print", command=excluir_print, width=15).pack(side=tk.LEFT, padx=5)

        # Frame para botões de controle
        controle_frame = tk.Frame(popup)
        controle_frame.pack(pady=10)

        tk.Button(controle_frame, text="❌ Cancelar", command=cancelar_processamento, 
                  bg="#ff6b6b", fg="white", width=15).pack(side=tk.LEFT, padx=5)
        tk.Button(controle_frame, text="✅ Incluir Todos", command=incluir_todos, 
                  bg="#4ecdc4", fg="white", width=15).pack(side=tk.LEFT, padx=5)

        def on_closing():
            cancelar_processamento()

        popup.protocol("WM_DELETE_WINDOW", on_closing)
        popup.grab_set()
        root.wait_window(popup)
        
        if self.processamento_cancelado:
            return False
        
        return resultado

    def salvar_docx(self):
        if self.template_path:
            # Usa o nome do template sem o prefixo "Evidencias_"
            nome_base = os.path.basename(self.template_path)
            # Remove a extensão .docx se existir
            if nome_base.lower().endswith('.docx'):
                nome_base = nome_base[:-5]
            nome_arquivo = f"{nome_base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        else:
            # Nome simples com timestamp
            nome_arquivo = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        caminho_save = os.path.join(self.output_dir, nome_arquivo)
        
        try:
            self.doc.save(caminho_save)
            self.saved_file_path = caminho_save
            
            # Função para abrir a pasta (será chamada após fechar o messagebox)
            def abrir_posta_apos_mensagem():
                if os.name == 'nt':
                    os.startfile(self.output_dir)
                elif os.name == 'posix':
                    import subprocess
                    if sys.platform == 'darwin':
                        subprocess.Popen(['open', self.output_dir])
                    else:
                        subprocess.Popen(['xdg-open', self.output_dir])
            
            # Mostra a mensagem e agenda a abertura da pasta para depois
            messagebox.showinfo("Concluído", f"Documento gerado com sucesso!\nSalvo em:\n{caminho_save}")
            
            # Agenda a abertura da pasta para depois de fechar the messagebox
            root.after(100, abrir_posta_apos_mensagem)
                
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao salvar documento: {str(e)}")

    # ---------- Editor de prints ----------
    # (Mantido igual - não houve mudanças no editor)
    def abrir_editor(self, caminho_print, parent):
        editor = tk.Toplevel(parent)
        editor.title("Editor de Evidência")
        editor.geometry("1200x800")
        
        # Frame principal
        main_frame = tk.Frame(editor)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Frame para ferramentas e opções
        tools_frame = tk.Frame(main_frame)
        tools_frame.pack(side=tk.TOP, fill=tk.X, pady=5)
        
        # Frame para a área de desenho
        canvas_frame = tk.Frame(main_frame)
        canvas_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        # Carrega a imagem original
        self.original_img = Image.open(caminho_print).convert("RGBA")
        img_w, img_h = self.original_img.size
        
        # Calcula o fator de escala para exibição
        max_w, max_h = 1000, 700
        scale = min(max_w / img_w, max_h / img_h)
        self.scale_factor = scale
        disp_w, disp_h = int(img_w * scale), int(img_h * scale)
        
        # Cria cópia da imagem para edição
        self.editing_img = self.original_img.copy()
        self.display_img = self.editing_img.resize((disp_w, disp_h), Image.LANCZOS)

        # Variáveis para controle
        self.current_tk_img = ImageTk.PhotoImage(self.display_img)
        self.elements = []  # Lista de elementos desenhados
        self.undo_stack = []  # PILHA PARA DESFAZER AÇÕES
        self.temp_element = None
        
        # Canvas para a imagem
        self.canvas = tk.Canvas(canvas_frame, width=disp_w, height=disp_h, cursor="cross", bg="gray")
        self.canvas.pack(padx=5, pady=5)
        self.canvas_img = self.canvas.create_image(0, 0, anchor="nw", image=self.current_tk_img)
        
        # Variáveis de controle - COR PADRÃO VERMELHA
        tool_var = tk.StringVar(value="rectangle")  # RETÂNGULO COMO PADRÃO
        color_var = tk.StringVar(value="#FF0000")   # VERMELHO COMO PADRÃO
        width_var = tk.IntVar(value=3)
        
        # Ferramentas - SUBSTITUINDO RADIOBUTTONS POR ÍCONES EMOJI
        tk.Label(tools_frame, text="Ferramenta:").pack(side=tk.LEFT, padx=5)
        
        # Frame para os botões de ícone
        icon_frame = tk.Frame(tools_frame)
        icon_frame.pack(side=tk.LEFT, padx=5)
        
        # Ícones emoji para cada ferramenta
        tool_icons = {
            "rectangle": "⬜",   # Retângulo
            "circle": "🔴",      # Círculo  
            "arrow": "👉",       # Seta - Mão apontando
            "text": "🆎"         # Texto - Botão AB
        }

        # Função para criar botões com estilo consistente
        def criar_botao_ferramenta(parent, texto, valor, variavel):
            btn = tk.Radiobutton(parent, text=texto, font=("Arial", 12), 
                               variable=variavel, value=valor, indicatoron=0, 
                               width=3, height=2, relief=tk.RAISED,
                               cursor="hand2")
            return btn

        # Cria os botões para cada ferramenta
        for tool_value, icon in tool_icons.items():
            btn = criar_botao_ferramenta(icon_frame, icon, tool_value, tool_var)
            btn.pack(side=tk.LEFT, padx=2)

        # Destacar o botão do retângulo (selecionado por padrão)
        for widget in icon_frame.winfo_children():
            if isinstance(widget, tk.Radiobutton) and widget.cget("value") == "rectangle":
                widget.config(relief=tk.SUNKEN, bg="#e3f2fd")  # Azul claro para selecionado
                break

        # Função para atualizar a apariência dos botões
        def update_button_appearance(*args):
            selected_tool = tool_var.get()
            for widget in icon_frame.winfo_children():
                if isinstance(widget, tk.Radiobutton):
                    if widget.cget("value") == selected_tool:
                        widget.config(relief=tk.SUNKEN, bg="#e3f2fd")  # Selecionado
                    else:
                        widget.config(relief=tk.RAISED, bg="SystemButtonFace")  # Normal

        tool_var.trace("w", update_button_appearance)
        
        # Controles de cor and espessura - APENAS CORES ESSENCIAIS
        color_frame = tk.Frame(tools_frame)
        color_frame.pack(side=tk.LEFT, padx=20)
        
        tk.Label(color_frame, text="Cor:").pack(side=tk.LEFT)
        
        # Paleta de cores reduzida (apenas as essenciais)
        colors = ["#FF0000", "#00FF00", "#FFFF00", "#000000", "#FFFFFF"]
        color_buttons_frame = tk.Frame(color_frame)
        color_buttons_frame.pack(side=tk.LEFT, padx=5)
        
        for color in colors:
            btn = tk.Button(color_buttons_frame, bg=color, width=2, height=1, 
                           command=lambda c=color: self.set_color(color_var, c, color_preview))
            btn.pack(side=tk.LEFT, padx=1)
        
        # Botão para cor personalizada
        custom_btn = tk.Button(color_frame, text="Personalizada", 
                              command=lambda: self.choose_custom_color(editor, color_var, color_preview))
        custom_btn.pack(side=tk.LEFT, padx=5)
        
        # Preview de cor
        color_preview = tk.Frame(color_frame, width=30, height=20, bg=color_var.get())
        color_preview.pack(side=tk.LEFT, padx=5)
        
        # Controle de espessura
        width_frame = tk.Frame(tools_frame)
        width_frame.pack(side=tk.LEFT, padx=20)
        
        tk.Label(width_frame, text="Espessura:").pack(side=tk.LEFT)
        tk.Scale(width_frame, from_=1, to=10, variable=width_var, orient=tk.HORIZONTAL, 
                length=100, showvalue=1).pack(side=tk.LEFT, padx=5)
        
        # BOTÃO DESFAZER
        def undo_action():
            if self.elements:  # Se houver elementos para desfazer
                # Remove o último elemento and adiciona à pilha de desfazer
                removed_element = self.elements.pop()
                self.undo_stack.append(removed_element)
                refresh_display()
        
        undo_btn = tk.Button(tools_frame, text="↩️ Desfazer (Ctrl+Z)", command=undo_action)
        undo_btn.pack(side=tk.LEFT, padx=20)
        
        # Variáveis para desenho
        start_xy = {"x": None, "y": None}
        
        def refresh_display():
            # Redesenha todos os elementos
            self.canvas.delete("all")
            self.canvas.create_image(0, 0, anchor="nw", image=self.current_tk_img)
            
            for element in self.elements:
                elem_type, coords, color, width, text = element
                scaled_coords = [int(c * self.scale_factor) for c in coords]
                
                if elem_type == "circle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_oval(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "rectangle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "arrow":
                    x1, y1, x2, y2 = scaled_coords
                    self.draw_arrow_on_canvas(x1, y1, x2, y2, color, width)
                elif elem_type == "text":
                    x, y = scaled_coords
                    self.canvas.create_text(x, y, text=text, fill=color, font=("Arial", 12), anchor="nw")
            
            # Desenha elemento temporário durante a criação
            if self.temp_element:
                elem_type, coords, color, width, text = self.temp_element
                scaled_coords = [int(c * self.scale_factor) for c in coords]
                
                if elem_type == "circle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_oval(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "rectangle":
                    x1, y1, x2, y2 = scaled_coords
                    self.canvas.create_rectangle(x1, y1, x2, y2, outline=color, width=width)
                elif elem_type == "arrow":
                    x1, y1, x2, y2 = scaled_coords
                    self.draw_arrow_on_canvas(x1, y1, x2, y2, color, width)
        
        def draw_arrow_on_canvas(x1, y1, x2, y2, color, width):
            # Desenha a linha da seta
            self.canvas.create_line(x1, y1, x2, y2, fill=color, width=width)
            
            # Calcula o ângulo da seta
            angle = math.atan2(y2 - y1, x2 - x1)
            
            # Desenha a ponta da seta (triângulo)
            arrow_size = 15
            x3 = x2 - arrow_size * math.cos(angle - math.pi/6)
            y3 = y2 - arrow_size * math.sin(angle - math.pi/6)
            x4 = x2 - arrow_size * math.cos(angle + math.pi/6)
            y4 = y2 - arrow_size * math.sin(angle + math.pi/6)
            
            self.canvas.create_polygon(x2, y2, x3, y3, x4, y4, fill=color, outline=color)
        
        def on_button_press(event):
            start_xy["x"], start_xy["y"] = event.x, event.y
        
        def on_motion(event):
            if start_xy["x"] is not None:
                # Desenho em tempo real
                sx, sy = start_xy["x"], start_xy["y"]
                ex, ey = event.x, event.y
                
                # Converte para coordenadas da imagem original
                ix1, iy1 = int(sx / self.scale_factor), int(sy / self.scale_factor)
                ix2, iy2 = int(ex / self.scale_factor), int(ey / self.scale_factor)
                
                tool = tool_var.get()
                color = color_var.get()
                width = width_var.get()
                
                if tool == "circle":
                    radius = int(((ix2 - ix1)**2 + (iy2 - iy1)**2)**0.5)
                    self.temp_element = ("circle", [ix1-radius, iy1-radius, ix1+radius, iy1+radius], color, width, "")
                elif tool == "rectangle":
                    # Garante que x2 >= x1 and y2 >= y1
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.temp_element = ("rectangle", [x1_norm, y1_norm, x2_norm, y2_norm], color, width, "")
                elif tool == "arrow":
                    self.temp_element = ("arrow", [ix1, iy1, ix2, iy2], color, width, "")
                
                refresh_display()
        
        def on_button_release(event):
            if start_xy["x"] is not None:
                # Converte coordenadas da tela para coordenadas da imagem original
                sx, sy = start_xy["x"], start_xy["y"]
                ex, ey = event.x, event.y
                
                ix1, iy1 = int(sx / self.scale_factor), int(sy / self.scale_factor)
                ix2, iy2 = int(ex / self.scale_factor), int(ey / self.scale_factor)
                
                tool = tool_var.get()
                color = color_var.get()
                width = width_var.get()
                
                # Limpa a pilha de desfazer quando uma nova ação é realizada
                self.undo_stack.clear()
                
                if tool == "circle":
                    radius = int(((ix2 - ix1)**2 + (iy2 - iy1)**2)**0.5)
                    self.elements.append(("circle", [ix1-radius, iy1-radius, ix1+radius, iy1+radius], color, width, ""))
                
                elif tool == "rectangle":
                    # Garante que x2 >= x1 and y2 >= y1
                    x1_norm = min(ix1, ix2)
                    y1_norm = min(iy1, iy2)
                    x2_norm = max(ix1, ix2)
                    y2_norm = max(iy1, iy2)
                    self.elements.append(("rectangle", [x1_norm, y1_norm, x2_norm, y2_norm], color, width, ""))
                
                elif tool == "arrow":
                    self.elements.append(("arrow", [ix1, iy1, ix2, iy2], color, width, ""))
                
                elif tool == "text":
                    # Para texto, pede o conteúdo e adiciona na posição clicada
                    text = simpledialog.askstring("Texto", "Digite o texto:", parent=editor)
                    if text:
                        self.elements.append(("text", [ix1, iy1], color, width, text))
                        # Atualiza a visualização para mostrar o texto imediatamente
                        refresh_display()
                
                self.temp_element = None
                refresh_display()
            
            start_xy["x"], start_xy["y"] = None, None
        
        # BIND DO CTRL+Z (atalho global dentro do editor)
        def on_key_press(event):
            undo_action()

        editor.bind_all('<Control-z>', on_key_press)
        editor.bind_all('<Control-Z>', on_key_press)
        
        # Bind events para o canvas
        self.canvas.bind("<Button-1>", on_button_press)
        self.canvas.bind("<B1-Motion>", on_motion)
        self.canvas.bind("<ButtonRelease-1>", on_button_release)
        
        # Para ferramenta de texto, muda o cursor
        def update_cursor(*args):
            if tool_var.get() == "text":
                self.canvas.config(cursor="xterm")
            else:
                self.canvas.config(cursor="cross")
        
        tool_var.trace("w", update_cursor)
        
        # Botões de ação
        action_frame = tk.Frame(editor)
        action_frame.pack(side=tk.BOTTOM, pady=10)
        
        def salvar_imagem():
            # Cria uma cópia da imagem original para aplicar as anotações
            final_img = self.original_img.copy()
            draw = ImageDraw.Draw(final_img)
            
            # Aplica todas as anotações
            for element in self.elements:
                elem_type, coords, color, width, text = element
                
                if elem_type == "circle":
                    x1, y1, x2, y2 = coords
                    draw.ellipse([x1, y1, x2, y2], outline=color, width=width)
                
                elif elem_type == "rectangle":
                    x1, y1, x2, y2 = coords
                    draw.rectangle([x1, y1, x2, y2], outline=color, width=width)
                
                elif elem_type == "arrow":
                    x1, y1, x2, y2 = coords
                    # Desenha a linha
                    draw.line([x1, y1, x2, y2], fill=color, width=width)
                    
                    # Desenha a ponta da seta (simplificado)
                    angle = math.atan2(y2 - y1, x2 - x1)
                    arrow_size = 15
                    x3 = x2 - arrow_size * math.cos(angle - math.pi/6)
                    y3 = y2 - arrow_size * math.sin(angle - math.pi/6)
                    x4 = x2 - arrow_size * math.cos(angle + math.pi/6)
                    y4 = y2 - arrow_size * math.sin(angle + math.pi/6)
                    
                    draw.polygon([x2, y2, x3, y3, x4, y4], fill=color, outline=color)
                
                elif elem_type == "text":
                    x, y = coords
                    # Usa uma fonte padrão
                    try:
                        font = ImageFont.truetype("arial.ttf", 20)
                    except:
                        font = ImageFont.load_default()
                    draw.text((x, y), text, fill=color, font=font)
            
            # Salva a imagem editada (sobrescreve a original)
            final_img.save(caminho_print)
            messagebox.showinfo("Sucesso", "Imagem salva com sucesso!")
            editor.destroy()
            parent.destroy()  # Fecha também a janela de comentário
        
        def cancelar_edicao():
            if messagebox.askyesno("Confirmar", "Descartar todas as alterações?"):
                editor.destroy()
        
        tk.Button(action_frame, text="💾 Salvar", command=salvar_imagem, width=15).pack(side=tk.LEFT, padx=10)
        tk.Button(action_frame, text="❌ Cancelar", command=cancelar_edicao, width=15).pack(side=tk.LEFT, padx=10)

    def set_color(self, color_var, color, preview_widget):
        color_var.set(color)
        preview_widget.config(bg=color)

    def choose_custom_color(self, parent, color_var, preview_widget):
        color = colorchooser.askcolor(initialcolor=color_var.get(), parent=parent)[1]
        if color:
            color_var.set(color)
            preview_widget.config(bg=color)

# ------------------ MAIN ------------------
if __name__ == "__main__":
    root = tk.Tk()
    root.title("PrintF - Gerador de Evidências")
    root.geometry("400x200")
    root.resizable(False, False)
    
    # Centraliza a janela
    root.eval('tk::PlaceWindow . center')
    
    # Frame principal
    main_frame = ttk.Frame(root, padding=30)
    main_frame.pack(fill=tk.BOTH, expand=True)
    
    # Título
    ttk.Label(main_frame, text="PrintF - Gerador de Evidências", 
             font=("Arial", 16, "bold")).pack(pady=20)
    
    # Botão para iniciar
    def iniciar_gerador():
        gerador = GeradorDocx()
        if gerador.mostrar_janela_configuracao():
            # O processamento continua automaticamente após a configuração
            pass
    
    ttk.Button(main_frame, text="Iniciar Gerador de Evidências", 
              command=iniciar_gerador, width=25).pack(pady=10)
    
    # Botão para sair
    ttk.Button(main_frame, text="Sair", command=root.quit, width=15).pack(pady=10)
    
    root.mainloop()