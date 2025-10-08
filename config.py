import os
from pathlib import Path
import json

class AppConfig:
    """Configurações globais completas do PrintF Unificado"""
    
    # Versão
    VERSION = "1.0.0"
    APP_NAME = "PrintF"
    
    # Diretórios
    BASE_DIR = Path(os.path.dirname(os.path.abspath(__file__)))
    OUTPUT_DIR = BASE_DIR / "Output"
    TEMPLATES_DIR = BASE_DIR / "Templates"
    CONFIG_DIR = BASE_DIR / "Config"
    LOGS_DIR = BASE_DIR / "Logs"
    
    # Configurações de tema
    THEME = 'liquid_glass'  # 'default' ou 'liquid_glass'
    
    # Cores do tema padrão (fallback)
    COLORS = {
        'primary': "#2c3e50",
        'secondary': "#3498db", 
        'success': "#27ae60",
        'warning': "#f39c12",
        'danger': "#e74c3c",
        'dark': "#34495e",
        'light': "#ecf0f1",
        'gray': "#95a5a6"
    }
    
    # Configurações de captura
    CAPTURE_SETTINGS = {
        'default_mode': 'ocultar',  # ou 'manter'
        'keep_evidence_files': True,
        'timestamp_position': {'x': 0.75, 'y': 0.90},
        'timestamp_color': '#FFFFFF',
        'timestamp_background': '#000000B2',
        'timestamp_size': 24,
        'click_marker_radius': 20,
        'click_marker_color': (255, 255, 0, 100)
    }
    
    # Configurações de documentos
    DOCUMENT_SETTINGS = {
        'default_template': 'template_evidencias.docx',
        'image_width_inches': 6.0,
        'auto_open_folder': True,
        'backup_metadata': True
    }
    
    # Atalhos de teclado
    HOTKEYS = {
        'start_recording': 'f8',
        'pause_resume': 'f6', 
        'stop_recording': 'f9',
        'close_app': 'f12',
        'manual_capture': 'f10'
    }
    
    # Extensões suportadas
    FILE_EXTENSIONS = {
        'images': ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif'],
        'documents': ['.docx', '.pdf'],
        'data': ['.csv', '.json', '.xml']
    }
    
    # Configurações de UI responsiva
    UI_SETTINGS = {
        'min_width': 800,
        'min_height': 600,
        'default_width': 1000,
        'default_height': 700,
        'responsive_breakpoints': {
            'small': 900,
            'medium': 1200,
            'large': 1600
        }
    }
    
    @classmethod
    def setup_directories(cls):
        """Cria todos os diretórios necessários"""
        directories = [
            cls.OUTPUT_DIR, cls.TEMPLATES_DIR, 
            cls.CONFIG_DIR, cls.LOGS_DIR
        ]
        
        for directory in directories:
            directory.mkdir(exist_ok=True)
        
        # Criar template padrão se não existir
        cls._create_default_template()
    
    @classmethod
    def _create_default_template(cls):
        """Cria template padrão se não existir"""
        default_template = cls.TEMPLATES_DIR / "template_evidencias.docx"
        if not default_template.exists():
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Título
                title = doc.add_heading('Evidências de Teste - Documentação', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Informações do projeto
                doc.add_heading('Informações do Projeto', level=1)
                doc.add_paragraph("Projeto: [NOME_DO_PROJETO]")
                doc.add_paragraph("Módulo: [MÓDULO]")
                doc.add_paragraph("Versão: [VERSÃO]")
                doc.add_paragraph("Responsável: [RESPONSÁVEL]")
                
                doc.add_paragraph()
                
                # Caso de teste
                doc.add_heading('Caso de Teste', level=1)
                doc.add_paragraph("Caso de Teste: [NOME_DO_CASO]")
                
                doc.add_paragraph()
                
                # Seção de evidências
                doc.add_heading('Evidências Coletadas', level=1)
                doc.add_paragraph("As evidências abaixo foram capturadas durante a execução do teste:")
                
                doc.save(str(default_template))
                print(f"✅ Template padrão criado: {default_template}")
                
            except Exception as e:
                print(f"⚠️ Não foi possível criar template padrão: {e}")
    
    @classmethod
    def load_user_settings(cls):
        """Carrega configurações do usuário"""
        settings_file = cls.CONFIG_DIR / "user_settings.json"
        
        if settings_file.exists():
            try:
                with open(settings_file, 'r', encoding='utf-8') as f:
                    user_settings = json.load(f)
                    # Garantir que o tema liquid_glass esteja habilitado por padrão
                    if 'theme' not in user_settings:
                        user_settings['theme'] = 'liquid_glass'
                    # Garantir compatibilidade com versões anteriores
                    user_settings.setdefault('responsive_layout', True)
                    return user_settings
            except Exception as e:
                print(f"⚠️ Erro ao carregar configurações: {e}")
                # Retornar configurações padrão em caso de erro
        
        # Configurações padrão com liquid_glass habilitado
        return {
            'recent_folders': [],
            'window_size': {'width': cls.UI_SETTINGS['default_width'], 'height': cls.UI_SETTINGS['default_height']},
            'window_position': {'x': 100, 'y': 100},
            'theme': 'liquid_glass',  # HABILITADO POR PADRÃO
            'language': 'pt-BR',
            'auto_save': True,
            'check_updates': True,
            'responsive_layout': True
        }
    
    @classmethod
    def save_user_settings(cls, settings):
        """Salva configurações do usuário"""
        settings_file = cls.CONFIG_DIR / "user_settings.json"
        
        try:
            # Garantir que o diretório existe
            cls.CONFIG_DIR.mkdir(exist_ok=True)
            
            with open(settings_file, 'w', encoding='utf-8') as f:
                json.dump(settings, f, indent=2, ensure_ascii=False)
            print("✅ Configurações salvas com sucesso!")
            return True
        except Exception as e:
            print(f"❌ Erro ao salvar configurações: {e}")
            return False

# Instância global de configuração
APP_CONFIG = AppConfig()

# Inicializar diretórios na importação
APP_CONFIG.setup_directories()